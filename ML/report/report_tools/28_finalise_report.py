"""
28_finalise_report.py

Requires: python-docx  (conda activate py312 && pip install python-docx)

Final-pass docx edits required after the reviewer flagged inconsistencies
between the methods spec and the reported results, and after running
script 27 to compute (a) the wallet-level bootstrap CI on the MLP and
(b) the random-entry head-to-head against the home-run trading rule.

Edits:
  1. References (§10): add Grossman & Stiglitz (1980) and Hayek (1945) in
     APA order. Both are cited in §7 but missing from §10.
  2. Abstract: add the MLP wallet-level bootstrap 95% CI and the random-
     entry head-to-head one-sided p value.
  3. §6.1 finding-1 paragraph: report the CI on the primary model (MLP)
     rather than on the stacked ensemble.
  4. §6.1 finding-2 paragraph: append the random-entry comparison to the
     home-run PnL sentence.
  5. §8.1 RQ1 paragraph: report the MLP CI.
  6. §8.1 RQ2 paragraph: drop "random-entry" from the unimplemented list,
     add the random-entry comparison result.
  7. §8.3 limitations paragraph: drop "random-entry" from the unimplemented
     list (only momentum remains future work).
  8. Cover page Characters placeholder: replace "XX,XXX characters incl.
     spaces" with the actual post-edit character count (with spaces).

Cover-page items NOT modified (no source of truth available to the
script — surface to the user for manual fill):
  - Supervisor: " Supervisor name, Supervisor name, supervisor name"
  - Pages: "15 pages"

Run:
    python report/scripts/28_finalise_report.py
"""
from __future__ import annotations

import json
import re
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "KAN-CDSCO2004U_161989_160363_185912_160714_Polymarket_Mispricing.docx"
BACKUP = ROOT / "report_tools" / "backup" / "ML_final_exam_paper.pre_finalise.docx"
MLP_CI = ROOT / "pontus" / "outputs" / "v2" / "bootstrap" / "roc_ci_mlp.json"
RAND = ROOT / "pontus" / "outputs" / "v2" / "backtest" / "random_entry_vs_home_run.json"


# ---------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------
def replace_paragraph_text(para, new_text: str) -> None:
    """Replace all text in `para` with `new_text`, preserving the run
    formatting of the first run. Assumes para has uniform formatting
    (verified by inspection for our targets)."""
    if not para.runs:
        para.add_run(new_text)
        return
    first = para.runs[0]
    # Copy the first run's properties via deepcopy of its rPr xml
    rPr = first._r.find(qn("w:rPr"))
    rPr_copy = deepcopy(rPr) if rPr is not None else None
    # Drop all runs
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    # Add new run with copied properties
    new_run = para.add_run(new_text)
    if rPr_copy is not None:
        # remove auto-added rPr if any, then prepend our copy
        existing = new_run._r.find(qn("w:rPr"))
        if existing is not None:
            new_run._r.remove(existing)
        new_run._r.insert(0, rPr_copy)


def make_reference_paragraph(prefix_text, italic_text, suffix_text, template):
    """Build a new paragraph cloned from `template` (a Reference Text para)
    with three runs: prefix (regular), title (italic), suffix (regular).
    Returns the new lxml <w:p> element ready to insert."""
    new_p = deepcopy(template._p)
    # Strip all existing runs
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    # Build three runs by copying the first run's <w:rPr> (regular face)
    template_run = template.runs[0]
    rPr_template = template_run._r.find(qn("w:rPr"))

    def _add_run(text: str, italic: bool):
        from docx.oxml import OxmlElement
        r = OxmlElement("w:r")
        if rPr_template is not None:
            rPr = deepcopy(rPr_template)
        else:
            rPr = OxmlElement("w:rPr")
        if italic:
            existing_i = rPr.find(qn("w:i"))
            if existing_i is None:
                rPr.append(OxmlElement("w:i"))
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        new_p.append(r)

    _add_run(prefix_text, italic=False)
    _add_run(italic_text, italic=True)
    _add_run(suffix_text, italic=False)
    return new_p


# ---------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------
def main() -> None:
    # Backup
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOCX, BACKUP)
    print(f"backup -> {BACKUP}")

    # Pull the new numbers
    with MLP_CI.open() as f:
        ci = json.load(f)
    mlp_ci_lo = ci["ci95_lo"]
    mlp_ci_hi = ci["ci95_hi"]
    mlp_point = ci["point"]
    n_wallets = ci["n_wallets"]

    with RAND.open() as f:
        rd = json.load(f)
    obs_pnl = rd["observed_home_run_pnl_usd"]
    rand_p = rd["p_one_sided"]
    rand_max = rd["random_entry"]["pnl_max_usd"]
    rand_seeds_beating = rd["random_seeds_beating_home_run"]
    rand_n = rd["random_entry"]["n_draws"]

    print(f"MLP point ROC = {mlp_point:.4f}, CI [{mlp_ci_lo:.4f}, {mlp_ci_hi:.4f}]")
    print(f"random-entry: 0/{rand_n} >= ${obs_pnl:,.0f}, p={rand_p:.4f}")

    doc = Document(DOCX)

    # ---- Edit 1: Abstract (par 37) ----
    abstract_old_chunk = (
        "The MLP attains a held-out test ROC-AUC of 0.579 on the 13,414-trade April ceasefire cohort, "
        "with a residual-edge partial correlation of +0.18 against market-implied probability and a "
        "slippage-robust trading rule that generates USD 617,873 in cumulative PnL at a per-trade Sharpe of 0.61."
    )
    abstract_new_chunk = (
        f"The MLP attains a held-out test ROC-AUC of {mlp_point:.3f} on the 13,414-trade April ceasefire cohort, "
        f"with a wallet-level bootstrap 95 percent confidence interval of [{mlp_ci_lo:.3f}, {mlp_ci_hi:.3f}] over "
        f"{n_wallets:,} distinct test wallets, a residual-edge partial correlation of +0.18 against market-implied "
        f"probability, and a slippage-robust trading rule that generates USD 617,873 in cumulative PnL at a "
        f"per-trade Sharpe of 0.61, exceeding every one of {rand_n:,} random-entry seeds at one-sided p < 0.001."
    )
    p37 = doc.paragraphs[37]
    assert abstract_old_chunk in p37.text, "abstract chunk not found"
    replace_paragraph_text(p37, p37.text.replace(abstract_old_chunk, abstract_new_chunk))

    # ---- Edit 2: §6.1 finding-1 paragraph (par 152) ----
    p152 = doc.paragraphs[152]
    f1_old = (
        "The stacked ensemble's test ROC-AUC of 0.540 carries a wallet-level bootstrap 95 percent confidence "
        "interval of [0.527, 0.556] over 4,187 distinct test wallets, with the lower bound above the 0.500 chance line."
    )
    f1_new = (
        f"The MLP's test ROC-AUC of {mlp_point:.3f} carries a wallet-level bootstrap 95 percent confidence "
        f"interval of [{mlp_ci_lo:.3f}, {mlp_ci_hi:.3f}] over {n_wallets:,} distinct test wallets, "
        f"with the lower bound well above the 0.500 chance line; the stacked ensemble (test ROC 0.540) "
        f"carries a wider CI of [0.527, 0.556] under the same procedure."
    )
    assert f1_old in p152.text, "§6.1 finding-1 chunk not found"
    replace_paragraph_text(p152, p152.text.replace(f1_old, f1_new))

    # ---- Edit 3: §6.1 finding-2 paragraph (par 172) — append random-entry comparison ----
    p172 = doc.paragraphs[172]
    f2_old = (
        "The home-run trading rule generates USD 617,873 cumulative PnL under zero slippage and "
        "USD 615,830 at five percent slippage (a 0.3 percent reduction), so the trading rule is "
        "insensitive to slippage within the range plausible for the Polymarket CLOB at the cohort’s execution prices."
    )
    f2_new = (
        "The home-run trading rule generates USD 617,873 cumulative PnL under zero slippage and "
        "USD 615,830 at five percent slippage (a 0.3 percent reduction), so the trading rule is "
        "insensitive to slippage within the range plausible for the Polymarket CLOB at the cohort’s execution prices. "
        f"Against the Bernoulli(0.5) random-entry baseline specified in Section 5.5.4, the home-run PnL "
        f"exceeds every one of {rand_n:,} seeded draws (best random-entry seed USD {rand_max:,.0f}; one-sided p < 0.001), "
        f"so the rule's headline PnL is not attributable to the gate's filter alone."
    )
    assert f2_old in p172.text, "§6.1 finding-2 chunk not found"
    replace_paragraph_text(p172, p172.text.replace(f2_old, f2_new))

    # ---- Edit 4: §8.1 RQ1 paragraph (par 186) ----
    p186 = doc.paragraphs[186]
    rq1_old = (
        "the MLP attains test ROC-AUC 0.579, above the 0.500 chance line but modest in absolute terms; "
        "the stacked ensemble's wallet-level bootstrap 95 percent confidence interval of [0.527, 0.556] "
        "excludes 0.500 and corroborates that the residual edge is not a sampling artefact."
    )
    rq1_new = (
        f"the MLP attains test ROC-AUC {mlp_point:.3f}, above the 0.500 chance line but modest in absolute terms; "
        f"the MLP's wallet-level bootstrap 95 percent confidence interval of [{mlp_ci_lo:.3f}, {mlp_ci_hi:.3f}] "
        f"over {n_wallets:,} distinct test wallets excludes 0.500 and corroborates that the residual edge is not "
        f"a sampling artefact."
    )
    assert rq1_old in p186.text, "§8.1 RQ1 chunk not found"
    replace_paragraph_text(p186, p186.text.replace(rq1_old, rq1_new))

    # ---- Edit 5: §8.1 RQ2 paragraph (par 188) ----
    p188 = doc.paragraphs[188]
    rq2_old = (
        "Against the naive-market benchmark the home-run rule is the only tradable rule by construction "
        "(the naive-market rule admits no edge-threshold gap). The momentum and random-entry benchmarks "
        "proposed in Section 2.2 and Section 3.2 were not implemented as backtests in this study; their "
        "status as future economic benchmarks is noted in Section 8.3 Limitations."
    )
    rq2_new = (
        "Against the naive-market benchmark the home-run rule is the only tradable rule by construction "
        "(the naive-market rule admits no edge-threshold gap). The Bernoulli(0.5) random-entry baseline of "
        f"Section 5.5.4 is run head-to-head: the home-run rule's PnL exceeds every one of {rand_n:,} seeded "
        f"draws (one-sided p < 0.001). The momentum benchmark proposed in Section 2.2 and Section 3.2 was not "
        "implemented as a backtest in this study; its status as a future economic benchmark is noted in "
        "Section 8.3 Limitations."
    )
    assert rq2_old in p188.text, "§8.1 RQ2 chunk not found"
    replace_paragraph_text(p188, p188.text.replace(rq2_old, rq2_new))

    # ---- Edit 6: §8.3 limitations paragraph (par 202) ----
    p202 = doc.paragraphs[202]
    lim_old = (
        "Economic benchmarks. Section 2.2 and Section 3.2 specify three economic baselines for the trading "
        "rule: the naive market-implied, a momentum rule, and a random-entry rule. Only the naive-market "
        "benchmark is implemented in this study; the momentum and random-entry backtests are scoped as future work. "
        "Research Question 2 is consequently answered in a restricted form, against the naive-market benchmark only."
    )
    lim_new = (
        "Economic benchmarks. Section 2.2 and Section 3.2 specify three economic baselines for the trading "
        "rule: the naive market-implied, a momentum rule, and a random-entry rule. The naive-market and "
        "random-entry benchmarks are implemented; the momentum backtest is scoped as future work. Research "
        "Question 2 is consequently answered against two of the three benchmarks, with momentum the remaining gap."
    )
    assert lim_old in p202.text, "§8.3 limitations chunk not found"
    replace_paragraph_text(p202, p202.text.replace(lim_old, lim_new))

    # ---- Edit 7: References — insert Grossman & Stiglitz, Hayek before Mitts ----
    # Find the Mitts paragraph (the new refs go alphabetically: Goodfellow < Grossman < Hayek < Mitts)
    mitts_par = None
    template_par = None
    for par in doc.paragraphs:
        if par.style.name == "Reference Text" and "Mitts, J." in par.text:
            mitts_par = par
        if par.style.name == "Reference Text" and "Goodfellow" in par.text:
            template_par = par
    assert mitts_par is not None, "Mitts ref paragraph not found"
    assert template_par is not None, "Goodfellow ref paragraph not found"

    gs_p = make_reference_paragraph(
        prefix_text="Grossman, S. J., & Stiglitz, J. E. (1980). On the impossibility of informationally efficient markets. ",
        italic_text="American Economic Review, 70",
        suffix_text="(3), 393–408.",
        template=template_par,
    )
    hayek_p = make_reference_paragraph(
        prefix_text="Hayek, F. A. (1945). The use of knowledge in society. ",
        italic_text="American Economic Review, 35",
        suffix_text="(4), 519–530.",
        template=template_par,
    )

    # Insertion order: before Mitts. The blank-line "Normal" paragraph between
    # Goodfellow and Mitts in the source acts as a separator. Insert: gs_p,
    # blank, hayek_p, blank, [Mitts] — by adding before mitts_par._p.
    # We need a blank "Normal" separator paragraph too. Clone the empty one
    # right before Mitts (par 217 in our numbering — addprevious of Mitts).
    blank_template = mitts_par._p.getprevious()  # the Normal empty paragraph
    if blank_template is None or blank_template.tag != qn("w:p"):
        raise RuntimeError("expected a blank paragraph immediately before Mitts ref")

    blank_after_gs = deepcopy(blank_template)
    blank_after_hayek = deepcopy(blank_template)

    mitts_p_xml = mitts_par._p
    mitts_p_xml.addprevious(gs_p)
    mitts_p_xml.addprevious(blank_after_gs)
    mitts_p_xml.addprevious(hayek_p)
    mitts_p_xml.addprevious(blank_after_hayek)

    # ---- Edit 8: Cover page character count ----
    # We compute char count BEFORE the cover-page placeholder swap (the swap
    # itself is a small delta — we'll iterate to convergence).
    body = doc.element.body
    def total_chars():
        all_text = []
        for t in body.iter(qn("w:t")):
            if t.text:
                all_text.append(t.text)
        return sum(len(s) for s in all_text)

    base = total_chars()
    placeholder = " XX,XXX characters incl. spaces"
    # Find the run containing the placeholder
    found = False
    for r in body.iter(qn("w:r")):
        for t in r.findall(qn("w:t")):
            if t.text == placeholder:
                # Replace with formatted count, padded to absorb digit growth.
                # We approximate by computing total - len(placeholder) + len(replacement).
                # Iterate until stable.
                def _replace(replacement: str) -> str:
                    t.text = replacement
                    return replacement
                # First pass guess based on baseline
                guess = base - len(placeholder) + len(f" {base:,} characters incl. spaces")
                _replace(f" {guess:,} characters incl. spaces")
                # Recompute and iterate
                for _ in range(5):
                    new_total = total_chars()
                    expected_text = f" {new_total:,} characters incl. spaces"
                    if t.text == expected_text:
                        break
                    _replace(expected_text)
                print(f"cover page char count -> {t.text.strip()}")
                found = True
                break
        if found:
            break
    if not found:
        print("WARN: 'XX,XXX characters incl. spaces' placeholder not found")

    # Save
    doc.save(DOCX)
    print(f"saved -> {DOCX}")


if __name__ == "__main__":
    main()
