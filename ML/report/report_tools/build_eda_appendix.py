"""Build a sidecar appendix docx for the wallet-joined EDA panels.

Requires: python-docx  (conda activate py312 && pip install python-docx)

Produces `outputs/eda/eda_appendix.docx` (and a matching .md preview).
Each panel from `outputs/eda/` lands as a captioned figure, formatted to
match the report's existing Appendix A style (A.10 onward).

This file is a STANDALONE document, it does NOT modify the submitted
docx. Once captions and figure ordering are reviewed, hand-merge into
the paper at the end of Appendix A.

Usage:
    python scripts/build_eda_appendix.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
EDA_DIR = ROOT / "submission" / "report_assets" / "figures" / "appendix" / "eda"
OUT_DOCX = EDA_DIR / "eda_appendix.docx"
OUT_MD = EDA_DIR / "eda_appendix.md"

FIG_W_IN = 6.3  # Design.md FIG_W for A4 text-width

# Each entry: appendix label, image filename in EDA_DIR, caption text.
# Captions mirror the academic style used in the existing report (no
# em-dashes, no contractions, lead with what the figure shows). Numbers
# pick up where the original A.9 left off.
APPENDIX_ITEMS: list[dict] = [
    {
        "label": "Figure A.10.",
        "image": "08_train_test_shift.png",
        "caption": (
            "Train-to-test mean shift for the top 15 numeric features, "
            "expressed as the absolute standardised mean difference "
            "(Cohen d convention: 0.2 small, 0.5 medium, 0.8 large). "
            "Time-to-deadline indicators show large shift "
            "(d above 0.9), reflecting the regime change between the "
            "Iran-strike and Iran-ceasefire countdowns. Wallet-history "
            "features show medium shift, motivating the per-market "
            "GroupKFold protocol."
        ),
    },
    {
        "label": "Figure A.11.",
        "image": "09_late_flow.png",
        "caption": (
            "Hit rate and trade-count share by time-to-deadline bucket "
            "for train (Iran strike) and test (Iran ceasefire) cohorts. "
            "Test base rate falls below 0.40 inside the final hour and "
            "below 0.34 in the last five minutes, replicating the "
            "late-concentrated informed-flow signature documented in "
            "Mitts and Ofir 2026. The corresponding train-cohort "
            "buckets stay near the 0.50 base rate, isolating the "
            "phenomenon to the ceasefire regime."
        ),
    },
    {
        "label": "Figure A.12.",
        "image": "10_wallet_strata.png",
        "caption": (
            "Bet-correct base rate stratified by three Layer 6 wallet "
            "features. Left: deciles of wallet age in days at the time "
            "of the trade. Centre: causal CEX-funding flag, defined as "
            "first USDC inbound from a known CEX hot wallet observed "
            "before the trade timestamp. Right: deciles of polygon "
            "nonce at trade time, a proxy for prior on-chain "
            "experience. The split structure motivates inclusion of "
            "the Layer 6 enrichment in the strict-branch feature set."
        ),
    },
    {
        "label": "Figure A.13.",
        "image": "11_per_market_bimodality.png",
        "caption": (
            "Per-market hit rate decomposed by trade side. The x-axis "
            "plots the share of YES-side trades that ended bet_correct "
            "in a market, and the y-axis plots the same share for "
            "NO-side trades. Marker size is proportional to log10 "
            "trade count per market. The 73 markets cleanly partition "
            "into two clusters: 49 YES-resolved markets in the "
            "lower-right (high YES-side hit rate, low NO-side) and 24 "
            "NO-resolved markets in the upper-left (mirror image). "
            "The clusters sit near (0.7, 0.3) rather than (1.0, 0.0) "
            "because bet_correct also reflects BUY vs SELL direction "
            "within each side, so YES SELL trades behave like NO "
            "bets. The two-cluster structure is the reason "
            "single-feature ROC inverts across markets and motivates "
            "the GroupKFold(market_id) evaluation protocol."
        ),
    },
    {
        "label": "Figure A.14.",
        "image": "12_feature_stability.png",
        "caption": (
            "Single-feature ROC-AUC heatmap across markets for the "
            "top eight features by absolute Pearson correlation with "
            "the target. ROC-AUC ranges from 0 to 1 and measures, "
            "for one feature alone, the probability that a randomly "
            "chosen positive trade (bet_correct=1) has a higher "
            "feature value than a randomly chosen negative trade "
            "(bet_correct=0). 0.5 is no signal, values above 0.5 "
            "mean the feature ranks positives higher, and values "
            "below 0.5 mean the relationship inverts. Per-market "
            "AUCs straddling 0.5 therefore indicate a feature whose "
            "direction flips between YES- and NO-resolved markets. "
            "Several features (log_payoff_if_correct, "
            "contrarian_score, risk_reward_ratio_pre) achieve median "
            "single-feature AUC near 0.31 with p95 above 0.67, the "
            "structural finding behind the per-market resolution "
            "split in Figure A.13 and the reason group-aware cross "
            "validation is required."
        ),
    },
    {
        "label": "Figure A.15.",
        "image": "13_mutual_information.png",
        "caption": (
            "Top 25 features by mutual information with bet_correct, "
            "computed on a 150,000-row stratified sample using "
            "scikit-learn's mutual_info_classif and colored by "
            "feature group. Mutual information captures non-linear "
            "dependence that the Pearson correlation in Figure A.5 "
            "cannot detect. The x-axis is zoomed to the relevant "
            "range so the rank gap between adjacent features is "
            "legible. The leading features are short-window "
            "market-state and price microstructure variables "
            "(realized volatility, jump component, order-flow "
            "imbalance, recent volume) rather than wallet-identity "
            "features, consistent with the literature emphasis on "
            "flow signals."
        ),
    },
    {
        "label": "Figure A.16.",
        "image": "14_feature_taxonomy.png",
        "caption": (
            "Numeric feature counts by engineering layer. Of the 79 "
            "numeric features in the wallet-joined cohort, the "
            "market-state rolling layer contributes the largest "
            "share, followed by price and volatility microstructure, "
            "the HF-internal taker aggregates, and the on-chain "
            "Layer 6 wallet features. The taxonomy frames how the "
            "report's permutation-importance and ablation analyses "
            "are organised."
        ),
    },
    {
        "label": "Figure A.17.",
        "image": "19_event_timing.png",
        "caption": (
            "Trade volume per day across the joined cohort, colored by "
            "split and annotated with the two events that anchor the "
            "cohort design. The train cohort runs from late December "
            "through the Iran strike on 28 February, ending in a single "
            "intraday spike of roughly ninety-three thousand trades on "
            "the day of the event. The test cohort runs from 1 March "
            "through 7 April and is distributed across the run-up to "
            "the ceasefire announcement, with no comparable spike. The "
            "calendar separation between the two regimes is the "
            "structural reason the cohort is split by event rather "
            "than randomly."
        ),
    },
    {
        "label": "Figure A.18.",
        "image": "16_temporal_drift.png",
        "caption": (
            "Daily bet_correct base rate during the strike-countdown "
            "(train) and ceasefire-countdown (test) windows, with a "
            "seven-day rolling mean overlay. Daily bins are filtered "
            "to days with at least 50 trades. The rolling means stay "
            "within a tight band around 0.50, confirming there is no "
            "temporal drift in the target within either cohort that "
            "would invalidate the static train and test split."
        ),
    },
]


# Inline tables that ride alongside the figures. Each entry: caption
# label, source CSV name, header row, body rows.
APPENDIX_TABLES: list[dict] = [
    {
        "label": "Table A.1.",
        "caption": (
            "Cohort sizing per split, summarising the wide trade-count "
            "spread across the 73 markets that motivates the "
            "GroupKFold(market_id) protocol used in the modeling "
            "section. Source: 07_cohort_sizing_table.csv."
        ),
        "headers": ["Split", "n markets", "Min trades / market",
                    "Median trades / market", "Max trades / market"],
        "rows": [
            ["Train", "63", "1,274", "9,753", "119,842"],
            ["Test", "10", "5,127", "14,692", "90,730"],
        ],
    },
]


def _add_caption(doc: Document, label: str, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label_run = p.add_run(label + " ")
    label_run.bold = True
    cap_run = p.add_run(caption)
    cap_run.italic = True


def _add_figure(doc: Document, image_path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(FIG_W_IN))


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val


def build_docx() -> None:
    doc = Document()
    # Default font tweak so output matches typical CBS body copy.
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_heading("Appendix A — Wallet-joined cohort EDA (sidecar)", level=1)
    doc.add_paragraph(
        "This sidecar groups the additional EDA figures produced after "
        "the wallet enrichment was extended to full coverage of Alex's "
        "1,371,180-trade idea1 cohort (100 percent of train and 100 percent "
        "of test). Numbering picks up after the existing Appendix A.9 in "
        "the main report. Each figure is independent and may be merged "
        "into the appendix at the editor's discretion."
    )
    doc.add_paragraph("")

    for item in APPENDIX_ITEMS:
        img = EDA_DIR / item["image"]
        if not img.exists():
            doc.add_paragraph(
                f"[missing image: {item['image']}]"
            )
            _add_caption(doc, item["label"], item["caption"])
            continue
        _add_figure(doc, img)
        _add_caption(doc, item["label"], item["caption"])
        doc.add_paragraph("")

    for tbl in APPENDIX_TABLES:
        _add_table(doc, tbl["headers"], tbl["rows"])
        _add_caption(doc, tbl["label"], tbl["caption"])
        doc.add_paragraph("")

    doc.save(str(OUT_DOCX))
    print(
        f"saved {OUT_DOCX.relative_to(ROOT)}  "
        f"({len(APPENDIX_ITEMS)} figures, {len(APPENDIX_TABLES)} tables)"
    )


def build_md() -> None:
    """Markdown twin so the captions are reviewable in plain text."""
    lines = ["# Appendix A — Wallet-joined cohort EDA (sidecar)", ""]
    lines.append(
        "Sidecar EDA captions for review before merging into the main "
        "report. Numbering picks up after the existing Appendix A.9."
    )
    lines.append("")
    for item in APPENDIX_ITEMS:
        img = EDA_DIR / item["image"]
        present = "✓" if img.exists() else "MISSING"
        lines.append(f"### {item['label']} ({present})")
        lines.append(f"`{item['image']}`")
        lines.append("")
        lines.append(item["caption"])
        lines.append("")
    for tbl in APPENDIX_TABLES:
        lines.append(f"### {tbl['label']}")
        lines.append("")
        lines.append("| " + " | ".join(tbl["headers"]) + " |")
        lines.append("|" + "|".join(["---"] * len(tbl["headers"])) + "|")
        for row in tbl["rows"]:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")
        lines.append(tbl["caption"])
        lines.append("")
    OUT_MD.write_text("\n".join(lines) + "\n")
    print(f"saved {OUT_MD.relative_to(ROOT)}")


def main() -> None:
    build_docx()
    build_md()


if __name__ == "__main__":
    main()
