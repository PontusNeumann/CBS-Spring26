"""Restructure ML_final_exam_paper.docx ahead of modelling work.

Steps:
  1. Drop page-break-before on Heading 1 so sections flow continuously.
  2. Promote "References" from Reference Heading to Heading 1 so it joins the
     auto-outline numbering and enters the regenerated TOC.
  3. Rename body "Conclusion" to "Conclusion and Future Work" to match TOC.
  4. Clear the body of every section after EDA; leave only the titles /
     subtitles as placeholders.
  5. Insert a one-line Normal paragraph before every heading that follows
     Normal content (the "headroom" pattern already used in Introduction).
  6. Insert two key tables inline in EDA: dataset summary and time-to-
     settlement bucket evidence for the home-run rule.
  7. Set w:updateFields=true in settings.xml so Word regenerates the TOC on
     open (picks up the Appendix entry and renumbered References).
"""
from __future__ import annotations

from pathlib import Path
import re
import shutil
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "ML_final_exam_paper.docx"

HEADING_STYLE_IDS = {"Heading1", "Heading2", "Heading3", "ReferenceHeading"}


def element_style_id(elem):
    if elem.tag != qn("w:p"):
        return None
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        return None
    ps = pPr.find(qn("w:pStyle"))
    return ps.get(qn("w:val")) if ps is not None else None


def is_blank_normal(elem) -> bool:
    if elem.tag != qn("w:p"):
        return False
    sid = element_style_id(elem)
    if sid not in (None, "Normal"):
        return False
    text = "".join(t.text or "" for t in elem.iter(qn("w:t")))
    return text.strip() == ""


def clear_section_after(heading_para) -> int:
    removed = 0
    nxt = heading_para._p.getnext()
    while nxt is not None:
        if nxt.tag == qn("w:p") and element_style_id(nxt) in HEADING_STYLE_IDS:
            break
        to_remove = nxt
        nxt = nxt.getnext()
        to_remove.getparent().remove(to_remove)
        removed += 1
    return removed


def new_paragraph(text: str = "", style_id_val: str = "Normal"):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), style_id_val)
    pPr.append(ps)
    p.append(pPr)
    if text:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
    return p


def insert_blank_before(elem, style_id_val: str = "Normal"):
    elem.addprevious(new_paragraph("", style_id_val))


def find_para_by_text(doc, text_fragment: str):
    for p in doc.paragraphs:
        if text_fragment in (p.text or ""):
            return p
    raise LookupError(f"paragraph containing {text_fragment!r} not found")


def find_heading(doc, text: str, style_name: str | None = None):
    for p in doc.paragraphs:
        if p.text.strip() == text and (style_name is None or p.style.name == style_name):
            return p
    raise LookupError(f"heading {text!r} not found (style={style_name})")


def build_table_xml(header, rows):
    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tblLook = OxmlElement("w:tblLook")
    tblLook.set(qn("w:val"), "04A0")
    tblPr.append(tblLook)
    tbl.append(tblPr)

    tblGrid = OxmlElement("w:tblGrid")
    for _ in header:
        tblGrid.append(OxmlElement("w:gridCol"))
    tbl.append(tblGrid)

    def make_cell(text: str, bold: bool = False):
        tc = OxmlElement("w:tc")
        tc.append(OxmlElement("w:tcPr"))
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        if bold:
            rPr = OxmlElement("w:rPr")
            rPr.append(OxmlElement("w:b"))
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    tr = OxmlElement("w:tr")
    for h in header:
        tr.append(make_cell(h, bold=True))
    tbl.append(tr)
    for row in rows:
        tr = OxmlElement("w:tr")
        for v in row:
            tr.append(make_cell(str(v)))
        tbl.append(tr)

    return tbl


def insert_table_after(anchor_para, label, caption, header, rows):
    anchor = anchor_para._p

    cap = OxmlElement("w:p")
    cap_pPr = OxmlElement("w:pPr")
    cap_ps = OxmlElement("w:pStyle")
    cap_ps.set(qn("w:val"), "Normal")
    cap_pPr.append(cap_ps)
    cap.append(cap_pPr)

    run_label = OxmlElement("w:r")
    rPr_b = OxmlElement("w:rPr")
    rPr_b.append(OxmlElement("w:b"))
    run_label.append(rPr_b)
    t_label = OxmlElement("w:t")
    t_label.text = f"{label} "
    t_label.set(qn("xml:space"), "preserve")
    run_label.append(t_label)
    cap.append(run_label)

    run_cap = OxmlElement("w:r")
    rPr_i = OxmlElement("w:rPr")
    rPr_i.append(OxmlElement("w:i"))
    run_cap.append(rPr_i)
    t_cap = OxmlElement("w:t")
    t_cap.text = caption
    t_cap.set(qn("xml:space"), "preserve")
    run_cap.append(t_cap)
    cap.append(run_cap)

    tbl = build_table_xml(header, rows)
    trailing_blank = new_paragraph("", "Normal")
    leading_blank = new_paragraph("", "Normal")

    anchor.addnext(trailing_blank)
    anchor.addnext(tbl)
    anchor.addnext(cap)
    anchor.addnext(leading_blank)


SECTIONS_TO_CLEAR = [
    ("Data Pre-Processing", "Heading 2"),
    ("Data Filtering, Transformation and Combination", "Heading 2"),
    ("Primary model - MLP for probability estimation", "Heading 3"),
    ("Trading rule on the probability gap", "Heading 3"),
    ("Unsupervised arm - autoencoder and Isolation Forest", "Heading 3"),
    ("Baselines and the market-implied benchmark", "Heading 3"),
    ("Evaluation Metrics", "Heading 2"),
    ("Model Complexity Analysis", "Heading 2"),
    ("Key Findings", "Heading 2"),
    ("Actionable Insights", "Heading 2"),
    ("Practical Outcomes", "Heading 2"),
    ("Ethical Consideration", "Heading 1"),
    ("Answers to the Research Questions", "Heading 2"),
    ("Implications and Learning Reflections", "Heading 2"),
    ("Limitations", "Heading 2"),
    ("Conclusion", "Heading 1"),
]


def set_update_fields(docx_path: Path) -> None:
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    updated_fragment = f'<w:updateFields xmlns:w="{W_NS}" w:val="true"/>'
    tmp = docx_path.with_suffix(".tmp.docx")

    with zipfile.ZipFile(docx_path, "r") as zin:
        buffers = {n: zin.read(n) for n in zin.namelist()}

    key = "word/settings.xml"
    if key not in buffers:
        minimal = (
            "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
            f'<w:settings xmlns:w="{W_NS}">{updated_fragment}</w:settings>'
        ).encode()
        buffers[key] = minimal
    else:
        xml_text = buffers[key].decode()
        if "<w:updateFields" in xml_text:
            xml_text = re.sub(r"<w:updateFields[^/>]*/?>", updated_fragment, xml_text)
        else:
            xml_text = xml_text.replace("</w:settings>", f"{updated_fragment}</w:settings>")
        buffers[key] = xml_text.encode()

    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, data in buffers.items():
            zout.writestr(n, data)
    shutil.move(str(tmp), str(docx_path))


def main() -> None:
    doc = Document(str(DOCX))

    # 1. Remove pageBreakBefore from Heading 1 style
    h1_style = doc.styles["Heading 1"]
    pPr = h1_style.element.find(qn("w:pPr"))
    if pPr is not None:
        pbb = pPr.find(qn("w:pageBreakBefore"))
        if pbb is not None:
            pPr.remove(pbb)
            print("[1] removed pageBreakBefore from Heading 1 style")

    # 2. Promote References to Heading 1 and rename Conclusion
    for p in doc.paragraphs:
        if p.text.strip() == "References" and p.style.name == "Reference Heading":
            p.style = doc.styles["Heading 1"]
            print("[2] References -> Heading 1")
            break
    for p in doc.paragraphs:
        if p.text.strip() == "Conclusion" and p.style.name == "Heading 1":
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
            p.add_run("Conclusion and Future Work")
            print("[3] Conclusion renamed to 'Conclusion and Future Work'")
            break

    # 4. Clear body of each target section
    cleared = 0
    for text, style_name in SECTIONS_TO_CLEAR:
        try:
            heading = find_heading(doc, text, style_name)
        except LookupError:
            print(f"  ! missing heading: {text!r}")
            continue
        cleared += clear_section_after(heading)
    print(f"[4] cleared {cleared} paragraphs/tables across {len(SECTIONS_TO_CLEAR)} sections")

    # 5. Insert inline tables in EDA
    summary_anchor = find_para_by_text(doc, "The dataset spans 1,209,787 resolved trades")
    insert_table_after(
        summary_anchor,
        label="Table 1.",
        caption="Dataset summary statistics.",
        header=["Metric", "Value"],
        rows=[
            ["Trades", "1,209,787"],
            ["Unique wallets (proxyWallet)", "109,080"],
            ["Resolved markets (condition_id)", "74"],
            ["Timespan", "2025-12-22 to 2026-04-19"],
            ["Mean bet_correct", "0.504"],
            ["Train rows (mean bet_correct)", "846,847 (0.495)"],
            ["Validation rows (mean bet_correct)", "181,454 (0.523)"],
            ["Test rows (mean bet_correct)", "181,486 (0.522)"],
        ],
    )
    timing_anchor = find_para_by_text(
        doc, "Aggregating trading activity and correctness by time-to-settlement bucket"
    )
    insert_table_after(
        timing_anchor,
        label="Table 2.",
        caption="Trade count, total USD volume, and mean bet_correct by "
                "time-to-settlement bucket (post-resolution filter applied).",
        header=["Bucket", "Trades", "Volume (USD)", "Mean bet_correct"],
        rows=[
            ["<1h",   "31,926",  "10,122,690", "0.5442"],
            ["1-6h",  "75,645",  "22,305,230", "0.5088"],
            ["6-24h", "209,170", "45,241,500", "0.5292"],
            ["1-7d",  "421,442", "73,156,270", "0.4872"],
            ["7-30d", "284,107", "40,551,150", "0.4967"],
            [">30d",  "62,096",  "5,528,511",  "0.5304"],
        ],
    )
    print("[5] inserted 2 inline EDA tables")

    # 6. Headroom blank before every heading preceded by non-blank Normal
    body = doc.element.body
    inserted_blanks = 0
    for elem in list(body.iterchildren()):
        if elem.tag != qn("w:p"):
            continue
        sid = element_style_id(elem)
        if sid not in HEADING_STYLE_IDS:
            continue
        prev = elem.getprevious()
        if prev is None:
            continue
        if prev.tag == qn("w:p"):
            prev_sid = element_style_id(prev)
            if prev_sid in HEADING_STYLE_IDS:
                continue
            if is_blank_normal(prev):
                continue
        insert_blank_before(elem, "Normal")
        inserted_blanks += 1
    print(f"[6] inserted {inserted_blanks} headroom blanks before headings")

    doc.save(str(DOCX))

    # 7. Force TOC refresh on next Word open
    set_update_fields(DOCX)
    print("[7] set w:updateFields=true")

    print(f"\nsaved {DOCX}")


if __name__ == "__main__":
    main()
