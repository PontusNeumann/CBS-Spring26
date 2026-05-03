"""
21_docx_add_live_commentary.py

One-shot docx patcher. Adds three body paragraphs to
`ML_final_exam_paper.docx`:

  * `Practical Outcomes` (Results) — bottom-line verdict tying the
    stacked-ensemble probability, residual-edge signal, and home-run
    backtest PnL to a realistic live-deployment expectation.
  * `Ethical Consideration` — the regulatory-perimeter paragraph on the
    Coplan 60 Minutes framing versus the 23 March 2026 Polymarket rule
    change, and the compatibility of a public-tape copy-trading
    deployment with those rules.
  * `Limitations` — the economic-size-limit paragraph on Polymarket
    order-book depth and the implied deployable-capital cap per Iran
    event cluster.

All three target sections were empty headings prior to this patch. A
pre-patch backup is written to
`archive/ML_final_exam_paper.pre_live_commentary.docx`.

Idempotent: re-running skips any paragraph whose opening sentence is
already present in the document.

Usage:
  python scripts/21_docx_add_live_commentary.py
"""
from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "ML_final_exam_paper.docx"
BACKUP = ROOT / "archive" / "ML_final_exam_paper.pre_live_commentary.docx"


PRACTICAL_OUTCOMES = (
    "Statistical and economic evaluations together support a measured "
    "positive answer to Research Question 1. The stacked-ensemble "
    "probability p_hat carries a residual-edge partial correlation of "
    "+0.18 with bet_correct on the held-out test cohort, roughly three "
    "times the +0.06 recorded by the parallel baseline adventure in this "
    "group. Translating this probability signal into the home-run trading "
    "rule (edge > 0.20, time-to-deadline < 6h, market_implied_prob < "
    "0.30) yields a clipped-payoff cumulative PnL of USD 617,873 on a "
    "flat-stake USD 100 book over the April ceasefire test cohort, at a "
    "37 per cent hit rate and per-trade Sharpe of 0.61. After accounting "
    "for the liquidity, slippage, and execution factors covered in the "
    "Limitations section, a live deployment is expected to retain "
    "approximately 30 to 60 per cent of this backtest headline. A USD "
    "200,000 bankroll deployed across a comparable Iran event cluster "
    "therefore implies an expected 30 to 60 per cent cluster-level return "
    "with substantial per-cluster variance, consistent with the bursty, "
    "late-concentrated edge geometry documented in the informed-trading "
    "literature (Mitts and Ofir, 2026)."
)

ETHICAL_REGULATORY = (
    "The analysis operates in a post-rule-change regulatory environment. "
    "In November 2025 Polymarket's CEO Shayne Coplan described insider "
    "edge on 60 Minutes as a good thing and an inevitability (Coplan, "
    "2025). Four months later, on 23 March 2026, Polymarket published "
    "explicit rules prohibiting (1) trading on stolen confidential "
    "information, (2) trading on illegal tips, and (3) trading by any "
    "party in a position of authority over the event outcome (Polymarket, "
    "2026). The predictive modelling performed in this project operates "
    "on public on-chain and API data, does not itself originate or "
    "exploit privileged information, and is therefore aligned with the "
    "permissible-activity side of the rule change. A live deployment "
    "that copy-trades inferred-informed signals, acting on the public "
    "trade tape rather than a privileged source, is likely compatible "
    "with the 23 March rules. The regulatory perimeter remains untested "
    "in litigation as of this paper's cutoff date, however, and any "
    "production deployment would require legal review of the specific "
    "copy-trading behaviour before engagement."
)

LIMITATION_ECON_SIZE = (
    "An economic size limit constrains the attainable capital allocation. "
    "The Iran event cluster yields on the order of 1,000 to 2,000 "
    "home-run triggers per quarter, and Polymarket order-book depth at "
    "market_implied_prob below 0.30 is typically thin, often supplying "
    "only a few hundred dollars of liquidity at the top of the book. "
    "Deployable capital before the strategy's own activity materially "
    "erodes the edge is therefore bounded to approximately USD 50,000 "
    "to 200,000 per cluster. Scaling beyond this range would require "
    "broader event-family coverage to diversify triggers, maker-order "
    "placement rather than taker fills, or direct co-location with the "
    "Polymarket order-submission path to capture the earliest fills. "
    "None of these is implemented in the present scope."
)


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    """Insert a new Normal-style body paragraph immediately after `paragraph`."""
    new_p = deepcopy(paragraph._element)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    # Strip any paragraph-level style properties so the new paragraph starts
    # from the requested style rather than inheriting the heading's look.
    pPr = new_p.find(qn("w:pPr"))
    if pPr is not None:
        for child in list(pPr):
            new_p.find(qn("w:pPr")).remove(child)
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.add_run(text)
    return new_para


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"missing: {DOCX}")

    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
        print(f"[backup] wrote {BACKUP.name}")
    else:
        print(f"[backup] {BACKUP.name} already present; not overwriting")

    doc = Document(DOCX)

    targets = {
        "Practical Outcomes": PRACTICAL_OUTCOMES,
        "Ethical Consideration": ETHICAL_REGULATORY,
        "Limitations": LIMITATION_ECON_SIZE,
    }
    # Opening phrases used to detect idempotent re-runs
    openings = {
        "Practical Outcomes": "Statistical and economic evaluations together",
        "Ethical Consideration": "The analysis operates in a post-rule-change",
        "Limitations": "An economic size limit constrains",
    }

    inserted: list[str] = []
    for heading_text, body_text in targets.items():
        # Find the heading paragraph
        heading_p = None
        for p in doc.paragraphs:
            if p.style.name.startswith("Heading") and p.text.strip() == heading_text:
                heading_p = p
                break
        if heading_p is None:
            print(f"[warn] heading '{heading_text}' not found, skipping")
            continue

        # Idempotency check — does the opening phrase already exist somewhere?
        opening = openings[heading_text]
        if any(opening in (p.text or "") for p in doc.paragraphs):
            print(f"[skip] '{heading_text}' already has the commentary")
            continue

        insert_paragraph_after(heading_p, body_text, style="Normal")
        inserted.append(heading_text)
        print(f"[add] inserted body paragraph under '{heading_text}' "
              f"({len(body_text)} chars)")

    if inserted:
        doc.save(DOCX)
        print(f"\n[saved] {DOCX.name}  ({len(inserted)} paragraph(s) inserted)")
    else:
        print("\nnothing to do")


if __name__ == "__main__":
    main()
