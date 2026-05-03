"""
22_docx_add_feature_inventory.py

Appends two tables to the docx Appendix:
  - Table A.6 — Retained columns (54) with role + definition / formula.
  - Table A.7 — Dropped columns (25) with tier + drop rationale.

Retained columns listed first (kept variables grouped by semantic family,
non-features clearly labelled), dropped columns last, as requested.

Idempotent: re-running detects the A.6 heading and skips the insert.

Usage:
  python scripts/22_docx_add_feature_inventory.py
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "ML_final_exam_paper.docx"
BACKUP = ROOT / "archive" / "ML_final_exam_paper.pre_feature_inventory.docx"

# -------------------------------------------------------------------
# Retained columns — grouped semantically. Non-features carry a
# "non-feature" role tag so a reader can see at a glance which rows
# actually enter the model.
# -------------------------------------------------------------------
RETAINED: list[tuple[str, str, str, str]] = [
    # Group, Variable, Role, Definition / formula
    # Identifiers and raw trade primitives
    ("Identifiers", "proxyWallet",         "id",            "Taker's 0x Polygon wallet address."),
    ("Identifiers", "asset",               "id",            "ERC-1155 outcome-token id (77-digit decimal)."),
    ("Identifiers", "transactionHash",     "id",            "Polygon transaction hash of the order fill."),
    ("Identifiers", "condition_id",        "id",            "Polymarket market id (0x-prefixed 64-char hash)."),
    ("Identifiers", "source",              "id",            "Data provenance: hf (HuggingFace mirror) or api (Polymarket Data API)."),
    ("Raw trade",   "timestamp",           "id",            "Trade execution time (UTC)."),
    ("Raw trade",   "size",                "raw",           "Trade size in token units (shares)."),
    ("Raw trade",   "price",               "raw",           "Trade execution price in [0, 1]."),
    # Market metadata
    ("Market meta", "question",            "meta",          "Market question text. Parsed to derive deadline_ts."),
    ("Market meta", "end_date",            "meta",          "Gamma API endDate. Stale on 34 / 74 markets. Kept as metadata, not a feature."),
    ("Market meta", "resolution_ts",       "filter input",  "Post-hoc CLOB-lock timestamp. Used only in the settlement filter, never as a feature."),
    ("Market meta", "deadline_ts",         "meta",          "Advertised deadline parsed from question (\"by <Month> <Day>, <Year>\"). Published at market creation."),
    ("Market meta", "winning_outcome_index","meta",         "Resolved winning outcomeIndex (0 or 1). Label-derivation only."),
    ("Market meta", "resolved",            "meta",          "Boolean. True for all rows in the 74-market cohort."),
    ("Market meta", "is_yes",              "label",         "1 iff outcomes[winning_outcome_index] == \"Yes\". Not a feature."),
    ("Market meta", "settlement_minus_trade_sec", "filter", "resolution_ts − timestamp (sec). §4 filter drops rows ≤ 0."),
    ("Market meta", "market_implied_prob", "benchmark",     "CLOB mid-price at t via merge_asof(backward, closed-left); falls back to price on HF markets. Trading-rule benchmark. Excluded from the feature set."),
    ("Label",       "bet_correct",         "target",        "TARGET. (outcomeIndex == winning_outcome_index) == (side == \"BUY\")."),
    # Trade-local features
    ("Trade-local", "log_size",            "feature",       "log1p(max(0, size))."),
    ("Trade-local", "trade_value_usd",     "feature",       "size × price. USD notional of this trade."),
    # Market context (normalised only after Tier 2 drop)
    ("Market context", "market_price_vol_last_1h", "feature", "rolling_std(market_implied_prob, 1h, closed=\"left\") per market. Bounded price-based volatility, not scale."),
    # Time (all from deadline_ts, not resolution_ts)
    ("Time",        "pct_time_elapsed",    "feature",       "(t − market_start) / (deadline_ts − market_start), clipped to [0, 1]."),
    ("Time",        "market_timing_known", "feature",       "1 iff pct_time_elapsed is defined."),
    # Wallet global
    ("Wallet global","wallet_prior_trades","feature",       "groupby(wallet).cumcount() strictly before t."),
    ("Wallet global","wallet_prior_volume_usd", "feature",  "cumsum(trade_value_usd) per wallet, minus current row."),
    ("Wallet global","wallet_first_minus_trade_sec", "feature", "(first_ever_trade_of_wallet − t).dt.total_seconds(). Always ≤ 0."),
    ("Wallet global","wallet_prior_win_rate_causal","feature","Cumulative mean of bet_correct restricted to priors whose resolution_ts < t."),
    ("Wallet global","wallet_has_resolved_priors","feature","1 iff ≥ 1 prior with resolution_ts < t exists."),
    ("Wallet global","wallet_has_prior_trades","feature",   "1 iff wallet_prior_trades > 0."),
    # Wallet-in-market bursting
    ("Bursting",    "wallet_trades_in_market_last_1min",  "feature", "rolling_count per (wallet, condition_id), 60s window, closed-left."),
    ("Bursting",    "wallet_trades_in_market_last_10min", "feature", "Same with 600s window."),
    ("Bursting",    "wallet_trades_in_market_last_60min", "feature", "Same with 3600s window."),
    ("Bursting",    "wallet_is_burst",                    "feature", "1 iff wallet_trades_in_market_last_10min ≥ 3."),
    ("Bursting",    "wallet_median_gap_in_market",        "feature", "Per (wallet, market) expanding median of consecutive-trade gaps, shifted by 1."),
    # Wallet-in-market directional (symmetric only)
    ("Directional", "wallet_spread_ratio",  "feature",     "min(cum_0, cum_1) / max(cum_0, cum_1). Symmetric under swapping outcomes 0 and 1."),
    # Wallet-in-market position
    ("Position",    "wallet_is_whale_in_market", "feature", "1 iff cumulative wallet volume ≥ expanding per-market p95. Causal, 20-wallet warmup."),
    # Wallet-in-market depth
    ("Depth",       "wallet_prior_trades_in_market", "feature", "groupby((wallet, condition_id)).cumcount() strictly before t."),
    ("Depth",       "wallet_has_prior_trades_in_market", "feature", "1 iff wallet_prior_trades_in_market > 0."),
    # Interactions
    ("Interactions","size_vs_wallet_avg",        "feature", "trade_value_usd / (wallet_prior_volume_usd / wallet_prior_trades)."),
    ("Interactions","size_vs_market_cumvol_pct", "feature", "trade_value_usd / market_volume_so_far_usd. Top permutation-importance feature (Δroc +0.049)."),
    ("Interactions","size_vs_market_avg",        "feature", "trade_value_usd / (market_volume_so_far_usd / market_trade_count_so_far)."),
    # Layer 6 — on-chain identity
    ("Layer 6",     "wallet_enriched",             "feature", "1 iff Etherscan V2 returned non-empty tokentx for the wallet."),
    ("Layer 6",     "wallet_polygon_age_at_t_days","feature", "max(0, t − polygon_first_tx_ts) / 86400."),
    ("Layer 6",     "wallet_polygon_nonce_at_t",   "feature", "searchsorted(outbound_ts, t, side=\"left\"). Count of prior outbound ERC-20 transfers."),
    ("Layer 6",     "wallet_log_polygon_nonce_at_t","feature","log1p(wallet_polygon_nonce_at_t)."),
    ("Layer 6",     "wallet_n_inbound_at_t",       "feature", "Count of prior inbound ERC-20 transfers."),
    ("Layer 6",     "wallet_log_n_inbound_at_t",   "feature", "log1p(wallet_n_inbound_at_t)."),
    ("Layer 6",     "wallet_n_cex_deposits_at_t",  "feature", "Count of prior inbound USDC transfers from a known CEX hot-wallet."),
    ("Layer 6",     "wallet_cex_usdc_cumulative_at_t","feature","Cumulative USD from CEX inflows strictly before t."),
    ("Layer 6",     "wallet_log_cex_usdc_cum",     "feature", "log1p(wallet_cex_usdc_cumulative_at_t)."),
    ("Layer 6",     "days_from_first_usdc_to_t",   "feature", "(t − first_usdc_inbound_ts) / 86400 iff first_usdc_inbound_ts < t, else NaN."),
    ("Layer 6",     "wallet_funded_by_cex_scoped", "feature", "1 iff the wallet was funded by a CEX AND first_usdc_inbound_ts < t."),
    # Layer 7 — cross-market
    ("Layer 7",     "wallet_market_category_entropy", "feature", "Shannon entropy (nats) over 8-category distribution of the wallet's prior distinct markets. NaN if < 2 prior markets."),
    ("Layer 7",     "wallet_has_cross_market_history","feature", "1 iff wallet_market_category_entropy is defined."),
]

# -------------------------------------------------------------------
# Dropped columns — by tier, with rationale. 25 items.
# -------------------------------------------------------------------
DROPPED: list[tuple[str, str, str]] = [
    # Tier 1 — confirmed leak or bug
    ("Tier 1 — leak / bug", "wallet_prior_win_rate",     "Temporal leak (P0-9). Cumulative mean over priors regardless of resolution status, so includes bet_correct of prior markets not yet resolved at t. +0.13 leak-driven correlation with target. Replaced by wallet_prior_win_rate_causal."),
    ("Tier 1 — leak / bug", "is_position_exit",          "Denominator bug (P0-2). max(|pos_before|, size) uses the current trade's size, so a wallet's first-ever SELL always fires ratio = size/size = 1.0 and flags as exit."),
    ("Tier 1 — leak / bug", "is_position_flip",          "Shares the signed-size family with is_position_exit. Inherits the denominator assumption; dropped pending re-verification."),
    ("Tier 1 — leak / bug", "wallet_funded_by_cex",      "Structurally leaky lifetime flag (defence-in-depth). Empirically 0 rows leak in this cohort but the definition is not scoped to t. The scoped variant wallet_funded_by_cex_scoped is retained."),
    # Tier 2 — market-identity absolute-scale (P0-8)
    ("Tier 2 — market identity", "time_to_settlement_s",     "Absolute-scale deadline distance. Lets a model memorise which sub-market a trade belongs to and shortcut to its resolution."),
    ("Tier 2 — market identity", "log_time_to_settlement",   "Log variant of time_to_settlement_s. Same P0-8 shortcut risk."),
    ("Tier 2 — market identity", "market_volume_so_far_usd", "Absolute USD volume. Feb 28 is roughly 10x the smaller strike markets — strong market-identity signal."),
    ("Tier 2 — market identity", "market_vol_1h_log",        "log1p(rolling 1h USD volume). Absolute scale."),
    ("Tier 2 — market identity", "market_vol_24h_log",       "Same for the 24h window."),
    ("Tier 2 — market identity", "market_trade_count_so_far", "Absolute trade count per market."),
    ("Tier 2 — market identity", "size_x_time_to_settlement", "log_size × log_time_to_settlement. Back-door leak of market identity through the interaction."),
    # Tier 3 — metadata bloat
    ("Tier 3 — metadata bloat", "conditionId",           "CamelCase duplicate of condition_id from alternate API paths."),
    ("Tier 3 — metadata bloat", "title",                 "Market title string (duplicate of question-derived content)."),
    ("Tier 3 — metadata bloat", "slug_x",                "URL slug from first API join."),
    ("Tier 3 — metadata bloat", "slug_y",                "URL slug from second API join."),
    ("Tier 3 — metadata bloat", "icon",                  "Market icon image URL."),
    ("Tier 3 — metadata bloat", "eventSlug",             "Parent-event URL slug."),
    ("Tier 3 — metadata bloat", "outcome",               "API-path \"is-YES-side\" flag. Confusingly named next to is_yes; dropped to avoid misuse."),
    ("Tier 3 — metadata bloat", "name",                  "User-profile name (API-path only)."),
    ("Tier 3 — metadata bloat", "pseudonym",             "User-profile pseudonym."),
    ("Tier 3 — metadata bloat", "bio",                   "User-profile bio string."),
    ("Tier 3 — metadata bloat", "profileImage",          "User avatar URL."),
    ("Tier 3 — metadata bloat", "profileImageOptimized", "User avatar URL variant."),
    ("Tier 3 — metadata bloat", "outcomes",              "Raw per-market outcomes array (e.g. \"Yes;No\"). Used once to derive is_yes, then redundant."),
    # Tier 4 — obsolete
    ("Tier 4 — obsolete", "split",                       "Legacy trade-timestamp quantile split. Replaced by market-cohort parquets under data/experiments/."),
    # Tier 5 — direction encoding (P0-11 + P0-12)
    ("Tier 5 — direction encoding", "side",              "P0-11. Pairs with outcomeIndex to deterministically encode bet_correct via (outcomeIndex == wi) == (side == BUY). Mapping flips across YES/NO resolution — a mixed training cohort learns an averaged rule that inverts on a single-resolution test cohort."),
    ("Tier 5 — direction encoding", "outcomeIndex",      "P0-11. See side."),
    ("Tier 5 — direction encoding", "wallet_position_size_before_trade", "P0-12. Signed cumulative position encodes direction."),
    ("Tier 5 — direction encoding", "trade_size_vs_position_pct", "P0-12. Uses signed position."),
    ("Tier 5 — direction encoding", "wallet_cumvol_same_side_last_10min", "P0-12. Explicit same-side filter."),
    ("Tier 5 — direction encoding", "wallet_directional_purity_in_market", "P0-12. outcomeIndex-share aggregate."),
    ("Tier 5 — direction encoding", "wallet_has_both_sides_in_market", "P0-12. Indicator on the outcomeIndex distribution."),
    ("Tier 5 — direction encoding", "market_buy_share_running", "P0-12. Running share of BUY per market. 30-point cohort shift between train (0.38) and test (0.67) re-opens the P0-11 channel."),
]


# -------------------------------------------------------------------
# Insertion helpers
# -------------------------------------------------------------------
def _set_cell_font(cell, size_pt: float = 9, bold: bool = False) -> None:
    """Force monospace-ish compact font on the new appendix tables so they
    fit on portrait A4 without column overflow. Uses the document default
    family."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(size_pt)
            if bold:
                run.bold = True
        if not p.runs:
            run = p.add_run()
            run.font.size = Pt(size_pt)
            if bold:
                run.bold = True


def _add_table(doc, header: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        _set_cell_font(hdr[i], size_pt=9, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _set_cell_font(cells[i], size_pt=8.5)
    return t


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"missing: {DOCX}")

    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
        print(f"[backup] wrote {BACKUP.name}")
    else:
        print(f"[backup] {BACKUP.name} already present; not overwriting")

    doc = Document(DOCX)

    # Idempotency — skip if Table A.6 already appended
    for p in doc.paragraphs:
        if "Feature inventory" in p.text and p.style.name == "Heading 2":
            print("[skip] Feature inventory already present")
            return

    # Append new subsection at end of document (end of Appendix by layout)
    print("[append] Feature inventory subsection at end of Appendix...")
    doc.add_heading("Feature inventory", level=2)
    doc.add_paragraph(
        "The following two tables document every column in the consolidated "
        "trade dataset. Table A.6 lists the retained columns — the 36 model "
        "features and the 18 non-feature utility columns (identifiers, "
        "labels, filter inputs, and the trading-rule benchmark). Table A.7 "
        "lists the 25 columns physically removed by scripts/20_finalize_"
        "dataset.py, grouped by the drop tier. The distinction between a "
        "kept feature, a non-feature utility column, and a dropped column "
        "is enforced at the file level — the final consolidated dataset "
        "contains only the rows of Table A.6."
    )

    doc.add_paragraph("Table A.6. Retained columns of the consolidated dataset "
                      "(54 total; 36 model features + 18 non-feature utility columns).")

    _add_table(
        doc,
        header=["Group", "Variable", "Role", "Definition / formula"],
        rows=[[g, v, r, d] for (g, v, r, d) in RETAINED],
    )

    doc.add_paragraph("")  # spacer
    doc.add_paragraph("Table A.7. Columns physically dropped from the consolidated "
                      "dataset by the finalisation step (25 total).")

    _add_table(
        doc,
        header=["Tier", "Variable", "Reason"],
        rows=[[t, v, r] for (t, v, r) in DROPPED],
    )

    doc.save(DOCX)
    print(f"[saved] {DOCX.name}  "
          f"(+2 tables, +{len(RETAINED) + len(DROPPED)} rows)")


if __name__ == "__main__":
    main()
