"""
27_docx_trim_sections.py

Two space-saving edits requested after the front-matter swap:

1. Shorten the "Contribution and LLM Usage Disclosure" section. Keep the
   Heading 1, replace the body paragraph with a one-sentence summary, and
   delete the multi-row table underneath. The numPr on Heading 1 auto
   numbers so the section stays as section 4.

2. Shorten the Abstract paragraph from ~270 words to ~210 words (guideline
   target is 200). Keywords section is kept — CBS project guidelines
   require at least five keywords.

Idempotent. Pre-patch backup at archive/ML_final_exam_paper.pre_trim.docx.
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "ML_final_exam_paper.docx"
BACKUP = ROOT / "archive" / "ML_final_exam_paper.pre_trim.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"

NEW_ABSTRACT = (
    "Prediction markets on Polymarket price real-world events continuously, "
    "yet whether those prices are informationally efficient remains an open "
    "question. Prior work documents USD 143 million in anomalous profits "
    "concentrated on geopolitical contracts, with named wallets active in "
    "Iran-strike markets in particular (Mitts and Ofir, 2026). This project "
    "tests whether a model trained on pre-execution market-state and "
    "behavioural features produces a probability estimate that systematically "
    "diverges from the contemporaneous market-implied probability, and "
    "whether the resulting gap is tradable. The scope covers seventy-four "
    "resolved sub-markets under Polymarket events 114242, 236884, 355299, "
    "and 357625. A market-embedding multilayer perceptron estimates a "
    "per-trade probability p_hat from standardised market-state and "
    "wallet-history features; logistic regression, random forest, and an "
    "undercomplete autoencoder serve as baselines. Evaluation proceeds on "
    "temporally held-out data against statistical criteria (ROC-AUC, Brier "
    "score, calibration) and economic criteria (PnL, Sharpe ratio, hit rate, "
    "drawdown). The MLP attains a held-out test ROC-AUC of 0.579 on the "
    "13,414-trade April ceasefire cohort, with a residual-edge partial "
    "correlation of +0.18 against market-implied probability and a "
    "slippage-robust trading rule that generates USD 617,873 in cumulative "
    "PnL at a per-trade Sharpe of 0.61. The results are consistent with a "
    "weak-form efficient market hosting a small but monetisable "
    "informed-flow residual, aligning with the concentration anomaly "
    "documented by Mitts and Ofir (2026)."
)


def _p_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(f"{W}t"))


def _p_style(el) -> str:
    s = el.find(f"{W}pPr/{W}pStyle")
    return s.get(f"{W}val") if s is not None else "Normal"


SHORT_CONTRIBUTION = (
    "All conceptual design - research question, scope, feature taxonomy, "
    "and trading-rule specification - was authored by the group without "
    "LLM assistance. Claude (Anthropic, Opus 4.7) was used for code "
    "drafting on the data-extraction, EDA, and modelling scripts and for "
    "prose drafting on parts of the Methodology and Appendix, under "
    "continuous author review, testing, and correction; no output was "
    "accepted without author validation."
)


def shorten_contribution_section(doc) -> None:
    """Keep the Heading 1 "Contribution and LLM Usage Disclosure", replace
    the body paragraph with SHORT_CONTRIBUTION, and delete the table plus
    any empty paragraphs between Heading 1 and the next Heading 1."""
    body = doc.element.body
    children = list(body)
    h1_idx = None
    next_h1_idx = None
    for i, el in enumerate(children):
        if el.tag == f"{W}p" and _p_style(el) == "Heading1" \
                and "Contribution" in _p_text(el):
            h1_idx = i
            for j in range(i + 1, len(children)):
                nxt = children[j]
                if nxt.tag == f"{W}p" and _p_style(nxt) == "Heading1":
                    next_h1_idx = j
                    break
            break
    if h1_idx is None or next_h1_idx is None:
        print("[contrib] Contribution section not located; skipping")
        return

    # Find the first <w:p Normal> after the H1 that has visible text.
    body_p_idx = None
    for k in range(h1_idx + 1, next_h1_idx):
        el = children[k]
        if el.tag == f"{W}p" and _p_text(el).strip():
            body_p_idx = k
            break

    # Import lxml helpers via python-docx's underlying library for building
    # a fresh run on an existing paragraph.
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if body_p_idx is not None:
        p = children[body_p_idx]
        # Remove all runs and other content children, keep pPr only.
        for child in list(p):
            if child.tag != f"{W}pPr":
                p.remove(child)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = SHORT_CONTRIBUTION
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
        # Ensure paragraph style is Normal
        pPr = p.find(f"{W}pPr")
        if pPr is not None:
            pStyle = pPr.find(f"{W}pStyle")
            if pStyle is not None:
                pStyle.set(f"{W}val", "Normal")
        print("[contrib] shortened body paragraph "
              f"({len(SHORT_CONTRIBUTION.split())} words)")

    # Remove everything else between the (now shortened) body paragraph and
    # the next Heading 1 - tables and trailing blank paragraphs.
    removed = 0
    keep_indices = {h1_idx}
    if body_p_idx is not None:
        keep_indices.add(body_p_idx)
    # Keep exactly one empty <w:p> right before the next Heading 1 as
    # separator - 23_docx_spacing_and_breaks.py style.
    separator_idx = None
    for k in range(next_h1_idx - 1, h1_idx, -1):
        el = children[k]
        if el.tag == f"{W}p" and not _p_text(el).strip() \
                and _p_style(el) == "Normal":
            separator_idx = k
            break
    if separator_idx is not None:
        keep_indices.add(separator_idx)

    for k in range(h1_idx + 1, next_h1_idx):
        if k in keep_indices:
            continue
        el = children[k]
        body.remove(el)
        removed += 1
    print(f"[contrib] removed {removed} table/blank element(s) from section")


def shorten_abstract(doc) -> bool:
    """Replace the single Normal paragraph immediately after the 'Abstract'
    Heading 1 with the shorter NEW_ABSTRACT text. Returns True if
    replaced."""
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.style.name == "Heading 1" and p.text.strip() == "Abstract":
            if i + 1 >= len(paras):
                return False
            target = paras[i + 1]
            # Guard: "named wallets active" appears in NEW_ABSTRACT and not
            # in the original (which has "named individual wallets operating").
            if "named wallets active" in target.text:
                print("[abstract] already shortened")
                return False
            # Clear runs, then add the shorter text as a single run in the
            # same paragraph style.
            for run in list(target.runs):
                run.text = ""
            target.add_run(NEW_ABSTRACT)
            words = len(NEW_ABSTRACT.split())
            print(f"[abstract] replaced with {words}-word version")
            return True
    print("[abstract] Abstract heading not found")
    return False


def main() -> None:
    assert DOCX.exists(), f"missing: {DOCX}"

    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
        print(f"[backup] wrote {BACKUP.name}")
    else:
        print(f"[backup] {BACKUP.name} already present; not overwriting")

    doc = Document(DOCX)
    shorten_contribution_section(doc)
    shorten_abstract(doc)
    doc.save(DOCX)
    print(f"[saved] {DOCX.name}")


if __name__ == "__main__":
    main()
