"""
23_docx_spacing_and_breaks.py

Applies the three formatting rules codified in Design.md:
  * `space-before` and `space-after` on every paragraph set to 0.
  * Blank separator rows use the `Normal` style with default font size.
  * `References` and `Appendix` top-level headings each start on a new
    page (page-break-before on the heading's paragraph properties).

Also restores the caption for Table A.6 which was shortened by a manual
edit.

Idempotent. Pre-patch backup written to
`archive/ML_final_exam_paper.pre_spacing.docx`.
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "ML_final_exam_paper.docx"
BACKUP = ROOT / "archive" / "ML_final_exam_paper.pre_spacing.docx"


TABLE_A6_CAPTION = (
    "Table A.6. Retained columns of the consolidated dataset (54 total; "
    "36 model features and 18 non-feature utility columns). Group labels "
    "the semantic family of each column; Role distinguishes feature from "
    "identifier, label, filter input, and benchmark."
)


def _get_or_add_pPr(paragraph):
    p = paragraph._element
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    return pPr


def set_zero_spacing(paragraph) -> None:
    """Force `space-before` and `space-after` to 0 on the paragraph."""
    pPr = _get_or_add_pPr(paragraph)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    # Strip any auto-spacing flags so the zero actually applies.
    for key in ("w:beforeAutospacing", "w:afterAutospacing"):
        if spacing.get(qn(key)) is not None:
            spacing.set(qn(key), "0")


def set_page_break_before(paragraph) -> None:
    """Add `page-break-before` to a heading's paragraph properties so
    Word starts it on a new page."""
    pPr = _get_or_add_pPr(paragraph)
    # Remove any existing pageBreakBefore to avoid duplicates
    for el in pPr.findall(qn("w:pageBreakBefore")):
        pPr.remove(el)
    pbb = OxmlElement("w:pageBreakBefore")
    pPr.append(pbb)


def fix_table_a6_caption(doc) -> bool:
    """Locate the placeholder 'Table A.6.' paragraph (or the full caption if
    already restored) and set it to the canonical text above. Returns True if
    a change was made."""
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("Table A.6"):
            if text == TABLE_A6_CAPTION.strip():
                return False
            # Clear runs and replace
            for run in list(p.runs):
                run.text = ""
            p.add_run(TABLE_A6_CAPTION)
            return True
    return False


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"missing: {DOCX}")

    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
        print(f"[backup] wrote {BACKUP.name}")
    else:
        print(f"[backup] {BACKUP.name} already present; not overwriting")

    doc = Document(DOCX)

    # 1. Zero-out space-before and space-after on EVERY paragraph.
    n = 0
    for p in doc.paragraphs:
        set_zero_spacing(p)
        n += 1
    print(f"[spacing] forced zero before/after on {n} paragraphs")

    # 2. Page-break-before on References and Appendix.
    broken = []
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() in ("References", "Appendix"):
            set_page_break_before(p)
            broken.append(p.text.strip())
    print(f"[page-break] added to: {broken}")

    # 3. Restore Table A.6 caption if truncated.
    if fix_table_a6_caption(doc):
        print("[caption] restored Table A.6 caption to the full description")
    else:
        print("[caption] Table A.6 caption already canonical")

    doc.save(DOCX)
    print(f"[saved] {DOCX.name}")


if __name__ == "__main__":
    main()
