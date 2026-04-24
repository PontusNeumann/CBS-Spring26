"""
25_docx_restore_front_matter.py

Restores the title page (page 1) and table of contents (page 2) to their
pre-formatting (CBS template) appearance while keeping the CBS-compliant
>=3 cm top margin on the body section.

Concretely:

1. Splits the document into two sections:
   * Section 1 (cover + TOC) uses the original template top margin of
     1304 twips (~2.30 cm).
   * Section 2 (body + references + appendix) keeps the current
     1701-twip (~3.00 cm) top margin required by the CBS formatting
     rule for body text.

2. Restores explicit TOC font sizes (TOC1 = 22 half-pt = 11 pt, TOC2 = 18,
   TOC3 = 18) so the table of contents fits on one page at the original
   template density. Without explicit sizes the TOC entries inherit 11 pt
   from Normal, which is what pushed the TOC onto two pages.

3. Shrinks section headings modestly so each Heading 1 does not dominate
   the page:
   * Heading 1: 72 -> 44 half-points (36 pt -> 22 pt); line height
     680 -> 440 twips.
   * Heading 2: 28 -> 26 half-points (14 pt -> 13 pt).
   * Heading 3: 26 -> 24 half-points (13 pt -> 12 pt).

4. Re-enforces the zero space-before / space-after rule on every paragraph
   so separation between paragraphs comes from explicit blank Normal
   paragraphs rather than style-level spacing.

Idempotent. Pre-patch backup at archive/ML_final_exam_paper.pre_restore.docx.
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "ML_final_exam_paper.docx"
BACKUP = ROOT / "archive" / "ML_final_exam_paper.pre_restore.docx"

FRONT_MATTER_TOP_MARGIN = "1304"  # twips, matches pre_formatting template

HEADING_SIZES = {
    # styleId -> (new sz half-points, optional new line twips)
    "Heading1": (44, 440),
    "Heading2": (26, None),
    "Heading3": (24, None),
}

TOC_SIZES = {
    "TOC1": 22,
    "TOC2": 18,
    "TOC3": 18,
}


def _pPr(p):
    el = p._element
    pPr = el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        el.insert(0, pPr)
    return pPr


def set_zero_spacing(p) -> None:
    pPr = _pPr(p)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    for key in ("w:beforeAutospacing", "w:afterAutospacing"):
        if spacing.get(qn(key)) is not None:
            spacing.set(qn(key), "0")


def update_heading_and_toc_sizes(styles_root) -> None:
    """Mutate the lxml <w:styles> tree, restoring TOC sizes and shrinking
    heading sizes. `styles_root` is already the CT_Styles element."""
    root = styles_root
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    for style in root.findall(f"{W}style"):
        sid = style.get(f"{W}styleId")
        # Heading size / line tweaks
        if sid in HEADING_SIZES:
            new_sz, new_line = HEADING_SIZES[sid]
            rPr = style.find(f"{W}rPr")
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                style.append(rPr)
            sz = rPr.find(f"{W}sz")
            if sz is None:
                sz = OxmlElement("w:sz")
                rPr.append(sz)
            sz.set(f"{W}val", str(new_sz))
            if new_line is not None:
                pPr = style.find(f"{W}pPr")
                if pPr is not None:
                    spacing = pPr.find(f"{W}spacing")
                    if spacing is not None:
                        spacing.set(f"{W}line", str(new_line))
                        spacing.set(f"{W}lineRule", "exact")
            print(f"[style] {sid}: sz -> {new_sz} half-pt"
                  + (f", line -> {new_line}" if new_line else ""))
        # TOC size restore
        if sid in TOC_SIZES:
            new_sz = TOC_SIZES[sid]
            rPr = style.find(f"{W}rPr")
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                style.append(rPr)
            sz = rPr.find(f"{W}sz")
            if sz is None:
                sz = OxmlElement("w:sz")
                rPr.append(sz)
            sz.set(f"{W}val", str(new_sz))
            print(f"[style] {sid}: sz -> {new_sz} half-pt")


def find_last_sectPr(body):
    """Return the final <w:sectPr> of the body (the section 2 properties)."""
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    # The final sectPr is a direct child of <w:body>, not inside a pPr.
    for child in reversed(list(body)):
        if child.tag == f"{W}sectPr":
            return child
    return None


def build_section1_sectPr(last_sectPr):
    """Clone the final sectPr and override top margin to the template value.
    Returns a new sectPr element that can be placed inside a paragraph's pPr
    to terminate section 1."""
    new = OxmlElement("w:sectPr")
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    # Copy all attributes and children
    for k, v in last_sectPr.attrib.items():
        new.set(k, v)
    for child in last_sectPr:
        new.append(child.__copy__())
    # Override top margin
    pgMar = new.find(f"{W}pgMar")
    if pgMar is None:
        pgMar = OxmlElement("w:pgMar")
        new.append(pgMar)
    pgMar.set(f"{W}top", FRONT_MATTER_TOP_MARGIN)
    # Section 1 should not restart page numbering; let section 2 keep its
    # start=1 so body page 1 corresponds to the first body page.
    pgNumType = new.find(f"{W}pgNumType")
    if pgNumType is not None:
        new.remove(pgNumType)
    # Explicit next-page section break
    typ = new.find(f"{W}type")
    if typ is None:
        typ = OxmlElement("w:type")
        new.insert(0, typ)
    typ.set(f"{W}val", "nextPage")
    return new


def insert_section_break_before_abstract(doc) -> bool:
    """Attach the section-1 sectPr to the paragraph immediately before the
    'Abstract' Heading 1. Returns True on success."""
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    body = doc.element.body
    last_sectPr = find_last_sectPr(body)
    if last_sectPr is None:
        print("[warn] no final sectPr found; skipping section split")
        return False
    s1 = build_section1_sectPr(last_sectPr)

    # Walk paragraphs to find the one right before 'Abstract' (Heading 1).
    paras = doc.paragraphs
    abstract_idx = None
    for i, p in enumerate(paras):
        if p.style.name == "Heading 1" and p.text.strip() == "Abstract":
            abstract_idx = i
            break
    if abstract_idx is None:
        print("[warn] Abstract heading not found")
        return False
    if abstract_idx == 0:
        print("[warn] Abstract is first paragraph; no room for section break")
        return False

    target = paras[abstract_idx - 1]
    # If this paragraph already carries a sectPr, replace it.
    pPr = _pPr(target)
    for existing in pPr.findall(qn("w:sectPr")):
        pPr.remove(existing)
    pPr.append(s1)
    print(f"[section] inserted section break on paragraph "
          f"#{abstract_idx - 1} ({target.style.name!r}): "
          f"top margin = {FRONT_MATTER_TOP_MARGIN} twips for front matter")
    return True


def remove_toc_manual_page_break(doc) -> None:
    """The existing TOC ends with a manual <w:br w:type='page'/>. With the
    section break now handling pagination, that manual break would leave an
    empty page. Remove it if present."""
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    removed = 0
    # Walk from start to Abstract heading
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() == "Abstract":
            break
        for r in p._element.findall(f"{W}r"):
            for br in r.findall(f"{W}br"):
                if br.get(f"{W}type") == "page":
                    r.remove(br)
                    removed += 1
    if removed:
        print(f"[page-break] removed {removed} manual page-break(s) in front matter")


def enforce_zero_spacing_everywhere(doc) -> None:
    n = 0
    for p in doc.paragraphs:
        set_zero_spacing(p)
        n += 1
    print(f"[spacing] zero before/after on {n} paragraphs")


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"missing: {DOCX}")

    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
        print(f"[backup] wrote {BACKUP.name}")
    else:
        print(f"[backup] {BACKUP.name} already present; not overwriting")

    doc = Document(DOCX)

    update_heading_and_toc_sizes(doc.styles.element)
    insert_section_break_before_abstract(doc)
    remove_toc_manual_page_break(doc)
    enforce_zero_spacing_everywhere(doc)

    doc.save(DOCX)
    print(f"[saved] {DOCX.name}")


if __name__ == "__main__":
    main()
