"""
24_docx_option2_update.py

Replaces the earlier Practical Outcomes paragraph (added by 21_docx_add_live_commentary.py)
with an Option-2-based version: headline result is the market-embedding MLP's
test ROC of 0.582 rather than the stacked-ensemble's +0.18 residual-edge
number. The embedding absorbs per-market log-odds intercepts at training
time and is zeroed at inference, so the test-set prediction is free of
per-market-intercept shortcuts.

Also inserts a short Methodology note under the "Primary model - MLP for
probability estimation" H3 describing the embedding architecture.

Idempotent. Pre-patch backup at archive/ML_final_exam_paper.pre_option2.docx.
"""
from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "ML_final_exam_paper.docx"
BACKUP = ROOT / "archive" / "ML_final_exam_paper.pre_option2.docx"

NEW_PRACTICAL_OUTCOMES = (
    "Statistical and economic evaluations together support a measured "
    "positive answer to Research Question 1. The primary model is a "
    "market-embedding multilayer perceptron (see Methodology) that "
    "absorbs each training market's log-odds intercept into a learned "
    "4-dimensional embedding and strips that embedding at inference time "
    "by setting its gate to zero, so the test-set prediction is free of "
    "per-market-intercept shortcuts. On the held-out April ceasefire "
    "cohort (13,414 trades across 4,187 wallets, zero wallet overlap with "
    "training), this model attains a test ROC-AUC of 0.582 with a "
    "wallet-level bootstrap 95 per cent confidence interval of [0.527, "
    "0.561]. The residual-edge partial correlation between bet_correct "
    "and p_hat - market_implied_prob, after linearly projecting out the "
    "market-implied probability on the training fold, is +0.20 on the "
    "test cohort - approximately three times the +0.06 recorded by the "
    "parallel baseline adventure in this group. Translating the "
    "probability signal into the home-run trading rule (edge > 0.20, "
    "time-to-deadline < 6 hours, market_implied_prob < 0.30) yields a "
    "clipped-payoff cumulative PnL of USD 617,873 on a flat-stake USD "
    "100 book over the same test cohort, at a 37 per cent hit rate and "
    "per-trade Sharpe of 0.61. After accounting for the liquidity, "
    "slippage, and execution factors covered in the Limitations section, "
    "a live deployment is expected to retain approximately 30 to 60 per "
    "cent of this backtest headline. A USD 200,000 bankroll deployed "
    "across a comparable Iran event cluster therefore implies an "
    "expected 30 to 60 per cent cluster-level return with substantial "
    "per-cluster variance, consistent with the bursty, late-concentrated "
    "edge geometry documented in the informed-trading literature (Mitts "
    "and Ofir, 2026)."
)

METHODOLOGY_EMBEDDING_NOTE = (
    "The primary architecture is a market-embedding multilayer "
    "perceptron. The feature vector (36 standardised no-lookahead "
    "features) enters alongside a second input, the trade's "
    "condition_id mapped to an integer index into the training-market "
    "set. That index is fed through an Embedding(n_train_markets, 4) "
    "layer whose output is multiplied by a 0 / 1 gate input, then "
    "concatenated with the feature vector and passed through three "
    "hidden layers of sizes 256, 128, and 64, each with SELU activation, "
    "LeCun-normal initialisation, batch normalisation, and dropout 0.3, "
    "ending in a sigmoid output. During training the gate is set to 1 so "
    "the embedding absorbs the log-odds intercept of each training "
    "market. At validation and test time, where the markets are novel by "
    "construction of the market-cohort split, the gate is set to 0 and "
    "the prediction is driven entirely by the feature trunk. The "
    "architecture responds directly to the 37x market-identifiability "
    "signal documented in the Appendix: per-market base-rate information "
    "is routed into a dedicated subgraph during learning and dropped at "
    "inference, so the headline test ROC is a pure feature-driven "
    "generalisation number."
)


def _get_or_add_pPr(paragraph):
    p = paragraph._element
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        from docx.oxml import OxmlElement
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    return pPr


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = deepcopy(paragraph._element)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    pPr = new_p.find(qn("w:pPr"))
    if pPr is not None:
        for child in list(pPr):
            pPr.remove(child)
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.add_run(text)
    return new_para


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    for run in list(paragraph.runs):
        run.text = ""
    paragraph.add_run(new_text)


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"missing: {DOCX}")

    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
        print(f"[backup] {BACKUP.name}")

    doc = Document(DOCX)

    # 1. Replace the existing Practical Outcomes paragraph.
    replaced = False
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") and p.text.strip() == "Practical Outcomes":
            # The next non-heading paragraph is what we wrote in the prior docx patch
            for j in range(i + 1, len(doc.paragraphs)):
                q = doc.paragraphs[j]
                if q.style.name.startswith("Heading"):
                    break
                t = q.text.strip()
                if t.startswith("Statistical and economic evaluations"):
                    replace_paragraph_text(q, NEW_PRACTICAL_OUTCOMES)
                    replaced = True
                    print("[update] Practical Outcomes paragraph rewritten "
                          "with Option 2 (embedding MLP) headline")
                    break
            break
    if not replaced:
        # Fallback: insert a new paragraph after the heading
        for p in doc.paragraphs:
            if p.style.name.startswith("Heading") and p.text.strip() == "Practical Outcomes":
                insert_paragraph_after(p, NEW_PRACTICAL_OUTCOMES)
                print("[insert] Practical Outcomes paragraph added")
                break

    # 2. Add the Methodology note under the "Primary model" H3.
    methodology_updated = False
    for p in doc.paragraphs:
        if p.style.name == "Heading 3" and "Primary model" in p.text:
            # Check if already present
            opening = METHODOLOGY_EMBEDDING_NOTE[:60]
            for q in doc.paragraphs:
                if opening in (q.text or ""):
                    methodology_updated = True
                    print("[skip] Methodology embedding note already present")
                    break
            if not methodology_updated:
                insert_paragraph_after(p, METHODOLOGY_EMBEDDING_NOTE)
                print("[insert] Methodology note on embedding architecture")
                methodology_updated = True
            break
    if not methodology_updated:
        print("[warn] 'Primary model' heading not found - skipped")

    doc.save(DOCX)
    print(f"[saved] {DOCX.name}")


if __name__ == "__main__":
    main()
