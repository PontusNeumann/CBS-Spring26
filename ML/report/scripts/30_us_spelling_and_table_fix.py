"""
30_us_spelling_and_table_fix.py

Two final corrections requested by the user after reviewing the docx:

1. **Bulletproof fix for orphaned Table A.6 and A.7 titles.** The previous
   keepNext + cantSplit attempt did not cure the orphan in Word. The
   reliable fix is to merge the title text INTO the first row of the
   table itself (single cell with gridSpan covering all columns, no
   borders), then delete the original standalone title paragraph. Word
   can never insert a page break inside a single row, so title and
   first data row are now physically inseparable; the rest of the long
   table is free to wrap across pages.

2. **UK -> US spelling pass.** The project standard is US English (now
   recorded in Design.md). Convert remaining -ise/-isation/-ised, -our,
   -re, and -ll- doublings in the docx body.

Run:
    python report/scripts/30_us_spelling_and_table_fix.py
"""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "ML_final_exam_paper.docx"
BACKUP = ROOT / "backup" / "ML_final_exam_paper.pre_us_table_fix.docx"


# ---------------------------------------------------------------------
# Spelling: ordered (longest-first to avoid prefix collisions)
# ---------------------------------------------------------------------
UK_US_PAIRS: list[tuple[str, str]] = [
    # -isation family (long forms first)
    ("parameterisations", "parameterizations"),
    ("parameterisation", "parameterization"),
    ("operationalisation", "operationalization"),
    ("characterisation", "characterization"),
    ("standardisation", "standardization"),
    ("residualisation", "residualization"),
    ("regularisation", "regularization"),
    ("normalisation", "normalization"),
    ("generalisation", "generalization"),
    ("initialisation", "initialization"),
    ("memorisation", "memorization"),
    ("monetisation", "monetization"),
    ("organisation", "organization"),
    ("realisation", "realization"),
    ("optimisation", "optimization"),
    ("visualisation", "visualization"),
    ("finalisation", "finalization"),
    ("categorisation", "categorization"),
    ("synthesisation", "synthesization"),
    ("hypothesisation", "hypothesization"),
    # -ised / -ising / -ise / -isable
    ("operationalised", "operationalized"),
    ("operationalising", "operationalizing"),
    ("operationalise", "operationalize"),
    ("characterised", "characterized"),
    ("characterising", "characterizing"),
    ("characterise", "characterize"),
    ("standardised", "standardized"),
    ("standardising", "standardizing"),
    ("standardise", "standardize"),
    ("residualised", "residualized"),
    ("residualising", "residualizing"),
    ("residualise", "residualize"),
    ("regularised", "regularized"),
    ("regularising", "regularizing"),
    ("regularise", "regularize"),
    ("unregularised", "unregularized"),
    ("normalised", "normalized"),
    ("normalising", "normalizing"),
    ("normalise", "normalize"),
    ("generalised", "generalized"),
    ("generalising", "generalizing"),
    ("generalise", "generalize"),
    ("initialised", "initialized"),
    ("initialise", "initialize"),
    ("memorised", "memorized"),
    ("memorising", "memorizing"),
    ("memorise", "memorize"),
    ("monetisable", "monetizable"),
    ("monetised", "monetized"),
    ("monetising", "monetizing"),
    ("monetise", "monetize"),
    ("organised", "organized"),
    ("organise", "organize"),
    ("realised", "realized"),
    ("realising", "realizing"),
    ("realise", "realize"),
    ("optimised", "optimized"),
    ("optimising", "optimizing"),
    ("optimise", "optimize"),
    ("visualised", "visualized"),
    ("visualising", "visualizing"),
    ("visualise", "visualize"),
    ("finalised", "finalized"),
    ("finalising", "finalizing"),
    ("finalise", "finalize"),
    ("categorised", "categorized"),
    ("categorise", "categorize"),
    ("annualised", "annualized"),
    ("annualise", "annualize"),
    ("summarised", "summarized"),
    ("summarising", "summarizing"),
    ("summarise", "summarize"),
    ("emphasised", "emphasized"),
    ("emphasising", "emphasizing"),
    ("emphasise", "emphasize"),
    ("recognised", "recognized"),
    ("recognising", "recognizing"),
    ("recognise", "recognize"),
    ("prioritised", "prioritized"),
    ("prioritise", "prioritize"),
    ("minimised", "minimized"),
    ("minimises", "minimizes"),
    ("minimising", "minimizing"),
    ("minimise", "minimize"),
    ("maximised", "maximized"),
    ("maximises", "maximizes"),
    ("maximising", "maximizing"),
    ("maximise", "maximize"),
    ("utilised", "utilized"),
    ("utilising", "utilizing"),
    ("utilise", "utilize"),
    ("synthesised", "synthesized"),
    ("synthesise", "synthesize"),
    ("hypothesised", "hypothesized"),
    ("hypothesise", "hypothesize"),
    # -our -> -or
    ("behavioural", "behavioral"),
    ("behaviour", "behavior"),
    ("colour", "color"),
    ("coloured", "colored"),
    ("favour", "favor"),
    ("favoured", "favored"),
    ("neighbour", "neighbor"),
    # -re -> -er  (only forms that show up in this corpus context)
    ("centred", "centered"),
    ("centring", "centering"),
    # Doubled-l (UK) -> single-l (US) for unstressed second syllable
    ("modelling", "modeling"),
    ("modelled", "modeled"),
    ("labelling", "labeling"),
    ("labelled", "labeled"),
    ("travelling", "traveling"),
    ("travelled", "traveled"),
    ("cancelling", "canceling"),
    ("cancelled", "canceled"),
]


def case_preserving_replace(text: str, uk: str, us: str) -> tuple[str, int]:
    """Replace `uk` with `us` in `text`, preserving the leading-letter case
    of each match (e.g. Normalisation -> Normalization, normalisation ->
    normalization). Returns (new_text, count)."""
    pattern = re.compile(r"\b" + re.escape(uk) + r"\b", re.IGNORECASE)
    count = 0

    def repl(m):
        nonlocal count
        count += 1
        match = m.group(0)
        if match[:1].isupper():
            return us[:1].upper() + us[1:]
        return us

    return pattern.sub(repl, text), count


def apply_spelling_to_runs(p_elem, replacements: list[tuple[str, str]]) -> dict[str, int]:
    """Apply spelling pairs to every <w:t> inside the paragraph. Tracks
    counts per pair."""
    counts: dict[str, int] = {}
    for t in p_elem.iter(qn("w:t")):
        if not t.text:
            continue
        new_text = t.text
        for uk, us in replacements:
            new_text, n = case_preserving_replace(new_text, uk, us)
            if n:
                counts[uk] = counts.get(uk, 0) + n
        if new_text != t.text:
            t.text = new_text
    return counts


# ---------------------------------------------------------------------
# Table title -> first row conversion
# ---------------------------------------------------------------------
def _set_no_borders(tcPr) -> None:
    """Set all cell borders to nil so the caption row reads as standalone text."""
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    # Remove any existing
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def _count_grid_cols(tbl) -> int:
    """Count columns from <w:tblGrid>/<w:gridCol>."""
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        # Fallback: count cells in first row
        first_tr = tbl.find(qn("w:tr"))
        return len(first_tr.findall(qn("w:tc"))) if first_tr is not None else 1
    return len(grid.findall(qn("w:gridCol")))


def convert_title_to_caption_row(title_para_elem, table_elem) -> None:
    """Move the title paragraph into a new merged-cell first row of the table,
    then delete the original title paragraph (and the empty paragraph immediately
    before it, if any, to avoid leaving a stray blank line)."""
    n_cols = _count_grid_cols(table_elem)

    # Build new row: trPr (cantSplit), single tc with gridSpan, no borders, copied paragraph
    new_tr = OxmlElement("w:tr")
    trPr = OxmlElement("w:trPr")
    trPr.append(OxmlElement("w:cantSplit"))
    new_tr.append(trPr)

    new_tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    gridSpan = OxmlElement("w:gridSpan")
    gridSpan.set(qn("w:val"), str(n_cols))
    tcPr.append(gridSpan)
    _set_no_borders(tcPr)
    new_tc.append(tcPr)

    # Clone the title paragraph as the cell's paragraph; ensure keepNext (so
    # the caption stays with the data row directly below it inside the table).
    para_clone = deepcopy(title_para_elem)
    pPr = para_clone.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para_clone.insert(0, pPr)
    if pPr.find(qn("w:keepNext")) is None:
        pPr.append(OxmlElement("w:keepNext"))
    if pPr.find(qn("w:keepLines")) is None:
        pPr.append(OxmlElement("w:keepLines"))
    new_tc.append(para_clone)

    new_tr.append(new_tc)

    # Insert the new row as the FIRST row of the table (after tblPr/tblGrid)
    insert_after = table_elem.find(qn("w:tblGrid"))
    if insert_after is None:
        insert_after = table_elem.find(qn("w:tblPr"))
    if insert_after is None:
        table_elem.insert(0, new_tr)
    else:
        insert_after.addnext(new_tr)

    # Remove the original title paragraph
    parent = title_para_elem.getparent()
    # Also remove a single immediately-preceding empty paragraph, if present
    prev = title_para_elem.getprevious()
    parent.remove(title_para_elem)
    if prev is not None and prev.tag == qn("w:p"):
        prev_text = "".join((t.text or "") for t in prev.iter(qn("w:t")))
        if prev_text.strip() == "":
            parent.remove(prev)


# ---------------------------------------------------------------------
# Char count (matches earlier convention: body + tables, excl ToC + refs + appendix)
# ---------------------------------------------------------------------
def count_body_chars(doc) -> int:
    body = doc.element.body
    chars = 0
    in_toc = True
    stop = False
    for elem in body.iterchildren():
        if stop:
            break
        if elem.tag == qn("w:p"):
            ppr = elem.find(qn("w:pPr"))
            style = None
            if ppr is not None:
                ps = ppr.find(qn("w:pStyle"))
                if ps is not None:
                    style = ps.get(qn("w:val"))
            text = "".join(t.text or "" for t in elem.iter(qn("w:t")))
            if in_toc:
                if style == "Heading1" and text.strip() == "Abstract":
                    in_toc = False
                else:
                    continue
            if style == "Heading1" and text.strip() in ("References", "Appendix"):
                stop = True
                break
            chars += len(text)
        elif elem.tag == qn("w:tbl"):
            if in_toc:
                continue
            for t in elem.iter(qn("w:t")):
                chars += len(t.text or "")
    return chars


def update_cover_char_count(doc, count: int) -> None:
    placeholder_pat = re.compile(r"^\s*\d[\d,]*\s+characters incl\. spaces$|^\s*XX,XXX characters incl\. spaces$")
    for r in doc.element.body.iter(qn("w:r")):
        for t in r.findall(qn("w:t")):
            if t.text and placeholder_pat.match(t.text):
                t.text = f" {count:,} characters incl. spaces"
                return


# ---------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------
def main() -> None:
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOCX, BACKUP)
    print(f"backup -> {BACKUP}")

    doc = Document(DOCX)
    body = doc.element.body
    elems = list(body.iterchildren())

    # ---- 1) Table title -> first row for A.6, A.7 ----
    targets = []  # (title_idx, table_idx, label)
    for i, e in enumerate(elems):
        if e.tag == qn("w:p"):
            text = "".join(t.text or "" for t in e.iter(qn("w:t"))).strip()
            if text.startswith("Table A.6.") or text.startswith("Table A.7."):
                # find next w:tbl
                j = i + 1
                while j < len(elems) and elems[j].tag != qn("w:tbl"):
                    if elems[j].tag == qn("w:p"):
                        nt = "".join(t.text or "" for t in elems[j].iter(qn("w:t"))).strip()
                        if nt:
                            break
                    j += 1
                if j < len(elems) and elems[j].tag == qn("w:tbl"):
                    targets.append((i, j, text[:12]))

    if len(targets) != 2:
        print(f"WARN: expected 2 targets, found {len(targets)}")

    # Process in reverse order so earlier indices remain valid after deletes
    for title_idx, table_idx, label in reversed(targets):
        convert_title_to_caption_row(elems[title_idx], elems[table_idx])
        print(f"  merged title->row: {label}")

    # ---- 2) Spelling pass on every paragraph (including in tables) ----
    total_counts: dict[str, int] = {}
    # Walk every paragraph in the document
    for p in doc.paragraphs:
        c = apply_spelling_to_runs(p._p, UK_US_PAIRS)
        for k, v in c.items():
            total_counts[k] = total_counts.get(k, 0) + v
    # Walk paragraphs inside table cells too
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    c = apply_spelling_to_runs(p._p, UK_US_PAIRS)
                    for k, v in c.items():
                        total_counts[k] = total_counts.get(k, 0) + v

    if total_counts:
        print("Spelling replacements:")
        for k in sorted(total_counts, key=lambda x: -total_counts[x]):
            print(f"  {k:30s} -> {dict(UK_US_PAIRS).get(k, '?'):30s} x{total_counts[k]}")
    else:
        print("No UK spellings replaced.")

    # ---- 3) Recompute and update cover-page char count ----
    body_chars = count_body_chars(doc)
    update_cover_char_count(doc, body_chars)
    print(f"cover-page char count -> {body_chars:,} characters incl. spaces")

    doc.save(DOCX)
    print(f"saved -> {DOCX}")


if __name__ == "__main__":
    main()
