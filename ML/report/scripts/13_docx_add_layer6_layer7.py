"""One-shot: add Layer 6, Layer 7, and missing-data coverage to the docx.

Surgical edits on ML_final_exam_paper.docx:
  - Para 89: replace the outdated "No Polygonscan or external on-chain enrichment is used." sentence.
  - Para 93: rewrite the features paragraph to describe seven layers including Layer 6 (on-chain
    identity) and Layer 7 (cross-market diversity), removing the deferred-to-future-work sentence.
  - After para 93: insert a new paragraph documenting the missingness treatment.

Run once:
  /Applications/anaconda3/envs/py312/bin/python scripts/13_docx_add_layer6_layer7.py
"""

from pathlib import Path
from copy import deepcopy

from docx import Document

DOCX = Path(__file__).resolve().parents[1] / "ML_final_exam_paper.docx"

OLD_89_SENTENCE = "No Polygonscan or external on-chain enrichment is used."
NEW_89_SENTENCE = (
    "An additional on-chain enrichment pipeline queries the Etherscan V2 API for "
    "each wallet's full Polygon token-transfer history, producing time-invariant "
    "scalars and per-wallet timestamp arrays that are bisected at trade time to "
    "build the Layer 6 features described below."
)

NEW_93_TEXT = (
    "Features are organised in seven layers, all strictly no-lookahead. "
    "Trade-local features cover log size, side, and outcome index; the "
    "contemporaneous price is deliberately excluded from the feature set so that "
    "p_hat is independent of the market's own belief and the residual p_hat "
    "minus market_implied_prob is a clean signal. Market-context features "
    "include cumulative trade count and volume, one-hour rolling price "
    "volatility, log-transformed trailing one-hour and twenty-four-hour USD "
    "volume, and the running buy share of the market. Time features are "
    "time-to-settlement with log variant and percent of market lifetime "
    "elapsed. Wallet-global features cover prior trades, prior volume, prior "
    "win rate, and a wallet-age-at-trade proxy. Wallet-in-market features form "
    "three behavioural clusters with specific informed-flow interpretations. "
    "The bet-slicing cluster (wallet_trades_in_market_last_1min, 10min, 60min, "
    "wallet_cumvol_same_side_last_10min, wallet_is_burst, "
    "wallet_prior_trades_in_market, wallet_median_gap_in_market) captures "
    "informed BUYs split across many small rapid-fire trades to avoid moving "
    "the price. The spread-builder cluster (wallet_directional_purity_in_market, "
    "wallet_has_both_sides_in_market, wallet_spread_ratio) captures "
    "volatility-trading wallets buying both sides, which provide an uninformed "
    "negative class. The whale-exit cluster (wallet_position_size_before_trade, "
    "trade_size_vs_position_pct, is_position_exit, is_position_flip, "
    "wallet_is_whale_in_market) distinguishes large full-position SELLs near "
    "deadline from retail profit-takes. Interaction features are trade size "
    "relative to wallet average, the log-size by log-time-to-settlement "
    "product, trade size as a fraction of prior market cumulative volume, and "
    "trade size relative to the market's running average. Layer 6 on-chain "
    "identity is derived from each wallet's full Polygon token-transfer "
    "history and contributes nine causal columns: wallet age at trade, "
    "Polygon transaction nonce and its log, count of inbound transfers and "
    "its log, count of CEX-originated USDC deposits, cumulative CEX USDC "
    "volume and its log, days since first USDC inflow, a binary indicator of "
    "first-funding from a known CEX hot wallet, and a trade-time-scoped "
    "variant of that indicator. Layer 7 cross-market diversity is a single "
    "feature, the Shannon entropy in nats over the coarse category "
    "distribution of a wallet's prior distinct markets across the full "
    "Polymarket universe, streamed from the HuggingFace mirror and bucketed "
    "into eight domains. GDELT news-lag features remain scoped as future "
    "work. Every feature at row t uses only rows with timestamp strictly "
    "before t; the market-implied probability is retained only as the "
    "trading-rule benchmark."
)

MISSINGNESS_TEXT = (
    "Missingness in the feature frame is handled by the underlying "
    "semantics rather than by imputation at the feature stage. Running and "
    "expanding features are NaN on the first occurrence of a wallet or a "
    "wallet-in-market pair; Layer 6 features are NaN on rows whose wallet "
    "failed Etherscan V2 enrichment; pct_time_elapsed is NaN on the small "
    "number of markets missing both resolution_ts and end_date metadata. "
    "Five binary indicators carry the missingness signal into the feature "
    "frame: wallet_has_prior_trades, wallet_has_prior_trades_in_market, "
    "wallet_has_cross_market_history, market_timing_known, and "
    "wallet_enriched. Imputation is applied only at the modelling stage and "
    "differs by classifier. Tree-based models accept NaN natively in "
    "scikit-learn 1.4+ and consume raw columns; logistic regression and the "
    "MLP receive train-split-only median imputation followed by "
    "standardisation. In every case the indicator columns are retained as "
    "features, so whatever value is imputed for the raw feature the model "
    "still sees the was-this-missing signal."
)


def main() -> None:
    doc = Document(str(DOCX))

    # --- Para 89: replace outdated sentence in run 0 ---
    p89 = doc.paragraphs[89]
    r0 = p89.runs[0]
    assert OLD_89_SENTENCE in r0.text, "expected sentence not found in para 89 run 0"
    r0.text = r0.text.replace(OLD_89_SENTENCE, NEW_89_SENTENCE)
    print("para 89: updated pipeline sentence")

    # --- Para 93: rewrite in place (single run, no inline formatting) ---
    p93 = doc.paragraphs[93]
    assert len(p93.runs) == 1, f"expected 1 run in para 93, got {len(p93.runs)}"
    p93.runs[0].text = NEW_93_TEXT
    print("para 93: rewrote features paragraph (6 -> 7 layers, Layer 6 + Layer 7 described)")

    # --- Insert a new paragraph after para 93 with the missingness text ---
    # Clone para 93's XML to inherit style, then replace text and insert.
    new_p = deepcopy(p93._element)
    # Strip all existing runs from the clone
    for r in new_p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
        new_p.remove(r)
    p93._element.addnext(new_p)
    # Re-wrap as a Paragraph object to use the high-level API
    from docx.text.paragraph import Paragraph
    inserted = Paragraph(new_p, p93._parent)
    inserted.add_run(MISSINGNESS_TEXT)
    print("inserted new paragraph (missingness treatment) after para 93")

    doc.save(str(DOCX))
    print(f"saved -> {DOCX}")


if __name__ == "__main__":
    main()
