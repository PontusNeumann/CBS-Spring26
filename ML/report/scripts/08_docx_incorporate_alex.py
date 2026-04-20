"""Incorporate pre-EDA sharpeners (approved plan, 2026-04-20).

Scope: everything before Exploratory Data Analysis, keeping paragraph
footprint roughly constant (15-page budget).

Edits:
  1. Abstract — one sentence grounding in Mitts & Ofir anomaly.
  2. Motivation (§2.1) — replace generic efficient-market opening with
     three concrete anchors: Columbia USD 143M, Magamyman USD 553K,
     and Iran-strikes USD 1.2M; one sentence on the Maduro case as
     pattern evidence.
  3. Related Work (§2.3) — add Vergence / Palantir / TWG AI novelty
     sentence (March 10, 2026).
  4. Dataset Description (§5.1) — insert compressed four-event table
     (event ID, theme, sub-markets, trades, volume, YES/NO mix).
  5. Target label (§5.1 bet_correct paragraph) — one-line defense of
     simple directional over profitable-PnL, informed-threshold, and
     wallet-level variants.
  6. Features (§5.1 six-layer paragraph) — adopt bet-slicing /
     spread-builder / whale-exit cluster labels (mapping Alex's
     naming onto our existing columns).
  7. Split paragraphs (both the Problem Statement and Methodology
     copies) — add explicit event-leakage acknowledgment and defend
     against the objection (4 events, within-market detection scope,
     not future-event forecasting).
  8. References — add Polymarket / Palantir / TWG AI partnership
     announcement (March 10, 2026) and Gambling Insider coverage of
     the Iran-strike insider profits (March 11, 2026).

One-shot script; re-running will duplicate insertions. Idempotency
guards check for sentinel strings before inserting.
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "ML_final_exam_paper.docx"


# ---------------------------------------------------------------------------
# Replacement texts
# ---------------------------------------------------------------------------
ABSTRACT_SENTINEL = "Prediction markets on Polymarket price real-world events"
ABSTRACT_NEW = (
    "Prediction markets on Polymarket price real-world events continuously, yet "
    "whether those prices are informationally efficient remains an open "
    "question. Prior work documents USD 143 million in anomalous profits "
    "concentrated on geopolitical contracts, with named individual wallets "
    "operating across Iran-strike markets in particular (Mitts and Ofir, 2026). "
    "This project tests whether a machine learning model trained on "
    "pre-execution market-state and behavioural features produces a probability "
    "estimate that systematically diverges from the contemporaneous "
    "market-implied probability on resolved Iran geopolitical markets, and "
    "whether the resulting gap is tradable. The scope covers all seventy-four "
    "resolved sub-markets under Polymarket events 114242, 236884, 355299, and "
    "357625, with trade correctness derived from official resolution. A "
    "multilayer perceptron estimates a per-trade probability p_hat from "
    "standardised features covering market state, recent activity, and wallet "
    "history. Logistic regression, random forest, and Isolation Forest serve as "
    "baselines, and an undercomplete autoencoder provides an unsupervised "
    "anomaly lens. Evaluation proceeds on temporally held-out data along two "
    "axes. Statistical quality uses ROC-AUC, Brier score, and calibration of "
    "p_hat. Economic quality evaluates a gap-threshold trading rule against "
    "naive market-implied, momentum, and random-entry baselines on cumulative "
    "PnL, Sharpe ratio, hit rate, and drawdown. Findings indicate whether "
    "Polymarket's Iran geopolitical markets admit a systematic mispricing "
    "signal recoverable from behavioural features alone, informing both "
    "market-efficiency theory and applied prediction-market trading strategy."
)

MOTIV_P1_SENTINEL = "efficient-market null states"
MOTIV_P1_NEW = (
    "Polymarket's efficient-market null states that the contemporaneous price "
    "of a binary contract already incorporates all public information available "
    "at that instant. The null is empirically contestable on this platform. "
    "Mitts and Ofir (2026) document USD 143 million in anomalous profits "
    "concentrated on geopolitical contracts and identify the individual wallets "
    "responsible, including one wallet earning USD 553 thousand on the "
    "Magamyman case. Approximately USD 1.2 million of those profits were "
    "realised across the Iran-strike markets studied here, with trades placed "
    "hours before strikes were publicly reported (Gambling Insider, 2026). A "
    "comparable Polymarket account earned USD 400 thousand on Venezuelan "
    "leadership markets shortly before reports of a United States operation "
    "targeting Maduro, indicating that documented insider trading on Polymarket "
    "has been a geopolitical-markets phenomenon rather than an isolated case. "
    "A systematic gap between a model-predicted probability p_hat and the "
    "market-implied probability, if it exists and predicts settlement, is "
    "direct evidence of this mispricing channel and can in principle be "
    "monetised."
)

RELATED_WORK_VERGENCE = (
    "The novelty gap is dated and verifiable. Polymarket announced a "
    "prediction-market integrity partnership with Palantir and TWG AI on 10 "
    "March 2026, and the accompanying Vergence surveillance system is "
    "explicitly scoped to sports contracts. No equivalent public machine "
    "learning baseline exists for geopolitical sub-markets, which is where "
    "every documented insider-trading episode on Polymarket has occurred. "
    "This project provides such a baseline for the Iran cluster."
)

DATASET_P1_SENTINEL = "dataset consists of Polymarket on-chain"
DATASET_P1_NEW = (
    "The dataset consists of Polymarket on-chain trade records for all "
    "seventy-four resolved sub-markets under four Iran-related events: "
    "114242 (US strikes Iran by [date]), 236884 (Iran x Israel or US conflict "
    "ends by [date]), 355299 (Trump announces US x Iran ceasefire end by "
    "[date]), and 357625 (US x Iran ceasefire extended by [date]). The full "
    "resolved universe is used rather than a hand-picked subset, covering "
    "both YES and NO outcomes. Aggregate composition is reported in Table 1."
)

# Event ID | Theme | Sub-markets | Trades | Volume (USD) | YES / NO
DATASET_TABLE = {
    "label": "Table 1.",
    "caption": "Iran cluster composition by Polymarket event. Aggregates "
               "computed on resolved trades only.",
    "header": ["Event ID", "Theme", "Sub-markets", "Trades", "Volume (USD)", "YES / NO"],
    "rows": [
        ["114242", "US strikes Iran by [date]",          "64", "1,145,584", "244.7M", "45 / 19"],
        ["236884", "Iran-Israel/US conflict ends",       "3",  "49,563",    "4.3M",   "3 / 0"],
        ["355299", "Trump announces ceasefire end",      "5",  "12,644",    "1.9M",   "5 / 0"],
        ["357625", "US-Iran ceasefire extended",         "2",  "1,996",     "0.3M",   "2 / 0"],
        ["Total",  "",                                   "74", "1,209,787", "251.1M", "55 / 19"],
    ],
}

LABEL_SENTINEL = "target variable is bet_correct"
LABEL_NEW = (
    "The target variable is bet_correct in {0, 1}, derived from the market "
    "resolution and the side of the trade, following the directional "
    "correctness convention used in prediction-market research (Wolfers and "
    "Zitzewitz, 2004). Three alternatives were considered and rejected. "
    "Profitable-PnL at resolution collapses to the same binary label at trade "
    "level, because any BUY executed below price one on the winning token is "
    "profitable by construction. Informed-threshold labels require an "
    "arbitrary entry-price cutoff and discard most trades from the positive "
    "class. Wallet-level propagation induces retroactive label leakage across "
    "a wallet's own earlier trades. The benchmark at trade time is "
    "market_implied_prob, taken from the CLOB mid-price or the price field of "
    "the trade itself."
)

FEATURES_SENTINEL = "Features are organised in six layers"
FEATURES_NEW = (
    "Features are organised in six layers, all strictly no-lookahead. "
    "Trade-local features cover log size, side, and outcome index; the "
    "contemporaneous price is deliberately excluded from the feature set so "
    "that p_hat is independent of the market's own belief and the residual "
    "p_hat minus market_implied_prob is a clean signal. Market-context "
    "features include cumulative trade count and volume, one-hour rolling "
    "price volatility, and log-transformed trailing one-hour and "
    "twenty-four-hour USD volume. Time features are time-to-settlement with "
    "log variant and percent of market lifetime elapsed. Wallet-global "
    "features cover prior trades, prior volume, prior win rate, and a "
    "wallet-age-at-trade proxy. Wallet-in-market features form three "
    "behavioural clusters with specific informed-flow interpretations. The "
    "bet-slicing cluster (wallet_trades_in_market_last_1min, 10min, 60min, "
    "wallet_cumvol_same_side_last_10min, wallet_is_burst, "
    "wallet_prior_trades_in_market) captures informed BUYs split across many "
    "small rapid-fire trades to avoid moving the price. The spread-builder "
    "cluster (wallet_directional_purity_in_market, "
    "wallet_has_both_sides_in_market, wallet_spread_ratio) captures "
    "volatility-trading wallets buying both sides, which provide an "
    "uninformed negative class. The whale-exit cluster "
    "(wallet_position_size_before_trade, trade_size_vs_position_pct, "
    "is_position_exit, is_position_flip, wallet_is_whale_in_market) "
    "distinguishes large full-position SELLs near deadline from retail "
    "profit-takes. Interaction features are trade size relative to wallet "
    "average, the log-size by log-time-to-settlement product, and trade size "
    "as a fraction of prior market cumulative volume. Polygon on-chain "
    "identity and GDELT news-lag layers are scoped as future work and not "
    "used in the current dataset. Every feature at row t uses only rows with "
    "timestamp strictly before t; the market-implied probability is retained "
    "only as the trading-rule benchmark."
)

SPLIT_METHODOLOGY_SENTINEL = (
    "The split is trade-timestamp temporal. Each trade is assigned to "
    "train, validation, or test based on its own execution timestamp, with "
    "quantile boundaries"
)
SPLIT_METHODOLOGY_NEW = (
    "The split is trade-timestamp temporal. Each trade is assigned to train, "
    "validation, or test based on its own execution timestamp, with quantile "
    "boundaries at 0.70 and 0.85 of the full trade-timestamp range. A "
    "trade-timestamp split was preferred over a settlement-date split because "
    "events 114242 and 355299 cluster their NO and YES resolutions on "
    "different calendar dates; a settlement-date split would place early "
    "NO-resolving markets in training and later YES-resolving markets in "
    "test, leaving the model with no in-sample YES outcomes. Overall "
    "bet_correct rate is 0.504, well inside the 35 to 65 percent band. The "
    "tradeoff is event leakage: because the four Polymarket events supply "
    "markets to every fold, the model can learn event-specific wallet and "
    "microstructure patterns rather than event-transferable insider-flow "
    "signatures. This is accepted because the research question is "
    "within-market detection of informed flow given observed event structure, "
    "not forecasting of novel geopolitical events; with only four events "
    "available, an event-holdout design would collapse the training set and "
    "foreclose the primary evaluation."
)

SPLIT_PROBLEM_SENTINEL = (
    "The split is trade-timestamp temporal. Each trade is assigned to "
    "train, validation, or test based on its own execution timestamp, not "
    "its market's settlement date"
)
SPLIT_PROBLEM_NEW = (
    "The split is trade-timestamp temporal. Each trade is assigned to train, "
    "validation, or test based on its own execution timestamp, not its "
    "market's settlement date. A settlement-date split would cluster early "
    "NO-resolving markets in training and late YES-resolving markets in test, "
    "leaving the model with no in-sample YES outcomes. A trade-timestamp "
    "split produces outcome-mixed folds because pre-deadline trades on "
    "eventually-YES markets exist well before those markets settle. Trades "
    "from the same event necessarily appear across folds, which is accepted "
    "on the basis that the research question is within-market detection of "
    "informed flow, not future-event forecasting; an event-holdout design is "
    "infeasible with only four events. Each trade retains its own market's "
    "outcome as the label. Rule parameters are tuned on the validation slice "
    "and frozen for the test slice."
)


# ---------------------------------------------------------------------------
# References to add
# ---------------------------------------------------------------------------
# APA 7, alphabetical. Inserted so the full list stays sorted.
# New entries: Gambling Insider (after Coplan, before Goodfellow) and
# Polymarket March 10 announcement (before the March 23 entry already present).
NEW_REFERENCES = {
    "Gambling Insider": [
        ("Gambling Insider. (2026, March 11). ", False),
        ("Polymarket partners with Palantir to curb suspicious trading on Iran-strike markets", True),
        (". Gambling Insider.", False),
    ],
    "Polymarket March 10": [
        ("Polymarket. (2026, March 10). ", False),
        ("Prediction-market integrity partnership with Palantir and TWG AI", True),
        (" [Company announcement]. Polymarket.", False),
    ],
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def set_para_text(para, text: str) -> None:
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    para.add_run(text)


def find_first(doc, sentinel: str):
    for p in doc.paragraphs:
        if sentinel in (p.text or ""):
            return p
    raise LookupError(f"paragraph with sentinel {sentinel!r} not found")


def find_all(doc, sentinel: str):
    return [p for p in doc.paragraphs if sentinel in (p.text or "")]


def insert_paragraph_after(anchor_para, text: str = "", style: str = "Normal"):
    new_p = OxmlElement("w:p")
    if style:
        pPr = OxmlElement("w:pPr")
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), style.replace(" ", ""))
        pPr.append(ps)
        new_p.append(pPr)
    if text:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        new_p.append(r)
    anchor_para._p.addnext(new_p)
    return new_p


def insert_paragraph_before(anchor_para, text: str = "", style: str = "Normal"):
    new_p = OxmlElement("w:p")
    if style:
        pPr = OxmlElement("w:pPr")
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), style.replace(" ", ""))
        pPr.append(ps)
        new_p.append(pPr)
    if text:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        new_p.append(r)
    anchor_para._p.addprevious(new_p)
    return new_p


def build_table(header, rows):
    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tblLook = OxmlElement("w:tblLook")
    tblLook.set(qn("w:val"), "04A0")
    tblPr.append(tblLook)
    tbl.append(tblPr)

    tblGrid = OxmlElement("w:tblGrid")
    for _ in header:
        tblGrid.append(OxmlElement("w:gridCol"))
    tbl.append(tblGrid)

    def make_cell(text: str, bold: bool = False):
        tc = OxmlElement("w:tc")
        tc.append(OxmlElement("w:tcPr"))
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        if bold:
            rPr = OxmlElement("w:rPr")
            rPr.append(OxmlElement("w:b"))
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    tr = OxmlElement("w:tr")
    for h in header:
        tr.append(make_cell(h, bold=True))
    tbl.append(tr)
    for row in rows:
        tr = OxmlElement("w:tr")
        for v in row:
            tr.append(make_cell(str(v)))
        tbl.append(tr)
    return tbl


def insert_reference_before(target_para, parts):
    style_id = "ReferenceText"
    try:
        _ = target_para.part.document.styles["Reference Text"]
    except KeyError:
        style_id = "Normal"
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), style_id)
    pPr.append(ps)
    new_p.append(pPr)
    for text, italic in parts:
        r = OxmlElement("w:r")
        if italic:
            rPr = OxmlElement("w:rPr")
            rPr.append(OxmlElement("w:i"))
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        new_p.append(r)
    target_para._p.addprevious(new_p)
    return new_p


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    doc = Document(str(DOCX))

    # 1. Abstract
    abs_p = find_first(doc, ABSTRACT_SENTINEL)
    if "USD 143 million" not in abs_p.text:
        set_para_text(abs_p, ABSTRACT_NEW)
        print("[1] Abstract sharpened")
    else:
        print("[1] Abstract already has anomaly sentence; skipped")

    # 2. Motivation paragraph 1
    motiv_p = find_first(doc, MOTIV_P1_SENTINEL)
    if "Mitts and Ofir (2026) document" not in motiv_p.text:
        set_para_text(motiv_p, MOTIV_P1_NEW)
        print("[2] Motivation paragraph sharpened")
    else:
        print("[2] Motivation already has anomaly figures; skipped")

    # 3. Related Work — add Vergence paragraph after Mitts & Ofir paragraph
    mitts_p = find_first(doc, "Work on informed trading in financial")
    next_p = mitts_p._p.getnext()
    next_is_vergence = (next_p is not None and next_p.tag == qn("w:p")
                        and "Vergence" in "".join(t.text or "" for t in next_p.iter(qn("w:t"))))
    if not next_is_vergence:
        # Insert a blank Normal + the Vergence paragraph after Mitts & Ofir
        insert_paragraph_after(mitts_p, "", style="Normal")
        # After inserting a blank, the Mitts paragraph's next-sibling is that
        # blank. We want the Vergence paragraph AFTER the blank, so insert
        # after the blank.
        blank = mitts_p._p.getnext()
        # Now insert the Vergence paragraph after the blank.
        vergence_xml = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), "Normal")
        pPr.append(ps)
        vergence_xml.append(pPr)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = RELATED_WORK_VERGENCE
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        vergence_xml.append(r)
        blank.addnext(vergence_xml)
        print("[3] Related Work — Vergence paragraph inserted")
    else:
        print("[3] Related Work already has Vergence paragraph; skipped")

    # 4. Dataset Description paragraph + table
    ds_p = find_first(doc, DATASET_P1_SENTINEL)
    if "Aggregate composition is reported in Table 1" not in ds_p.text:
        set_para_text(ds_p, DATASET_P1_NEW)

        # Insert caption + table after this paragraph (before the next one).
        next_sib = ds_p._p.getnext()  # typically a blank Normal

        # Build: blank, caption, table, blank, then reinsert original next_sib
        # We insert directly after ds_p in reverse order (so final order is
        # blank / caption / table / blank).
        trailing_blank = OxmlElement("w:p")
        tb_pPr = OxmlElement("w:pPr"); tb_ps = OxmlElement("w:pStyle")
        tb_ps.set(qn("w:val"), "Normal"); tb_pPr.append(tb_ps); trailing_blank.append(tb_pPr)

        tbl = build_table(DATASET_TABLE["header"], DATASET_TABLE["rows"])

        cap = OxmlElement("w:p")
        c_pPr = OxmlElement("w:pPr"); c_ps = OxmlElement("w:pStyle")
        c_ps.set(qn("w:val"), "Normal"); c_pPr.append(c_ps); cap.append(c_pPr)
        # Caption: bold label + italic body
        r_lbl = OxmlElement("w:r")
        rPr_b = OxmlElement("w:rPr"); rPr_b.append(OxmlElement("w:b"))
        r_lbl.append(rPr_b)
        t_lbl = OxmlElement("w:t"); t_lbl.text = f'{DATASET_TABLE["label"]} '
        t_lbl.set(qn("xml:space"), "preserve"); r_lbl.append(t_lbl); cap.append(r_lbl)
        r_body = OxmlElement("w:r")
        rPr_i = OxmlElement("w:rPr"); rPr_i.append(OxmlElement("w:i"))
        r_body.append(rPr_i)
        t_body = OxmlElement("w:t"); t_body.text = DATASET_TABLE["caption"]
        t_body.set(qn("xml:space"), "preserve"); r_body.append(t_body); cap.append(r_body)

        leading_blank = OxmlElement("w:p")
        lb_pPr = OxmlElement("w:pPr"); lb_ps = OxmlElement("w:pStyle")
        lb_ps.set(qn("w:val"), "Normal"); lb_pPr.append(lb_ps); leading_blank.append(lb_pPr)

        ds_p._p.addnext(trailing_blank)
        ds_p._p.addnext(tbl)
        ds_p._p.addnext(cap)
        ds_p._p.addnext(leading_blank)
        print("[4] Dataset Description paragraph + event-breakdown table inserted")
    else:
        print("[4] Dataset Description table already inserted; skipped")

    # 5. Target label
    lbl_p = find_first(doc, LABEL_SENTINEL)
    if "directional correctness convention" not in lbl_p.text:
        set_para_text(lbl_p, LABEL_NEW)
        print("[5] Target label paragraph sharpened")
    else:
        print("[5] Target label already sharpened; skipped")

    # 6. Features — six-layer paragraph with behavioural cluster labels
    feat_p = find_first(doc, FEATURES_SENTINEL)
    if "bet-slicing cluster" not in feat_p.text:
        set_para_text(feat_p, FEATURES_NEW)
        print("[6] Feature clusters renamed to behavioural labels")
    else:
        print("[6] Feature clusters already renamed; skipped")

    # 7. Split paragraphs — both copies
    try:
        split_m = find_first(doc, SPLIT_METHODOLOGY_SENTINEL)
        if "event leakage" not in split_m.text:
            set_para_text(split_m, SPLIT_METHODOLOGY_NEW)
            print("[7a] Methodology split paragraph: event-leakage defense added")
        else:
            print("[7a] Methodology split already addresses event leakage; skipped")
    except LookupError:
        print("[7a] Methodology split sentinel not found; skipped")

    try:
        split_p = find_first(doc, SPLIT_PROBLEM_SENTINEL)
        if "event-holdout" not in split_p.text:
            set_para_text(split_p, SPLIT_PROBLEM_NEW)
            print("[7b] Problem-Statement split paragraph: event-leakage defense added")
        else:
            print("[7b] Problem-Statement split already addresses event leakage; skipped")
    except LookupError:
        print("[7b] Problem-Statement split sentinel not found; skipped")

    # 8. References — add 2 entries, alphabetically sorted
    ref_paras = [p for p in doc.paragraphs if p.style.name == "Reference Text"]
    ref_texts = [(p.text or "").strip() for p in ref_paras]

    # Gambling Insider inserted before Goodfellow
    if not any("Gambling Insider" in t for t in ref_texts):
        for p in ref_paras:
            if p.text.strip().startswith("Goodfellow"):
                insert_reference_before(p, NEW_REFERENCES["Gambling Insider"])
                print("[8a] Reference added: Gambling Insider (2026, March 11)")
                break
    else:
        print("[8a] Gambling Insider reference already present; skipped")

    # Polymarket (March 10) inserted before Polymarket (March 23)
    if not any("March 10" in t for t in ref_texts):
        for p in ref_paras:
            if p.text.strip().startswith("Polymarket. (2026, March 23)"):
                insert_reference_before(p, NEW_REFERENCES["Polymarket March 10"])
                print("[8b] Reference added: Polymarket (2026, March 10)")
                break
    else:
        print("[8b] Polymarket March 10 reference already present; skipped")

    doc.save(str(DOCX))
    print(f"\nsaved {DOCX}")


if __name__ == "__main__":
    main()
