"""
26_docx_restore_cover_and_toc.py

The user asked to literally replace the cover page + table-of-contents in
ML_final_exam_paper.docx with the versions from the pre_formatting archive,
preserving the four extra fact rows that were added to the cover table
after the template was imported (Type of paper, Supervisor, Characters,
Pages).

Steps:

1. Load XML bytes from
   archive/ML_final_exam_paper.pre_formatting.docx (document.xml).
2. Inject the four fact paragraphs (extracted verbatim from the current
   document) into the pre_formatting cover SDT, immediately after the
   "Semester" paragraph.
3. Extract from pre_formatting the region [cover SDT .. first <w:p> of
   Abstract heading) — this is the cover SDT plus the TOC region.
4. Replace the same region in the current document with that block.
5. Re-attach the section-1 sectPr (top margin 1304 twips = 2.3 cm, matching
   the pre_formatting template) to the blank paragraph immediately before
   the Abstract heading so the front matter keeps its original tight
   margin while the body retains the CBS-compliant 3 cm top margin.

Idempotent. Pre-patch backup at archive/ML_final_exam_paper.pre_swap.docx.
"""
from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "ML_final_exam_paper.docx"
PRE_DOCX = ROOT / "archive" / "ML_final_exam_paper.pre_formatting.docx"
BACKUP = ROOT / "archive" / "ML_final_exam_paper.pre_swap.docx"

FRONT_MATTER_TOP_MARGIN = "1304"  # twips (~2.30 cm), pre_formatting template top
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _read_document_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as z:
        return z.read("word/document.xml").decode("utf-8")


def _write_document_xml(docx_path: Path, new_xml: str) -> None:
    """Replace word/document.xml inside the docx zip."""
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, \
         zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                data = new_xml.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(docx_path)


def _find_cover_sdt_bounds(xml: str) -> tuple[int, int]:
    marker = xml.find('<w:docPartGallery w:val="Cover Pages"/>')
    assert marker > 0, "Cover Pages gallery not found"
    start = xml.rfind("<w:sdt>", 0, marker)
    if start < 0:
        start = xml.rfind("<w:sdt ", 0, marker)
    depth = 0
    i = start
    while i < len(xml):
        if xml.startswith("<w:sdt>", i) or xml.startswith("<w:sdt ", i):
            depth += 1
            i += 6
        elif xml.startswith("</w:sdt>", i):
            depth -= 1
            i += 8
            if depth == 0:
                return start, i
        else:
            i += 1
    raise ValueError("Cover SDT not closed")


def _find_abstract_p_start(xml: str) -> int:
    abs_idx = xml.find(">Abstract<")
    assert abs_idx > 0, "Abstract heading not found"
    return xml.rfind("<w:p ", 0, abs_idx)


def _extract_fact_paragraphs(current_sdt_xml: str) -> str:
    """Find and return the XML for the four fact paragraphs (Type of paper,
    Supervisor, Characters, Pages) in the current cover SDT."""
    # The paragraphs are consecutive <w:p> blocks. Anchor the search on the
    # Type-of-paper paragraph and capture the next four <w:p> blocks.
    idx = current_sdt_xml.find("Type of paper")
    assert idx > 0, "Type of paper not found in current cover SDT"
    p_start = current_sdt_xml.rfind("<w:p ", 0, idx)
    # Consume four paragraphs
    cursor = p_start
    for _ in range(4):
        close = current_sdt_xml.find("</w:p>", cursor) + len("</w:p>")
        cursor = close
    return current_sdt_xml[p_start:cursor]


def _inject_facts_into_pre_sdt(pre_sdt_xml: str, facts_xml: str) -> str:
    """Insert facts_xml into the pre cover SDT immediately after the
    'Semester' paragraph."""
    # Locate the 'Semester' text inside a <w:t> within the SDT
    sem_idx = pre_sdt_xml.find("Semester")
    assert sem_idx > 0, "Semester not found in pre cover SDT"
    # Advance to the end of the paragraph containing Semester
    p_end = pre_sdt_xml.find("</w:p>", sem_idx) + len("</w:p>")
    return pre_sdt_xml[:p_end] + facts_xml + pre_sdt_xml[p_end:]


def swap_front_matter() -> None:
    cur_xml = _read_document_xml(DOCX)
    pre_xml = _read_document_xml(PRE_DOCX)

    cur_cs, cur_ce = _find_cover_sdt_bounds(cur_xml)
    pre_cs, pre_ce = _find_cover_sdt_bounds(pre_xml)
    cur_abs = _find_abstract_p_start(cur_xml)
    pre_abs = _find_abstract_p_start(pre_xml)

    print(f"[bounds] cur cover SDT: [{cur_cs}, {cur_ce}] cur Abstract p: {cur_abs}")
    print(f"[bounds] pre cover SDT: [{pre_cs}, {pre_ce}] pre Abstract p: {pre_abs}")

    cur_sdt = cur_xml[cur_cs:cur_ce]
    pre_sdt = pre_xml[pre_cs:pre_ce]
    pre_toc_region = pre_xml[pre_ce:pre_abs]

    # Extract the four fact paragraphs from the current cover SDT and inject
    # them into the pre cover SDT after Semester.
    facts = _extract_fact_paragraphs(cur_sdt)
    print(f"[facts] extracted {len(facts)} bytes of Type/Supervisor/Characters/Pages")
    merged_sdt = _inject_facts_into_pre_sdt(pre_sdt, facts)

    # Replace current [2629, cur_abs) with merged_sdt + pre_toc_region.
    new_xml = cur_xml[:cur_cs] + merged_sdt + pre_toc_region + cur_xml[cur_abs:]
    print(f"[swap] old front-matter length: {cur_abs - cur_cs}; "
          f"new length: {len(merged_sdt) + len(pre_toc_region)}")

    _write_document_xml(DOCX, new_xml)
    print(f"[saved] {DOCX.name}")


def reapply_section_break() -> None:
    """Place a section-1 sectPr (top margin 1304 twips) on the blank
    paragraph immediately before the Abstract heading."""
    doc = Document(DOCX)
    W = f"{{{W_NS}}}"
    body = doc.element.body

    last_sectPr = None
    for child in reversed(list(body)):
        if child.tag == f"{W}sectPr":
            last_sectPr = child
            break
    assert last_sectPr is not None, "document missing final sectPr"

    new_sect = OxmlElement("w:sectPr")
    for k, v in last_sectPr.attrib.items():
        new_sect.set(k, v)
    for child in last_sectPr:
        new_sect.append(child.__copy__())
    # Override top margin and set next-page break
    pgMar = new_sect.find(f"{W}pgMar")
    pgMar.set(f"{W}top", FRONT_MATTER_TOP_MARGIN)
    pgNumType = new_sect.find(f"{W}pgNumType")
    if pgNumType is not None:
        new_sect.remove(pgNumType)
    typ = new_sect.find(f"{W}type")
    if typ is None:
        typ = OxmlElement("w:type")
        new_sect.insert(0, typ)
    typ.set(f"{W}val", "nextPage")

    paras = doc.paragraphs
    abstract_idx = None
    for i, p in enumerate(paras):
        if p.style.name == "Heading 1" and p.text.strip() == "Abstract":
            abstract_idx = i
            break
    assert abstract_idx is not None and abstract_idx > 0

    target = paras[abstract_idx - 1]
    pPr = target._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        target._element.insert(0, pPr)
    for existing in pPr.findall(qn("w:sectPr")):
        pPr.remove(existing)
    pPr.append(new_sect)

    # Also zero space-before / space-after on every paragraph (idempotent).
    n = 0
    for p in doc.paragraphs:
        pp = p._element.find(qn("w:pPr"))
        if pp is None:
            pp = OxmlElement("w:pPr")
            p._element.insert(0, pp)
        spacing = pp.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pp.append(spacing)
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        n += 1

    # Remove manual page break in front matter (from the restored TOC field).
    removed = 0
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() == "Abstract":
            break
        for r in p._element.findall(f"{W}r"):
            for br in r.findall(f"{W}br"):
                if br.get(f"{W}type") == "page":
                    r.remove(br)
                    removed += 1

    doc.save(DOCX)
    print(f"[section] reapplied section-1 sectPr on paragraph "
          f"#{abstract_idx - 1} ({target.style.name!r})")
    print(f"[spacing] zero before/after on {n} paragraphs")
    print(f"[page-break] removed {removed} manual page-break(s) in front matter")


def main() -> None:
    assert DOCX.exists(), f"missing: {DOCX}"
    assert PRE_DOCX.exists(), f"missing: {PRE_DOCX}"

    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
        print(f"[backup] wrote {BACKUP.name}")
    else:
        print(f"[backup] {BACKUP.name} already present; not overwriting")

    swap_front_matter()
    reapply_section_break()


if __name__ == "__main__":
    main()
