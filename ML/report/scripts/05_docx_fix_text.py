"""Align ML_final_exam_paper.docx with the actual project state (work up to
and including EDA), populate the Contribution and LLM Usage Disclosure table,
add a GenAI declaration at the top of Methodology, and build the References
list in APA 7.

Idempotent-ish: re-running replaces the targeted paragraph text but will
duplicate the GenAI paragraph, LLM table, and reference entries if run twice.
One-shot.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "ML_final_exam_paper.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def set_para_text(para, text: str) -> None:
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    para.add_run(text)


def insert_paragraph_before(target_para, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    target_para._p.addprevious(new_p)
    para = target_para.__class__(new_p, target_para._parent)
    if style:
        para.style = target_para.part.document.styles[style]
    if text:
        para.add_run(text)
    return para


def insert_reference_before(target_para, parts: list[tuple[str, bool]]):
    """parts: list of (text, italic). Inserts a Reference Text paragraph."""
    try:
        new_p = insert_paragraph_before(target_para, style="Reference Text")
    except KeyError:
        new_p = insert_paragraph_before(target_para, style="Normal")
    for text, italic in parts:
        run = new_p.add_run(text)
        if italic:
            run.italic = True
    return new_p


def insert_table_before(target_para, rows: int, cols: int):
    """Create a Word table and move it to just before `target_para`."""
    doc = target_para.part.document  # type: ignore
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = "Table Grid"
    # The add_table call appended the tbl element (and a trailing empty
    # paragraph) at the end of the body. Move the tbl element to before target.
    target_para._p.addprevious(tbl._tbl)
    return tbl


# ---------------------------------------------------------------------------
# Replacement text
# ---------------------------------------------------------------------------
NEW_PARA_80 = (
    "Features are organised in six layers, all strictly no-lookahead: "
    "trade-local (log size, side, outcome index), market context (cumulative "
    "trade count and volume, one-hour rolling price volatility, and "
    "log-transformed trailing one-hour and twenty-four-hour USD volume), time "
    "(time-to-settlement with log variant and percent of market lifetime "
    "elapsed), wallet global (prior trades, prior volume, prior win rate, and "
    "wallet-age-at-trade proxy), wallet-in-market in four clusters "
    "(bet-slicing bursts, directional purity, position-aware stance, and "
    "per-market depth), and interactions (trade size relative to wallet "
    "average, the log-size by log-time-to-settlement product, and trade size "
    "as a fraction of prior market cumulative volume). Every feature at row t "
    "uses only rows with timestamp strictly before t. The contemporaneous "
    "market-implied probability is deliberately excluded from the feature set "
    "so that p_hat is independent of the market's own belief and the residual "
    "p_hat minus market_implied_prob is a clean signal; the market-implied "
    "probability is retained in the dataset only as the trading-rule "
    "benchmark."
)

NEW_PARA_82 = (
    "The split is trade-timestamp temporal. Each trade is assigned to train, "
    "validation, or test based on its own execution timestamp, with quantile "
    "boundaries at 0.70 and 0.85 of the full trade-timestamp range. A "
    "trade-timestamp split was preferred over a settlement-date split because "
    "events 114242 and 355299 cluster their NO and YES resolutions on "
    "different calendar dates; a settlement-date split would place early "
    "NO-resolving markets in training and later YES-resolving markets in "
    "test, leaving the model with no in-sample YES outcomes. Overall "
    "bet_correct rate is 0.504, well inside the 35 to 65 percent band."
)

NEW_PARA_110 = (
    "The feature set is entirely numeric after the transformations above. "
    "Market identifier is retained as metadata only and is not one-hot "
    "encoded, since the seventy-four resolved sub-markets would introduce a "
    "sparse high-dimensional encoding with a high risk of overfitting to "
    "market-level idiosyncrasies; the trade-timestamp split already mixes "
    "all markets across folds. Calendar-derived features are limited to the "
    "continuous time-to-settlement family described in Section 5.1."
)

NEW_PARA_113 = (
    "Raw trades are pulled from two complementary sources. The HuggingFace "
    "mirror SII-WANGZJ/Polymarket_data carries the complete on-chain trade "
    "history for the sixty-seven sub-markets whose resolution predates the "
    "mirror's snapshot cutoff of 2026-03-31. The mirror's trades.parquet, "
    "thirty-nine gigabytes compressed, is too large to store locally, so it "
    "is read remotely in row-group chunks via pyarrow and fsspec, filtered "
    "on the fly to the target condition identifiers, and materialised to a "
    "local parquet subset. Each row-group read is wrapped in a "
    "retry-and-reopen loop so that a transient network or decompression "
    "fault does not abort the full transfer. For the seven ceasefire "
    "sub-markets created after the HF cutoff, trade history is pulled from "
    "the Polymarket Data API with side-split pagination; all seven markets "
    "contain fewer than five thousand trades each, well below the API's "
    "offset ceiling, so no truncation occurs. Market metadata and resolution "
    "status come from the Polymarket Gamma API for all seventy-four markets. "
    "Per-trade enrichment derives a true resolution timestamp from the first "
    "lock of the winning-token price, the contemporaneous market-implied "
    "probability by as-of merge against CLOB mid-prices where available and "
    "the trade-execution price otherwise, and the running market and wallet "
    "features described in Section 5.1. All running features use only rows "
    "with timestamp strictly before the target row. On-chain wallet-identity "
    "enrichment via Polygonscan and news-timing enrichment via GDELT are "
    "scoped as future work and are not used in the current dataset."
)

GENAI_DECLARATION = (
    "Implementation code and sections of the written report were co-authored "
    "with Claude (Anthropic, Opus 4.7) under continuous author review. The "
    "Contribution and LLM Usage Disclosure section documents the scope and "
    "validation of that assistance."
)

LLM_TABLE_INTRO = (
    "The group used Claude (Anthropic, Opus 4.7) for code drafting and "
    "report prose under continuous author review. All outputs were reviewed, "
    "tested, and corrected by the group; no component was accepted without "
    "author validation. The table below summarises the scope and validation "
    "of LLM assistance across the main project components. Research question, "
    "scope, feature taxonomy, and trading-rule design were authored by the "
    "group without LLM assistance."
)

LLM_TABLE_HEADER = [
    "ID",
    "Component / Module",
    "Contributor(s)",
    "LLM Tool (Model)",
    "Type of LLM Usage",
    "% LLM-Generated Code",
    "Validation / Modification",
    "Remarks",
]

LLM_TABLE_ROWS = [
    [
        "1",
        "Research question, scope, feature taxonomy, trading-rule design",
        "Authors",
        "—",
        "None",
        "0%",
        "—",
        "Conceptual design authored by the group.",
    ],
    [
        "2",
        "Data extraction pipeline (01_polymarket_api.py, 02_build_dataset.py)",
        "Authors + LLM",
        "Claude (Anthropic, Opus 4.7)",
        "Code drafting and refactoring",
        "70%",
        "Authors reviewed code against the plan, ran sample tests, and identified two silent data bugs during EDA review.",
        "Bugs caught and patched: HF token-label mismatch and uint64-seconds-as-nanoseconds timestamp.",
    ],
    [
        "3",
        "Exploratory data analysis script (04_eda.py)",
        "Authors + LLM",
        "Claude (Anthropic, Opus 4.7)",
        "Code drafting to the Design.md specification",
        "75%",
        "Authors verified column references against the dataset and checked figure conventions against Design.md.",
        "Nine figures and five tables produced for the report Appendix.",
    ],
    [
        "4",
        "Report prose (Methodology narrative, EDA commentary, Appendix captions)",
        "Authors + LLM",
        "Claude (Anthropic, Opus 4.7)",
        "Prose drafting and structural editing",
        "50%",
        "Authors verified every factual and numerical claim against project artefacts and edited prose for CBS academic style.",
        "—",
    ],
]


# ---------------------------------------------------------------------------
# APA 7 references — alphabetical
# ---------------------------------------------------------------------------
REFERENCES = [
    # Coplan, S.
    [
        ("Coplan, S. (2025, November). Interview with Cecilia Vega. ", False),
        ("60 Minutes", True),
        (". CBS News.", False),
    ],
    # Goodfellow et al. (book title italic)
    [
        ("Goodfellow, I., Bengio, Y., & Courville, A. (2016). ", False),
        ("Deep learning", True),
        (". MIT Press.", False),
    ],
    # Mitts & Ofir
    [
        ("Mitts, J., & Ofir, A. (2026). ", False),
        ("Informed trading in prediction markets", True),
        (" [Working paper]. Columbia Law School.", False),
    ],
    # Polymarket corporate announcement
    [
        ("Polymarket. (2026, March 23). ", False),
        ("Updated trading rules and enforcement policy", True),
        (" [Company announcement]. Polymarket.", False),
    ],
    # Wolfers & Zitzewitz
    [
        ("Wolfers, J., & Zitzewitz, E. (2004). Prediction markets. ", False),
        ("Journal of Economic Perspectives, 18", True),
        ("(2), 107–126. https://doi.org/10.1257/0895330041371321", False),
    ],
]


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    doc = Document(str(DOCX))

    # Capture paragraph references by text so index shifts don't matter.
    def find(text: str, style_prefix: str | None = None):
        for p in doc.paragraphs:
            if p.text.strip() == text and (style_prefix is None or p.style.name.startswith(style_prefix)):
                return p
        raise LookupError(f"paragraph with text {text!r} not found")

    def find_containing(substring: str):
        for p in doc.paragraphs:
            if substring in (p.text or ""):
                return p
        raise LookupError(f"paragraph containing {substring!r} not found")

    abstract_para = find_containing("Prediction markets on Polymarket price real-world events")
    llm_placeholder = find_containing("Table summarising per-member contributions")
    methodology_heading = find("Methodology", style_prefix="Heading 1")
    dataset_desc_heading = find("Dataset Description", style_prefix="Heading 2")
    features_para = find_containing("Features are organised in six layers")
    split_para = find_containing("temporal by market settlement date")
    winsor_para = find_containing("wallet_total_volume_usd are winsorised")
    encoding_para = find_containing("one-hot encoded")
    data_filter_para = find_containing(
        "Raw trades are pulled from two complementary sources"
    )
    ref_heading = find("References")

    # --- 1. Abstract ---
    new_abs = abstract_para.text
    new_abs = new_abs.replace(
        "Logistic regression and random forest serve as baselines",
        "Logistic regression, random forest, and Isolation Forest serve as baselines",
    )
    new_abs = new_abs.replace(
        "temporally held-out markets along two axes",
        "temporally held-out data along two axes",
    )
    set_para_text(abstract_para, new_abs)

    # --- 2-6. Text replacements in Methodology ---
    set_para_text(features_para, NEW_PARA_80)
    set_para_text(split_para, NEW_PARA_82)
    set_para_text(
        winsor_para,
        winsor_para.text.replace(
            "wallet_total_volume_usd", "wallet_prior_volume_usd"
        ),
    )
    set_para_text(encoding_para, NEW_PARA_110)
    set_para_text(data_filter_para, NEW_PARA_113)

    # --- 7. GenAI declaration at top of Methodology ---
    insert_paragraph_before(dataset_desc_heading, GENAI_DECLARATION, style="Normal")
    insert_paragraph_before(dataset_desc_heading, "", style="Normal")

    # --- 8. Contribution and LLM Usage Disclosure placeholder + table ---
    set_para_text(llm_placeholder, LLM_TABLE_INTRO)
    # Insert a blank, then the table, before the Methodology heading.
    insert_paragraph_before(methodology_heading, "", style="Normal")
    tbl = insert_table_before(methodology_heading, rows=1 + len(LLM_TABLE_ROWS),
                              cols=len(LLM_TABLE_HEADER))
    for j, h in enumerate(LLM_TABLE_HEADER):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for i, row in enumerate(LLM_TABLE_ROWS):
        for j, v in enumerate(row):
            tbl.rows[i + 1].cells[j].text = v
    insert_paragraph_before(methodology_heading, "", style="Normal")

    # --- 9. References list (APA 7, alphabetical) ---
    # Insert each entry before the next major heading that follows References
    # (Appendix). Find Appendix heading first; if missing, insert before end.
    appendix_heading = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() == "Appendix":
            appendix_heading = p
            break
    insertion_target = appendix_heading if appendix_heading is not None else None
    if insertion_target is None:
        raise LookupError("Appendix heading not found; cannot place references.")
    for parts in REFERENCES:
        insert_reference_before(insertion_target, parts)
    insert_paragraph_before(insertion_target, "", style="Normal")

    doc.save(str(DOCX))
    print(f"Saved {DOCX}")
    print(f"  - 1 Abstract edit")
    print(f"  - 5 Methodology text replacements")
    print(f"  - 1 GenAI declaration inserted")
    print(f"  - 1 LLM Usage Disclosure intro + {len(LLM_TABLE_ROWS)}-row table")
    print(f"  - {len(REFERENCES)} APA 7 reference entries")


if __name__ == "__main__":
    main()
