"""Insert EDA body narrative and an Appendix with all figures and tables into
`ML_final_exam_paper.docx`.

Body paragraphs are inserted immediately after the existing Heading 2
"Exploratory Data Analysis" (before the "Data Pre-Processing" heading). Each
paragraph comments on the associated artefact and points to its appendix entry.
All figures and their captioned tables live in a new Heading 1 "Appendix"
appended at the end of the document.

One-shot script — running it twice will duplicate content.
"""

from __future__ import annotations

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "ML_final_exam_paper.docx"
EDA = ROOT / "outputs" / "eda"

FIG_W_IN = 6.3  # Design.md FIG_W for A4 text-width


# ---------------------------------------------------------------------------
# Insertion helpers (python-docx has no "insert before" on the doc object;
# operate on the underlying lxml)
# ---------------------------------------------------------------------------
def insert_paragraph_before(target_p, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    target_p._p.addprevious(new_p)
    para = target_p.__class__(new_p, target_p._parent)
    if style:
        para.style = target_p.part.document.styles[style]
    if text:
        para.add_run(text)
    return para


def insert_picture_before(target_p, image_path: Path, width_in: float):
    # Create an empty paragraph first, then add the picture run.
    para = insert_paragraph_before(target_p, "", style=None)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(image_path), width=Inches(width_in))
    return para


# ---------------------------------------------------------------------------
# Body content — 9 paragraphs, each pointing to its appendix figure/table
# ---------------------------------------------------------------------------
EDA_BODY_PARAGRAPHS = [
    "The dataset spans 1,209,787 resolved trades across 74 Iran-event sub-markets, "
    "placed by 109,080 distinct wallets between 22 December 2025 and 19 April 2026. "
    "Overall bet_correct rate is 0.504, within the 35 to 65 percent band and "
    "requiring no resampling at this stage. The trade-timestamp split produces "
    "846,847 training, 181,454 validation, and 181,486 test rows, with per-split "
    "correctness rates of 0.495, 0.523, and 0.522 respectively. Per-split and "
    "per-market correctness distributions are reported in Appendix Figure A.2 and "
    "Appendix Table A.1.",

    "Missingness concentrates in Gamma metadata fields such as title, icon, and "
    "slug variants, which are informational rather than predictive, and in "
    "features that are structurally undefined on a wallet's first trade within a "
    "market. The in-market directional features wallet_directional_purity_in_market "
    "and wallet_spread_ratio are null for 22.6 percent of rows, corresponding to "
    "first-trade events. Null rates for all other modelling features remain below "
    "10 percent. The distribution of missingness by column is shown in Appendix "
    "Figure A.1.",

    "Univariate feature distributions split by bet_correct reveal limited "
    "separation between correct and incorrect trades. The clearest visible shift "
    "appears on time-to-settlement and percent-of-lifetime-elapsed, consistent "
    "with the hypothesis that informed flow concentrates near market deadlines. "
    "Count and USD-denominated features display long right tails: skewness "
    "exceeds 500 for size_vs_market_cumvol_pct and size_vs_wallet_avg, and "
    "exceeds 80 for trade_value_usd. Log1p transforms applied at feature "
    "construction compress these distributions to usable ranges. Feature "
    "distributions and skewness rankings are reported in Appendix Figure A.3 and "
    "Appendix Table A.2.",

    "Box plots of the eight most-skewed numeric features, percentile-clipped at "
    "the 1st and 99th for display readability, confirm that concentration in the "
    "upper tail is the dominant outlier pattern. Winsorisation rather than "
    "row-level removal is applied during pre-processing because extreme trades, "
    "in particular abnormally large bets relative to market activity, are part "
    "of the hypothesised informed-flow signal. See Appendix Figure A.4.",

    "Pairwise Pearson correlation is block-diagonal along the six feature "
    "layers: market-context variables correlate within themselves, wallet-global "
    "variables within themselves, and so on. The strongest pair is "
    "market_trade_count_so_far with market_volume_so_far_usd at 0.918, followed "
    "by wallet_directional_purity_in_market with wallet_spread_ratio at 0.917. "
    "No cross-layer pair exceeds 0.82. This level of redundancy is tolerable "
    "for the MLP under L2 regularisation and dropout, and motivates a Lasso "
    "pass in the baseline logistic regression to prune clearly duplicate "
    "signals. Correlation structure and the top-ranked pairs are reported in "
    "Appendix Figure A.5 and Appendix Table A.3.",

    "A two-component principal component projection of 38,570 wallets with at "
    "least five trades, computed over six behavioural aggregates, does not "
    "reveal a cleanly separated high-correctness cluster. Colouring by mean "
    "wallet correctness shows overlapping distributions across the projection. "
    "Any signal therefore resides in finer-grained interactions rather than in "
    "a coarse behavioural taxonomy, which motivates the non-linear MLP over "
    "simpler discrete-rule alternatives. See Appendix Figure A.6.",

    "Per-market price paths show the expected mixture of clean resolutions, "
    "with price approaching zero well before deadline for markets resolving NO, "
    "and late-stage information-driven swings for markets where the underlying "
    "event status was genuinely uncertain until close to settlement. The latter "
    "pattern is the environment in which mispricing is most plausible and in "
    "which the home-run trading rule is expected to operate. See Appendix "
    "Figure A.7.",

    "Aggregating trading activity and correctness by time-to-settlement bucket "
    "provides direct empirical support for the home-run gating rule. Trades "
    "placed less than one hour before settlement, representing 10.1 million USD "
    "in volume across 31,926 trades, show a mean bet_correct of 0.544. This is "
    "5.7 percentage points above the one-to-seven-day bucket at 0.487, which "
    "carries the bulk of total volume. The late-concentrated correctness "
    "pattern is consistent with documented informed-flow dynamics on Iran "
    "prediction markets and justifies the home-run rule's gate at "
    "time-to-settlement below six hours. Volume and correctness by bucket are "
    "reported in Appendix Figure A.8 and Appendix Table A.4.",

    "Partitioning wallets by directional purity and intraday burst rate yields "
    "four behavioural quadrants. Mean correctness within each quadrant sits "
    "between 0.470 and 0.501, close to the overall 0.504 baseline. The absence "
    "of a clean quadrant effect reinforces the conclusion from the PCA "
    "projection: a hand-specified discrete taxonomy does not produce an edge "
    "on its own, and the value-add of the MLP lies in learning non-linear "
    "interactions across the continuous feature set. Quadrant composition and "
    "correctness rates are reported in Appendix Figure A.9 and Appendix Table "
    "A.5.",
]


# ---------------------------------------------------------------------------
# Appendix content: (caption_label, caption_text, image_path_or_None, table_spec_or_None)
# Table spec = ("Table A.N.", caption, rows-of-lists, column-labels)
# ---------------------------------------------------------------------------
def build_appendix_items():
    items: list[dict] = []

    # -------- Figure A.1 missingness --------
    items.append({
        "type": "figure",
        "label": "Figure A.1.",
        "caption": "Share of missing values by column. Only columns with at least "
                   "one null are shown. Rates above 98 percent correspond to Gamma "
                   "metadata fields present only on ceasefire-market rows.",
        "image": EDA / "01_missingness.png",
    })

    # -------- Table A.1 summary stats --------
    summary_rows = [
        ["Trades", "1,209,787"],
        ["Unique wallets (proxyWallet)", "109,080"],
        ["Resolved markets (condition_id)", "74"],
        ["Timespan", "2025-12-22 to 2026-04-19"],
        ["Mean bet_correct", "0.504"],
        ["Train rows (mean bet_correct)", "846,847 (0.495)"],
        ["Validation rows (mean bet_correct)", "181,454 (0.523)"],
        ["Test rows (mean bet_correct)", "181,486 (0.522)"],
    ]
    items.append({
        "type": "table",
        "label": "Table A.1.",
        "caption": "Dataset summary statistics.",
        "header": ["Metric", "Value"],
        "rows": summary_rows,
    })

    # -------- Figure A.2 class balance --------
    items.append({
        "type": "figure",
        "label": "Figure A.2.",
        "caption": "Left: share of correct and incorrect trades per temporal split. "
                   "Right: mean bet_correct per market, sorted ascending, with the "
                   "0.5 baseline marked.",
        "image": EDA / "02_class_balance.png",
    })

    # -------- Figure A.3 distributions + Table A.2 skewness --------
    items.append({
        "type": "figure",
        "label": "Figure A.3.",
        "caption": "Histograms of twelve pre-processed features, separated by "
                   "bet_correct outcome, 1st to 99th percentile clipped for "
                   "display readability. Warmer tones represent correct trades.",
        "image": EDA / "03_feature_distributions.png",
    })

    skew_df = pd.read_csv(EDA / "03_skewness_table.csv").head(15)
    skew_df.columns = ["Feature", "Skewness"]
    skew_rows = [[r["Feature"], f"{r['Skewness']:.2f}"]
                 for _, r in skew_df.iterrows()]
    items.append({
        "type": "table",
        "label": "Table A.2.",
        "caption": "Top fifteen numeric features ranked by absolute skewness.",
        "header": ["Feature", "Skewness"],
        "rows": skew_rows,
    })

    # -------- Figure A.4 outliers --------
    items.append({
        "type": "figure",
        "label": "Figure A.4.",
        "caption": "Box plots of the eight most-skewed numeric features, "
                   "percentile-clipped at the 1st and 99th for display readability.",
        "image": EDA / "04_outlier_boxplots.png",
    })

    # -------- Figure A.5 correlation + Table A.3 top pairs --------
    items.append({
        "type": "figure",
        "label": "Figure A.5.",
        "caption": "Pairwise Pearson correlation among the numeric features, "
                   "estimated on a 100,000-row random sample.",
        "image": EDA / "05_correlation_heatmap.png",
    })

    corr_rows: list[list[str]] = []
    with open(EDA / "05_top_correlations.txt") as f:
        for line in f.read().splitlines()[2:17]:  # skip heading lines, take top 15
            parts = line.rsplit(None, 1)
            if len(parts) != 2:
                continue
            pair, r = parts
            pair_tokens = pair.split()
            if len(pair_tokens) < 2:
                continue
            split_at = len(pair_tokens) // 2
            f1 = " ".join(pair_tokens[:split_at])
            f2 = " ".join(pair_tokens[split_at:])
            corr_rows.append([f1, f2, r])
    items.append({
        "type": "table",
        "label": "Table A.3.",
        "caption": "Top fifteen feature pairs by absolute Pearson correlation.",
        "header": ["Feature 1", "Feature 2", "|r|"],
        "rows": corr_rows,
    })

    # -------- Figure A.6 PCA --------
    items.append({
        "type": "figure",
        "label": "Figure A.6.",
        "caption": "Two-component principal component projection of 38,570 "
                   "wallets with five or more trades, over six behavioural "
                   "aggregates. Points are coloured by wallet mean bet_correct.",
        "image": EDA / "06_pca_wallets.png",
    })

    # -------- Figure A.7 price trajectories --------
    items.append({
        "type": "figure",
        "label": "Figure A.7.",
        "caption": "Per-market trade price trajectories across the 74 resolved "
                   "sub-markets.",
        "image": EDA / "07_price_trajectories.png",
    })

    # -------- Figure A.8 event timing + Table A.4 --------
    items.append({
        "type": "figure",
        "label": "Figure A.8.",
        "caption": "Left: total traded USD volume by time-to-settlement bucket. "
                   "Right: mean bet_correct per bucket, with the 0.5 baseline marked. "
                   "Both computed with the post-resolution filter "
                   "settlement_minus_trade_sec > 0.",
        "image": EDA / "08_event_timing.png",
    })

    et_rows = [
        ["<1h",   "31,926",  "10,122,690",  "0.5442"],
        ["1-6h",  "75,645",  "22,305,230",  "0.5088"],
        ["6-24h", "209,170", "45,241,500",  "0.5292"],
        ["1-7d",  "421,442", "73,156,270",  "0.4872"],
        ["7-30d", "284,107", "40,551,150",  "0.4967"],
        [">30d",  "62,096",  "5,528,511",   "0.5304"],
    ]
    items.append({
        "type": "table",
        "label": "Table A.4.",
        "caption": "Trade count, total USD volume, and mean bet_correct by "
                   "time-to-settlement bucket.",
        "header": ["Bucket", "Trades", "Volume (USD)", "Mean bet_correct"],
        "rows": et_rows,
    })

    # -------- Figure A.9 quadrants + Table A.5 --------
    items.append({
        "type": "figure",
        "label": "Figure A.9.",
        "caption": "Wallet behavioural quadrants. Top row: histograms of mean "
                   "directional purity and burst rate across wallets with five or "
                   "more trades. Bottom left: log-transformed wallet trade count. "
                   "Bottom right: mean bet_correct per quadrant.",
        "image": EDA / "09_wallet_quadrants.png",
    })

    q_rows = [
        ["pure + bursty",     "1,370",  "0.501"],
        ["pure, not bursty",  "27,480", "0.484"],
        ["mixed + bursty",    "1,246",  "0.470"],
        ["mixed, not bursty", "8,474",  "0.497"],
    ]
    items.append({
        "type": "table",
        "label": "Table A.5.",
        "caption": "Wallet behavioural quadrants and mean bet_correct. Purity "
                   "threshold 0.7; burst-rate threshold 0.2. Filter: wallets with "
                   "five or more trades.",
        "header": ["Quadrant", "Wallets", "Mean bet_correct"],
        "rows": q_rows,
    })

    return items


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    doc = Document(str(DOCX))

    # 1. Locate the "Exploratory Data Analysis" heading and the following
    #    "Data Pre-Processing" heading; insert body paragraphs before the latter.
    eda_idx = None
    dp_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 2":
            if p.text.strip() == "Exploratory Data Analysis":
                eda_idx = i
            elif (p.text.strip() == "Data Pre-Processing"
                  and eda_idx is not None and dp_idx is None):
                dp_idx = i
                break
    if eda_idx is None or dp_idx is None:
        raise SystemExit("Could not locate EDA / Data Pre-Processing headings.")

    # Guard against double-run.
    for p in doc.paragraphs[eda_idx + 1: dp_idx]:
        if p.text.strip().startswith("The dataset spans 1,209,787"):
            raise SystemExit("EDA body text already inserted; aborting.")

    target = doc.paragraphs[dp_idx]
    for body in EDA_BODY_PARAGRAPHS:
        insert_paragraph_before(target, body, style="Normal")
        insert_paragraph_before(target, "", style="Normal")  # blank separator

    # 2. Append Appendix section at the end of the document.
    appendix_heading = doc.add_heading("Appendix", level=1)

    intro = doc.add_paragraph(
        "The following figures and tables accompany the Exploratory Data "
        "Analysis section. Each is cross-referenced from the main body where "
        "its findings are discussed."
    )
    doc.add_paragraph("")

    items = build_appendix_items()
    for it in items:
        if it["type"] == "figure":
            # Place the image first, caption line beneath (per Design.md style).
            fig_para = doc.add_paragraph()
            fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_para.add_run().add_picture(str(it["image"]), width=Inches(FIG_W_IN))
            cap = doc.add_paragraph()
            run = cap.add_run(f"{it['label']} ")
            run.bold = True
            cap.add_run(it["caption"]).italic = True
            doc.add_paragraph("")
        else:
            # Caption above the table.
            cap = doc.add_paragraph()
            run = cap.add_run(f"{it['label']} ")
            run.bold = True
            cap.add_run(it["caption"]).italic = True

            tbl = doc.add_table(rows=1 + len(it["rows"]), cols=len(it["header"]))
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            for j, h in enumerate(it["header"]):
                hdr[j].text = ""
                run = hdr[j].paragraphs[0].add_run(h)
                run.bold = True
            for r_idx, row in enumerate(it["rows"]):
                cells = tbl.rows[r_idx + 1].cells
                for j, v in enumerate(row):
                    cells[j].text = str(v)
            doc.add_paragraph("")

    doc.save(str(DOCX))
    print(f"saved {DOCX} with {len(EDA_BODY_PARAGRAPHS)} EDA body paragraphs "
          f"and {len(items)} appendix items")


if __name__ == "__main__":
    main()
