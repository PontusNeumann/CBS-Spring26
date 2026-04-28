"""Build a sidecar appendix docx for the wallet-joined EDA panels.

Produces `pontus/outputs/eda_walletjoined/eda_appendix.docx` (and a
matching .md preview). Each panel from `eda_walletjoined/` lands as a
captioned figure, formatted to match the report's existing Appendix A
style (A.10 onward).

This file is a STANDALONE document — it does NOT modify the submitted
docx. Once captions and figure ordering are reviewed, hand-merge into
the paper at the end of Appendix A.

Usage:
    python pontus/scripts/build_eda_appendix_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
EDA_DIR = ROOT / "pontus" / "outputs" / "eda_walletjoined"
OUT_DOCX = EDA_DIR / "eda_appendix.docx"
OUT_MD = EDA_DIR / "eda_appendix.md"

FIG_W_IN = 6.3  # Design.md FIG_W for A4 text-width

# Each entry: appendix label, image filename in EDA_DIR, caption text.
# Captions mirror the academic style used in the existing report (no
# em-dashes, no contractions, lead with what the figure shows). Numbers
# pick up where the original A.9 left off.
APPENDIX_ITEMS: list[dict] = [
    {
        "label": "Figure A.10.",
        "image": "02_wallet_coverage.png",
        "caption": (
            "Wallet enrichment coverage by split. After two PolygonScan "
            "extension passes, every taker in Alex's idea1 cohort is "
            "matched to a row in the wallet enrichment table, giving "
            "100.0 percent train and 100.0 percent test coverage on the "
            "1,371,180-trade cohort. The Layer 6 NaN rate falls to 0 "
            "percent after the retry passes."
        ),
    },
    {
        "label": "Figure A.11.",
        "image": "08_train_test_shift.png",
        "caption": (
            "Train-to-test mean shift for the top 15 numeric features, "
            "expressed as the absolute standardised mean difference "
            "(Cohen d convention: 0.2 small, 0.5 medium, 0.8 large). "
            "Time-to-deadline indicators show large shift "
            "(d above 0.9), reflecting the regime change between the "
            "Iran-strike and Iran-ceasefire countdowns. Wallet-history "
            "features show medium shift, motivating the per-market "
            "GroupKFold protocol."
        ),
    },
    {
        "label": "Figure A.12.",
        "image": "09_late_flow.png",
        "caption": (
            "Hit rate and trade-count share by time-to-deadline bucket "
            "for train (Iran strike) and test (Iran ceasefire) cohorts. "
            "Test base rate falls below 0.40 inside the final hour and "
            "below 0.34 in the last five minutes, replicating the "
            "late-concentrated informed-flow signature documented in "
            "Mitts and Ofir 2026. The corresponding train-cohort "
            "buckets stay near the 0.50 base rate, isolating the "
            "phenomenon to the ceasefire regime."
        ),
    },
    {
        "label": "Figure A.13.",
        "image": "10_wallet_strata.png",
        "caption": (
            "Bet-correct base rate stratified by three Layer 6 wallet "
            "features. Left: deciles of wallet age in days at the time "
            "of the trade. Centre: causal CEX-funding flag, defined as "
            "first USDC inbound from a known CEX hot wallet observed "
            "before the trade timestamp. Right: deciles of polygon "
            "nonce at trade time, a proxy for prior on-chain "
            "experience. The split structure motivates inclusion of "
            "the Layer 6 enrichment in the strict-branch feature set."
        ),
    },
    {
        "label": "Figure A.14.",
        "image": "11_per_market_bimodality.png",
        "caption": (
            "Per-market bet_correct base rate, split by train and test "
            "cohorts. The single-event resolution of each market "
            "produces a bimodal distribution: markets resolving with "
            "the consensus side cluster near 1, markets resolving "
            "against consensus cluster near 0. The shape is the "
            "structural reason single-feature ROC across markets is "
            "highly variable and motivates the GroupKFold(market_id) "
            "evaluation protocol."
        ),
    },
    {
        "label": "Figure A.15.",
        "image": "12_feature_stability.png",
        "caption": (
            "Single-feature ROC-AUC heatmap across markets for the "
            "top eight features by absolute Pearson correlation with "
            "the target. Several features (log_payoff_if_correct, "
            "contrarian_score, risk_reward_ratio_pre) achieve median "
            "single-feature AUC near 0.31 with p95 above 0.67, "
            "indicating a single transferable rule that inverts "
            "between YES- and NO-resolved markets. This is the "
            "structural finding behind the per-market bimodality "
            "and reinforces the use of group-aware cross validation."
        ),
    },
    {
        "label": "Figure A.16.",
        "image": "13_mutual_information.png",
        "caption": (
            "Top 20 features by mutual information with bet_correct, "
            "computed on a 150,000-row stratified sample using "
            "scikit-learn's mutual_info_classif. Mutual information "
            "captures non-linear dependence that the Pearson "
            "correlation in Figure A.5 cannot detect. The leading "
            "features are short-window microstructure variables "
            "(realized volatility, jump component, order-flow "
            "imbalance) rather than wallet-identity features, "
            "consistent with the literature emphasis on flow signals."
        ),
    },
    {
        "label": "Figure A.17.",
        "image": "14_feature_taxonomy.png",
        "caption": (
            "Numeric feature counts by engineering layer. Of the 79 "
            "numeric features in the wallet-joined cohort, the "
            "market-state rolling layer contributes the largest "
            "share, followed by price and volatility microstructure, "
            "the HF-internal taker aggregates, and the on-chain "
            "Layer 6 wallet features. The taxonomy frames how the "
            "report's permutation-importance and ablation analyses "
            "are organised."
        ),
    },
    {
        "label": "Figure A.18.",
        "image": "15_tail_diagnostics.png",
        "caption": (
            "Excess kurtosis ranking for the 15 most fat-tailed "
            "numeric features, with the |kurt|=3 fat-tail threshold "
            "marked. The accompanying CSV (15_tail_diagnostics.csv) "
            "reports the 1st, 5th, 95th, and 99th percentiles plus "
            "tail-conditional means for each feature, supporting the "
            "winsorisation choice during pre-processing."
        ),
    },
    {
        "label": "Figure A.19.",
        "image": "16_temporal_drift.png",
        "caption": (
            "Daily bet_correct base rate during the strike-countdown "
            "(train) and ceasefire-countdown (test) windows, with a "
            "seven-day rolling mean overlay. Daily bins are filtered "
            "to days with at least 50 trades. The rolling means stay "
            "within a tight band around 0.50, confirming there is no "
            "temporal drift in the target within either cohort that "
            "would invalidate the static train and test split."
        ),
    },
]


def _add_caption(doc: Document, label: str, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label_run = p.add_run(label + " ")
    label_run.bold = True
    cap_run = p.add_run(caption)
    cap_run.italic = True


def _add_figure(doc: Document, image_path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(FIG_W_IN))


def build_docx() -> None:
    doc = Document()
    # Default font tweak so output matches typical CBS body copy.
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_heading("Appendix A — Wallet-joined cohort EDA (sidecar)", level=1)
    doc.add_paragraph(
        "This sidecar groups the additional EDA figures produced after "
        "the wallet enrichment was extended to full coverage of Alex's "
        "1,371,180-trade idea1 cohort (100 percent of train and 100 percent "
        "of test). Numbering picks up after the existing Appendix A.9 in "
        "the main report. Each figure is independent and may be merged "
        "into the appendix at the editor's discretion."
    )
    doc.add_paragraph("")

    for item in APPENDIX_ITEMS:
        img = EDA_DIR / item["image"]
        if not img.exists():
            doc.add_paragraph(
                f"[missing image: {item['image']}]"
            )
            _add_caption(doc, item["label"], item["caption"])
            continue
        _add_figure(doc, img)
        _add_caption(doc, item["label"], item["caption"])
        doc.add_paragraph("")

    doc.save(str(OUT_DOCX))
    print(f"saved {OUT_DOCX.relative_to(ROOT)}  ({len(APPENDIX_ITEMS)} figures)")


def build_md() -> None:
    """Markdown twin so the captions are reviewable in plain text."""
    lines = ["# Appendix A — Wallet-joined cohort EDA (sidecar)", ""]
    lines.append(
        "Sidecar EDA captions for review before merging into the main "
        "report. Numbering picks up after the existing Appendix A.9."
    )
    lines.append("")
    for item in APPENDIX_ITEMS:
        img = EDA_DIR / item["image"]
        present = "✓" if img.exists() else "MISSING"
        lines.append(f"### {item['label']} ({present})")
        lines.append(f"`{item['image']}`")
        lines.append("")
        lines.append(item["caption"])
        lines.append("")
    OUT_MD.write_text("\n".join(lines) + "\n")
    print(f"saved {OUT_MD.relative_to(ROOT)}")


def main() -> None:
    build_docx()
    build_md()


if __name__ == "__main__":
    main()
