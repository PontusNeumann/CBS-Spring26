"""
patch_docx.py

In-place surgical patch to the latest root docx.

Targets only:
  1. Section 1 Introduction: split RQ paragraph into two so the unifying
     hypothesis sits with the research question, and the audience note +
     theoretical lenses + argument move into a separate paragraph.
  2. Section 2 Background: inject the Revolut/Nordnet/Avanza fintech-pressure
     sentence in the second paragraph after "high regulatory demands."
  3. Section 5.4 Implications: replace the first paragraph with the
     sharpened LO4 deliverable language.
  4. Page-break design: keep page-break-before only on References, Appendix,
     and the back-page paragraph. Strip it from body Heading 1 paragraphs.
  5. References: merge multi-paragraph entries (journal-name continuations)
     into one paragraph per reference, with exactly one blank Reference Text
     paragraph between consecutive references.

Does not touch:
  - The cover SDT / first page block.
  - Heading numbering. The CBS template auto-numbers Heading 1; do not add
    a manual "N. " prefix anywhere.
  - Any other paragraph, table, footer, header, or media element.

Idempotent on a successful run: the body-edit anchors are checked before
patching, the page-break logic is set-based, and the reference rebuild
produces a canonical layout regardless of starting state.

Backs up the source file to backup/ before saving.
"""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
TARGET = ROOT / "KAN-CDSCO2401U_185912_DPD_Spring2026.docx"
BACKUP_DIR = ROOT / "backup"


# ---------- new content ----------

NEW_RQ_PARA = (
    "The paper addresses the following research question: why are European "
    "universal banks outsourcing wealth-platform operations to FNZ, and "
    "what does FNZ's end-to-end, assets-linked platform model imply for "
    "bank differentiation and concentration risk in European wealth "
    "management? The underlying business problem is how incumbent banks "
    "can capture the efficiency benefits of FNZ-style outsourcing without "
    "losing strategic differentiation or creating unmanaged concentration "
    "risk. The unifying hypothesis is that FNZ outsourcing is "
    "strategically rational only if banks treat it as a "
    "platform-governance and capability-rebuilding problem, not merely as "
    "a cost-saving IT project."
)

NEW_AUDIENCE_PARA = (
    "The primary advisory audience is the incumbent European universal "
    "bank executive committee; the secondary audience is the European "
    "wealth-management sector and regulator. The analysis uses three "
    "theoretical lenses: transaction-cost economics, multi-sided platform "
    "economics, and dynamic capabilities. The argument is that platform "
    "outsourcing solves a real operating problem, but it also shifts "
    "competition away from operational excellence and toward advice, "
    "distribution, client ownership, and the strategic use of data."
)

FINTECH_INSERT = (
    " Tech-led retail-finance competitors such as Revolut, Nordnet, and "
    "Avanza intensify the pressure by raising customer expectations for "
    "speed, cost, and digital service quality; this makes modernization "
    "more urgent, even though FNZ itself operates as B2B wealth "
    "infrastructure rather than as a direct consumer-fintech rival."
)

NEW_5_4_P1 = (
    "For an incumbent universal bank already executing a wealth-platform "
    "outsourcing decision, three implications follow. First, the "
    "standardized operational layers are the ones to outsource, while the "
    "layers that create differentiation are the ones to retain and "
    "strengthen: client relationship, advisory quality, distribution, "
    "brand, and proprietary data capabilities. Second, governance "
    "protections should target the operational layer through four "
    "concrete commitments embedded in contract and operating procedure: "
    "data portability, exit planning, roadmap influence, and resilience "
    "oversight. Third, the bank's internal rebuild should be funded as a "
    "strategic capability programme tied to advice, distribution, and the "
    "data layer, and treated as the principal source of differentiation "
    "once the operational backbone is shared."
)


# ---------- helpers ----------

def get_or_add_pPr(p_el):
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    return pPr


def make_paragraph(style_id: str, text: str = ""):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)
    p.append(pPr)
    if text:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
    return p


def replace_text(p, new_text):
    """Replace all runs in a paragraph with a single plain-text run."""
    p_el = p._element
    for r in p_el.findall(qn("w:r")):
        p_el.remove(r)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = new_text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p_el.append(r)


def remove_page_break_before(p):
    pPr = p._element.find(qn("w:pPr"))
    if pPr is not None:
        for el in pPr.findall(qn("w:pageBreakBefore")):
            pPr.remove(el)


def add_page_break_before(p):
    pPr = get_or_add_pPr(p._element)
    if pPr.find(qn("w:pageBreakBefore")) is None:
        pPr.append(OxmlElement("w:pageBreakBefore"))


def find_para(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p):
            return p
    return None


# ---------- main ----------

REF_START = re.compile(r"\((\d{4}[a-z]?|n\.d\.)\)")


def patch():
    if not TARGET.exists():
        raise SystemExit(f"missing target: {TARGET}")

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    backup = BACKUP_DIR / f"{TARGET.stem}_{stamp}_pre-patch.docx"
    shutil.copy2(TARGET, backup)
    print(f"[backup] {backup.name}")

    doc = Document(str(TARGET))

    # ---- 1. §1 Introduction: split RQ paragraph + insert audience paragraph ----
    rq_p = find_para(
        doc,
        lambda p: p.style.name == "Normal"
        and p.text.startswith("The paper addresses the following research question"),
    )
    if rq_p is None:
        raise SystemExit("could not locate the research-question paragraph in §1")

    replace_text(rq_p, NEW_RQ_PARA)

    # Insert blank Normal + audience paragraph immediately after rq_p, in that
    # order so the final sequence is: rq_p -> blank -> audience -> existing-blank
    # -> Heading 1 Background.
    audience_p_el = make_paragraph("Normal", NEW_AUDIENCE_PARA)
    rq_p._element.addnext(audience_p_el)
    blank_p_el = make_paragraph("Normal", "")
    rq_p._element.addnext(blank_p_el)
    print("[§1] split RQ paragraph; inserted unifying-hypothesis + audience text")

    # ---- 2. §2 Background P2: inject fintech-pressure sentence ----
    bg_p2 = find_para(
        doc,
        lambda p: p.style.name == "Normal"
        and p.text.startswith("The model addresses a structural tension"),
    )
    if bg_p2 is None:
        raise SystemExit("could not locate §2 P2 paragraph")

    original = bg_p2.text
    anchor = "high regulatory demands."
    if anchor not in original:
        raise SystemExit("fintech anchor not found in §2 P2")
    if "Revolut" in original:
        print("[§2] fintech sentence already present; skipping")
    else:
        new_bg_p2 = original.replace(
            anchor,
            anchor + FINTECH_INSERT,
            1,
        )
        replace_text(bg_p2, new_bg_p2)
        print("[§2] injected Revolut/Nordnet/Avanza fintech-pressure sentence")

    # ---- 3. §5.4 Implications P1: sharpen ----
    impl_p1 = find_para(
        doc,
        lambda p: p.style.name == "Normal"
        and p.text.startswith("For an incumbent universal bank already executing"),
    )
    if impl_p1 is None:
        raise SystemExit("could not locate §5.4 P1")
    replace_text(impl_p1, NEW_5_4_P1)
    print("[§5.4] replaced P1 with sharpened LO4 deliverable language")

    # ---- 4. Page breaks: only References, Appendix, back page ----
    body_h1_titles = {
        "Introduction",
        "Background",
        "Method",
        "Results",
        "Discussion",
        "Conclusion and limitations",
    }
    stripped = []
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() in body_h1_titles:
            if p._element.find(qn("w:pPr")) is not None and p._element.find(
                qn("w:pPr")
            ).find(qn("w:pageBreakBefore")) is not None:
                remove_page_break_before(p)
                stripped.append(p.text.strip())
    print(f"[page-break] removed pageBreakBefore from body H1: {stripped}")

    added = []
    for p in doc.paragraphs:
        if p.style.name == "Reference Heading" and p.text.strip() == "References":
            add_page_break_before(p)
            added.append("References")
        elif p.style.name == "Heading 1" and p.text.strip() == "Appendix":
            add_page_break_before(p)
            added.append("Appendix")
    print(f"[page-break] ensured pageBreakBefore on: {added}")

    # ---- 5. References: merge multi-paragraph entries; one blank between each ----
    in_refs = False
    ref_section_paras = []
    appendix_h1_el = None
    for p in doc.paragraphs:
        sname = p.style.name
        text = p.text.strip()
        if sname == "Reference Heading" and text == "References":
            in_refs = True
            continue
        if in_refs and sname == "Heading 1" and text == "Appendix":
            appendix_h1_el = p._element
            break
        if in_refs and sname == "Reference Text":
            ref_section_paras.append(p)

    if appendix_h1_el is None:
        raise SystemExit("could not locate Appendix H1 to anchor reference rebuild")

    # group source paragraphs into refs
    groups: list[list] = []
    current: list = []
    for p in ref_section_paras:
        text = p.text.strip()
        if not text:
            continue  # blanks are not boundaries; structure is rebuilt
        if REF_START.search(text):
            if current:
                groups.append(current)
            current = [p]
        else:
            if current:
                current.append(p)
            else:
                current = [p]
    if current:
        groups.append(current)

    print(f"[refs] discovered {len(groups)} merged reference entries")

    # Build new paragraphs by deepcopying runs from each group's source paras.
    new_para_els = []
    for i, group in enumerate(groups):
        new_p = make_paragraph("ReferenceText")
        for j, src_p in enumerate(group):
            for r in src_p._element.findall(qn("w:r")):
                new_p.append(deepcopy(r))
            if j < len(group) - 1:
                # join across paragraph boundary with a single space
                sp_r = OxmlElement("w:r")
                sp_t = OxmlElement("w:t")
                sp_t.text = " "
                sp_t.set(qn("xml:space"), "preserve")
                sp_r.append(sp_t)
                new_p.append(sp_r)
        new_para_els.append(new_p)
        if i < len(groups) - 1:
            new_para_els.append(make_paragraph("ReferenceText"))

    # Remove old Reference Text paragraphs.
    for p in ref_section_paras:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

    # Insert new paragraphs immediately before the Appendix Heading 1.
    for el in new_para_els:
        appendix_h1_el.addprevious(el)

    print("[refs] rebuilt reference list with one blank between every entry")

    # ---- save ----
    doc.save(str(TARGET))
    print(f"[saved] {TARGET.name}")


if __name__ == "__main__":
    patch()
