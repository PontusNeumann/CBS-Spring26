"""
patch_docx_round4.py

FROZEN SNAPSHOT — DO NOT RUN.

This script holds the body content (Introduction, Background, Method, all
of section 5 Discussion, Conclusion and limitations, and the References
list) as it stood on 2026-04-29, BEFORE the 2026-05-03 interview redesign.
Re-running it would revert: (a) the section 3 Method interview paragraph
to the stale three-FNZ-personas wording, (b) the section 5 paragraphs to
the round-4 prose without the round-6 in-text citations to FNZ Interview
A/B, Danske Bank Interview, ECB (2025), and Bell Harley & Bryman (2022),
(c) the Conclusion and limitations paragraph to the pre-Danske wording,
and (d) the references list to the pre-Bell, pre-ECB, pre-Danske state.

The current canonical patcher is scripts/patch_docx_round6.py.

If the docx ever needs to be rebuilt from a clean template, this file's
content remains the source of truth for the body text scaffolding, and
should be brought back online only after its METHOD_PARAS, DISC_5_x_PARAS,
CONCLUSION_PARAS, and REFERENCES_ENTRIES are re-synchronized with the
2026-05-03+ state.

To re-enable for that scenario, comment out the SystemExit guard below.
"""
import sys
sys.exit("patch_docx_round4.py is a frozen snapshot — do not run. Use patch_docx_round6.py.")
"""
Round-4 surgical patch. Replaces the body content of the active .docx with
the substantively rewritten and syllabus-aligned content from
paper_body_draft.md (as of 2026-04-29).

Sections replaced:
  - §1 Introduction (full body content)
  - §2 Background
  - §3 Method
  - §5.1 - §5.4 Discussion subsections
  - §6 Conclusion and limitations
  - References list (paragraphs under "Reference Heading", before "Appendix" H1)

Sections preserved as-is per user instruction:
  - §4 Results (interview tables remain as placeholders)
  - §5 Discussion heading itself (just the H1 "Discussion")
  - Appendix A (placeholder)
  - Appendix B (interview transcript skeleton)

The patcher:
  - Takes a timestamped backup of the active .docx in backup/ before any change.
  - Patches `word/document.xml` only via python-docx; no other package parts
    are touched.
  - Applies the spacing rules from Design.md §1.2 to every inserted paragraph
    (space-before=0, space-after=0, both autospacing flags zeroed, single-line
    spacing line=240 lineRule=auto on Normal-styled paragraphs).
  - Inserts a single blank Normal paragraph as separator between content
    paragraphs (per Design.md §1.2 default).
  - Front page <w:sdt> and back-page Back-Text paragraphs are not touched.
  - Em-dashes are stripped from inserted text per the in-class instruction
    "No use of m-dash" (course_overview.md).
"""
# `from __future__ import annotations` removed because this script is hard-
# disabled at the top of the file; it never executes past the sys.exit guard.

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
TARGET = ROOT / "KAN-CDSCO2401U_185912_DPD_Spring2026.docx"
BACKUP_DIR = ROOT / "backup"


# ============================== content (paragraphs by section) ==============================

INTRO_PARAS = [
    "European wealth management is being reorganized around shared digital infrastructure. Functions that were historically built, operated, and defended inside the universal bank are increasingly being standardized on external platforms that combine technology, operations, and regulated market connectivity. This shift matters because it changes both the economics of delivery and the basis on which banks compete. The immediate attraction is clear. A shared platform can lower duplication, simplify fragmented operating models, and help banks respond to regulatory pressure, margin pressure, and rising demands for digital service quality. The strategic consequence is less clear. Once the operational backbone is outsourced, the bank no longer competes through the same capabilities as before.",
    "This paper examines that shift through the case of FNZ. FNZ is a global wealth-management platform that partners with banks, insurers, and wealth managers to provide integrated technology, operations, and market infrastructure. As of April 2026, FNZ's public site describes the firm as supporting over US$2.4 trillion in assets on platform and nearly 30 million end investors (FNZ, n.d.). The scale itself is not the paper's focus. The analytical interest lies in what this scale says about the changing organization of wealth management, especially in Europe where incumbent banks are rethinking how much of the value chain should remain in-house.",
    "The case is timely because recent public mandates show large institutions choosing FNZ not only for software but for operating infrastructure. Swedbank's 2021 platform investment framed the decision in terms of efficiency, simplification, scalability, and stronger governance (FNZ, 2021). UniCredit's 2025 partnership in Germany similarly emphasized standardization, operational simplification, and a scalable post-trade platform (FNZ, 2025c). These decisions suggest that a platform such as FNZ is not being bought merely as a technology layer. It is being adopted as part of the bank's operating model.",
    "The paper addresses the following research question:",
    "**Why are European universal banks outsourcing wealth-platform operations to FNZ, and what does FNZ's end-to-end, assets-linked platform model imply for bank differentiation and concentration risk in European wealth management?**",
    "The underlying business problem is how incumbent banks can capture the efficiency benefits of FNZ-style outsourcing without losing strategic differentiation or creating unmanaged concentration risk. The paper's unifying hypothesis is that FNZ outsourcing is strategically rational only if banks treat it as a platform-governance and capability-rebuilding problem, not merely as a cost-saving IT project.",
    "The purpose is twofold. First, the paper explains why the outsourcing logic is economically attractive to incumbent banks. Second, it evaluates the managerial and sector-level consequences of that choice, including what banks must rebuild internally once a shared operational backbone is in place. The primary advisory audience is the incumbent European universal bank executive committee; the secondary audience is the European wealth-management sector and regulator. The analysis uses three theoretical lenses, each chosen to address one question the case raises. **Transaction-cost economics** explains why the bank-FNZ relationship settles into a long-duration hybrid form rather than a market or a hierarchy. **Multi-sided platform economics** explains why the wealth-platform layer concentrates on a small number of providers despite the absence of classic consumer network effects. **Dynamic capabilities** explains what banks must protect or rebuild internally once operational excellence is no longer a source of differentiation. The argument throughout is that platform outsourcing solves a real operating problem, but it also shifts competition away from operational excellence and toward advice, distribution, client ownership, and the strategic use of data.",
]

BACKGROUND_PARAS = [
    "FNZ's position in the industry helps explain why it is analytically useful. The firm presents itself as an end-to-end wealth-management platform rather than a point-solution provider. Its public material emphasizes the consolidation of fragmented systems, the integration of technology with business and investment operations, and the ability to support banks, wealth managers, insurers, and asset managers on a common infrastructure (FNZ, n.d.). In practical terms, this means that the value proposition extends beyond software implementation. The platform becomes the operating backbone through which institutions process advice workflows, onboarding, portfolio administration, custody arrangements, reporting, and other core activities.",
    "This model addresses a structural tension in wealth management. Banks face growing pressure to modernize digital journeys and reduce cost-to-serve, but they do so in a sector marked by legacy systems, jurisdictional variation, and high regulatory demands. Tech-led retail-finance competitors such as Revolut, Nordnet, and Avanza intensify the pressure by raising customer expectations for speed, cost, and digital service quality; this makes modernization more urgent, even though FNZ itself operates as B2B wealth infrastructure rather than as a direct consumer-fintech rival. A platform partner becomes attractive when it can absorb a large share of this complexity while spreading investment across many institutional clients.",
    "The competitive pressure is not only about service speed. Iansiti and Lakhani (2020) characterize the current era as one in which competitive advantage in financial services accrues to firms that can convert routine operations into data-driven learning loops at scale. Tech-led entrants of the Revolut, Nordnet, and Avanza type are built around such loops from inception; legacy universal banks are not. Modernizing the operating backbone is therefore a precondition for competing on the layers (advisory quality, personalization, data-driven cross-sell) where AI-era competition is actually decided. A platform such as FNZ becomes attractive in this framing because it standardizes the layers that no longer differentiate, freeing the bank to invest in the layers that do. The strategic risk is that the bank standardizes the backbone without making the corresponding investment, ending up cheaper to run but no more competitive against tech-led entrants than before. Swedbank's public rationale illustrates this logic clearly. The bank described its investment in a new savings platform as part of a strategic effort to improve efficiency, simplify processes, and create a stable, standardized, and scalable foundation, while preserving the bank's direct customer relationship (FNZ, 2021). The same case also highlighted increased regulatory pressure and the benefit of sharing platform costs and technological development across institutions (FNZ, 2021).",
    "Recent mandates suggest that this logic has strengthened rather than weakened. In September 2025, UniCredit selected FNZ to transform its securities-services operations in Germany, explicitly citing standardization, efficiency, and a scalable, cloud-based post-trade platform (FNZ, 2025c). In June 2025, Raymond James Ltd. announced a strategic partnership with FNZ to deliver a next-generation wealth-management platform, framing the decision around end-to-end infrastructure, straight-through processing, and digital capability (FNZ, 2025b). In November 2025, FNZ reported US$2.1 trillion of assets on platform while raising US$650 million from existing institutional shareholders (FNZ, 2025a). These examples differ in geography and business model, but they point in the same direction: institutions are no longer buying isolated digital tools. They are redesigning the infrastructure on which the wealth proposition rests.",
    "That shift creates the core tension of the paper. If the platform solves complexity, lowers duplication, and offers scale economies, outsourcing becomes economically persuasive. However, if several banks adopt the same operational backbone, the traditional sources of competitive differentiation become weaker. The strategic problem is not whether outsourcing can improve operating efficiency. It is whether banks can continue to differentiate once more of the underlying process architecture is shared. This is where the FNZ case becomes analytically rich. It is a case about efficiency, but it is also a case about boundary choice, platform concentration, and the changing location of competitive advantage.",
]

METHOD_PARAS = [
    "The paper is a theory-driven single-case analysis. FNZ is the case; three syllabus-grounded theoretical lenses are applied in depth to interpret it. The first lens is transaction-cost economics, limited to the coordination-cost versus production-cost distinction, the move-to-the-middle logic (Clemons, Reddi, & Row, 1993), and asset specificity (Williamson, 1985), refined for ICT-mediated transactions by Cordella (2006). The second is multi-sided platform economics, limited to indirect network effects, switching and multi-homing costs, and the conditions under which a regulated B2B market tends toward winner-take-most outcomes (McIntyre & Chintakananda, 2014). The third is dynamic capabilities, limited to sensing, seizing, and reconfiguring (Teece, 2007), supplemented by analogical reasoning and cognitive search (Gavetti & Rivkin, 2007) to explain how bank executives frame the outsourcing decision. These three lenses address three distinct but linked questions: why banks outsource the wealth-platform layer, why activity concentrates on a small number of providers, and what incumbent banks must protect or rebuild internally afterward.",
    "Evidence is drawn primarily from public-record material: FNZ corporate publications and client-mandate press releases, peer-reviewed research on financial-services outsourcing and platform economics, and European regulator publications on outsourcing and operational resilience (EBA, 2019; European Parliament & Council, 2022). The case selection criterion is analytical fit. FNZ is used because its end-to-end model, AUA-linked monetization, and recent large-bank mandates expose all three theoretical mechanisms in the same firm. Public mandates outside Europe (Raymond James Canada) are used only where they sharpen a mechanism already visible inside the European context.",
    "A supplementary evidence layer of three semi-structured elite interviews at FNZ (Strategy, Operations, and Client Management) is reported in §4. The sampling logic is theoretical: each respondent informs one analytical block more strongly than the others. Interview material, when present, deepens the public-record argument; the analysis in §5 does not depend on it.",
    "Prior operations work at FNZ informs interpretation of generic operating mechanisms such as reconciliation, transfer and transaction workflow automation, and regulatory-reporting validation. This is treated as reflexive case familiarity, not as data: no client data, named internal checks, colleagues, incidents, screenshots, or non-public operating details are reported. Where this familiarity sharpens a theoretical claim, it appears as a clearly bounded first-person reflection.",
    "The scope is intentionally narrow. The paper concerns the outsourcing of the wealth-management operating backbone and the competitive and governance consequences that follow. It does not map the global wealth-platform market or evaluate every component of FNZ's product portfolio.",
]

DISC_5_1_PARAS = [
    "Transaction-cost economics frames the bank-FNZ relationship as a boundary choice between hierarchy and market (Williamson, 1985). The traditional European universal bank ran wealth operations as a hierarchy: in-house custody, in-house reconciliation, in-house regulatory reporting, on bank-owned legacy systems. The coordination costs of that arrangement are high and rising. Each jurisdiction, product, and regulatory change must be absorbed by the bank's own operating estate, where production costs are dominated by manual hand-offs, brittle integrations across legacy systems, and dual-keying between front and back office. Clemons, Reddi, and Row (1993) argued that information technology lowers coordination costs faster than it lowers production costs, pulling firms away from full hierarchy without pushing them all the way to spot-market relationships. FNZ is a textbook instance of that move to the middle: an explicit, long-duration, contract-governed bilateral relationship that internalizes the operating backbone of multiple banks under one platform.",
    "Cordella (2006) refined the Clemons et al. argument by showing that ICT can simultaneously lower coordination costs *and* raise production costs, because complexity migrates rather than disappearing. FNZ inverts this risk by combining technology, operations, and regulated market connectivity in a single integrated stack. Varian (2010) generalizes the underlying mechanism: when ICT lowers the cost of mediating, monitoring, and contracting transactions, firms increasingly buy bundles of contract-governed services rather than build them, and the boundary of the firm shifts accordingly. Public client mandates make the move-to-the-middle structure visible. Swedbank framed its 2021 platform investment around efficiency, simplification, scalability, and stronger governance, language that maps directly to coordination-cost reduction (FNZ, 2021). UniCredit's 2025 securities-services mandate in Germany and Raymond James Canada's 2025 partnership use the same vocabulary of standardization and post-trade simplification (FNZ, 2025c, 2025b). Across these cases, the bank does not buy software. It pays a long-term partner to absorb fragmented operations, in exchange for a relational hybrid contract priced on assets under administration.",
    "Asset specificity is the residual TCE concern (Williamson, 1985). Once a bank's operating model is built around FNZ workflows, custody arrangements, and reporting outputs, the underlying investments are difficult to redeploy: process maps, change-management muscle, vendor-management routines, and regulatory filings are all FNZ-shaped. This produces lock-in, and lock-in is what makes the hybrid form unstable in the long run unless it is governed deliberately. The implication is that banks selecting FNZ are not making a procurement decision; they are choosing a long-horizon governance partner. From prior FNZ operations work, the lived version of this is concrete: I observed that automating reconciliation and transfer workflows reduces manual hand-offs and surfaces breaks earlier, but it does not remove the need for bank-side control, dual validation, and client-side accountability. The hybrid is real on both sides of the boundary.",
]

DISC_5_2_PARAS = [
    "Multi-sided platform economics explains why the wealth-platform layer concentrates on a small number of providers rather than supporting many parallel infrastructures. FNZ is not a classic two-sided consumer marketplace; it is a regulated B2B platform whose sides are institutional clients (banks, insurers, wealth managers) and the regulated market participants on which their workflows depend (custodians, exchanges, transfer agents, regulators). The indirect network effects are weaker than in consumer platforms but still material. Each new institutional client raises the marginal value of FNZ's regulatory connectivity, jurisdictional coverage, and standardized integrations, because the cost of building those integrations is incurred once and amortized across the installed base (Iansiti & Lakhani, 2020). Recent client wins illustrate the cumulative effect: by November 2025, FNZ reported US$2.1 trillion of assets on platform and a US$650 million capital raise from existing institutional shareholders (FNZ, 2025a), and as of April 2026 its public site lists over US$2.4 trillion in assets on platform and nearly 30 million end investors (FNZ, n.d.).",
    "The forces that produce winner-take-most outcomes in this market are the multi-homing and switching dimensions identified by McIntyre and Chintakananda (2014). Multi-homing is operationally infeasible for the core wealth backbone: a universal bank cannot run two parallel custody, reconciliation, and reporting stacks for the same book of business without duplicating cost and creating reconciliation breaks at every boundary. Switching costs are amplified by the asset specificity discussed in §5.1 and by the regulatory cost of re-certifying a new platform under EBA outsourcing guidance (EBA, 2019). Demand-side variety is also low: bank executives explicitly value standardization and predictability in this layer, as the public mandate language at Swedbank, UniCredit, and Raymond James shows. The combination is decisive. High switching cost, infeasible multi-homing, and low appetite for operational variety push the European wealth-platform market toward a small number of providers serving most of the AUA pool.",
    "This is rational at the firm level and creates concentration risk at the sector level. The European Banking Authority's outsourcing guidelines explicitly require institutions to assess concentration and substitutability for critical or important outsourced functions (EBA, 2019), and the Digital Operational Resilience Act extends this to a sector-wide oversight regime for critical ICT third-party providers (European Parliament & Council, 2022). The implication is that platform economics that make FNZ individually attractive create sectoral fragility that cannot be solved at the level of a single bank's contract.",
]

DISC_5_3_PARAS = [
    "When the operating backbone moves to a shared platform, the basis of competition for the bank moves with it. The dynamic-capabilities perspective identifies sensing, seizing, and reconfiguring as the higher-order capabilities that determine sustained advantage in environments where the underlying productive capabilities can be acquired in the market (Teece, 2007). FNZ adoption is the moment at which the wealth-management operating layer becomes acquirable in the market. The capability that previously differentiated stronger banks, the ability to run wealth operations cleanly across jurisdictions and regulatory regimes, is now table stakes on the FNZ platform and therefore no longer differentiating.",
    "What remains differentiating shifts toward the layers FNZ cannot replicate: client relationship and trust, advisory quality, distribution reach, brand, and the proprietary use of bank-owned client data. Iansiti and Lakhani (2020) describe the modern digital firm as built around an AI factory in which data pipelines and experimentation platforms convert routine operations into learning loops. For a bank that has outsourced the operating backbone, the data pipeline is partially co-owned with FNZ, and the experimentation platform must therefore be rebuilt at the layers the bank still controls: advisory tooling, personalization, cross-sell, and client lifecycle management. Constantiou, Joshi, and Stelmaszak (2023) extend this point by showing that data assets in platform ecosystems generate value only when the data-using firm builds the organizational capability to route, recombine, and act on them; without that capability, outsourcing the backbone leaves the bank with shared infrastructure and no proprietary data leverage. Stelmaszak, Joshi, and Constantiou (2025) develop the same logic for AI specifically, framing AI capability as an organizing capability that arises only from the deliberate cultivation of human-algorithm relations inside the firm, exactly the cultivation the bank cannot delegate to a platform partner.",
    "The cognitive dimension is what most often goes wrong. Gavetti and Rivkin (2007) show that strategic decisions are routinely framed by analogy to past decisions, and that the choice of analogy shapes both the options considered and the capabilities rebuilt afterward. Bank executive committees that frame an FNZ adoption as *another IT outsourcing*, analogous to a SaaS swap or an infrastructure migration, typically scope the program to cost reduction and run-rate stability. They under-invest in the capability rebuild that the dynamic-capabilities framework predicts is essential. A more accurate analogy is the foundational supply-chain partnership of automotive original-equipment manufacturers, where a deep, long-horizon relationship with a Tier-1 supplier coexists with deliberate internal investment in design, brand, and customer ownership. Banks that adopt this second analogy frame the decision as a triple package: operational outsourcing, platform-governance arrangement, and capability-rebuilding program. Banks that adopt the first analogy capture the cost saving and surrender the strategic position.",
]

DISC_5_4_PARAS = [
    "The three lenses converge on a clear recommendation. For an incumbent European universal bank, FNZ-style outsourcing is strategically rational only when treated as a triple package: an operational outsourcing decision, a platform-governance arrangement, and a deliberate capability-rebuilding program. Each leg fails on its own.",
    "The operational outsourcing decision should externalize the layers where standardization dominates differentiation: custody arrangements, reconciliation, transfer and transaction processing, regulatory reporting infrastructure, and core post-trade workflows. These are precisely the layers where Cordella's (2006) production-cost concern is mitigated by FNZ's integrated stack and where, in Varian's (2010) terms, the cost of mediating and monitoring the transaction has fallen far enough to favor a contract-governed bundle over an in-house build. The layers retained internally are those that carry the bank's competitive identity: client relationship, advisory quality, distribution and brand, and the proprietary use of bank-owned client data.",
    "The platform-governance arrangement must internalize the concentration risk that the bank's individual contract creates for the sector. Drawing on EBA (2019) and the Digital Operational Resilience Act (European Parliament & Council, 2022), the bank should require contractual portability of client data and process artefacts, a tested exit plan, structured roadmap influence proportional to its share of FNZ's AUA, and continuous resilience oversight aligned with DORA's critical-third-party regime. Operationally, this means FNZ adoptions should be governed by a permanent vendor-management capability inside the bank, not by a one-time procurement function.",
    "The capability-rebuilding program is the leg that the cognitive analogy most often hides. Banks should fund, from day one, a deliberate buildout of advisory tooling, data and AI capability on the layers they still control, and an experimentation platform on the client-facing edge (Iansiti & Lakhani, 2020; Constantiou, Joshi, & Stelmaszak, 2023; Stelmaszak, Joshi, & Constantiou, 2025). Without this leg, the bank captures the run-rate cost saving and arrives in five years with reduced operational headcount, an FNZ-shaped operating model, and no new differentiating capability, the worst possible outcome under the dynamic-capabilities lens.",
    "In concrete operating terms, the rebuild has four components. First, a client-data layer that consolidates the bank's view of each client across products, channels, and the FNZ-managed back office, governed by the bank rather than co-owned with the platform. Second, an advisory-tooling stack (recommendation engines, suitability and product-fit logic, scenario tools) that converts the data layer into adviser productivity and client outcomes. Third, an experimentation platform that runs controlled tests on advisory journeys, pricing, and onboarding, so the layers the bank still owns improve faster than competitors who only outsource. Fourth, a permanent vendor-management capability that operates the FNZ relationship as a governance partnership rather than a procurement contract, including roadmap influence, exit testing, and resilience oversight. Each component is owned, staffed, and budgeted as a distinct program from day one of the FNZ adoption; deferring any of the four reproduces the cognitive trap above.",
    "At the sector level, the concentration externality identified in §5.2 cannot be solved by any single bank's contract. The implication for the European regulator is that the DORA critical-third-party regime should be applied to wealth-platform infrastructure with the same seriousness as to cloud hyperscalers, with substitutability assessments, stress testing, and orderly resolution planning at the platform layer (EBA, 2019; European Parliament & Council, 2022). Without that layer of supervision, the rationality of each individual outsourcing decision aggregates into a sectoral fragility no bank has the standing to fix alone.",
]

CONCLUSION_PARAS = [
    "European universal banks are outsourcing wealth-platform operations to FNZ because the move-to-the-middle logic of transaction-cost economics makes the bilateral hybrid form the most efficient response to a coordination-cost problem the legacy hierarchy can no longer absorb. The same outsourcing concentrates the regulated B2B platform layer on a small number of providers, because multi-homing is operationally infeasible, switching costs are amplified by asset specificity and re-certification requirements, and demand for operational variety is low. For the bank, the consequence is that operational excellence is no longer a source of competitive differentiation; advantage shifts to advisory quality, client relationship, brand, and the proprietary use of client data, and is sustained only by a deliberate capability-rebuilding program of the type the dynamic-capabilities framework predicts. For the sector, individually rational outsourcing decisions aggregate into a concentration externality that requires platform-level supervision under the DORA framework. The paper's recommendation follows: incumbent banks should treat FNZ adoption as a triple package of operational outsourcing, platform-governance arrangement, and capability-rebuilding program, and the European regulator should treat the platform layer itself as critical third-party infrastructure.",
    "The analysis is limited by its single-case design and by its reliance on public-record material complemented by reflexive case familiarity. The supplementary FNZ-side interviews reported in §4 and Appendix B add depth where the public record is thin, but no analytical claim, theoretical move, or recommendation in §5 depends on them. Further research could test the capability-rebuilding argument across a larger sample of post-FNZ European banks and could quantify the concentration externality at the sector level using the European supervisor's third-party register.",
]

# References. Each entry is a single Reference Text paragraph.
REFERENCES_ENTRIES = [
    "Clemons, E. K., Reddi, S. P., & Row, M. C. (1993). The impact of information technology on the organization of economic activity: The move to the middle hypothesis. Journal of Management Information Systems, 10(2), 9-35.",
    "Constantiou, I., Joshi, M., & Stelmaszak, M. (2023). Organizations as digital enactment systems. Journal of the Association for Information Systems, 24(6), 1770-1798.",
    "Cordella, A. (2006). Transaction costs and information systems: Does IT add up? Journal of Information Technology, 21(3), 195-202.",
    "European Banking Authority. (2019). EBA guidelines on outsourcing arrangements (EBA/GL/2019/02). European Banking Authority.",
    "European Parliament & Council. (2022). Regulation (EU) 2022/2554 of the European Parliament and of the Council of 14 December 2022 on digital operational resilience for the financial sector (DORA). Official Journal of the European Union, L 333.",
    "FNZ. (n.d.). About FNZ. Retrieved April 28, 2026, from https://www.fnz.com/about",
    "FNZ. (2021, June 9). Swedbank selects FNZ to develop a new investment platform for its customers [Press release]. https://www.fnz.com/news",
    "FNZ. (2025a, November). FNZ raises USD 650 million from existing institutional shareholders [Press release]. https://www.fnz.com/news",
    "FNZ. (2025b, June). Raymond James Ltd. partners with FNZ to deliver next-generation wealth platform in Canada [Press release]. https://www.fnz.com/news",
    "FNZ. (2025c, September). UniCredit selects FNZ to transform securities-services operations in Germany [Press release]. https://www.fnz.com/news",
    "Gavetti, G., & Rivkin, J. W. (2007). On the origin of strategy: Action and cognition over time. Organization Science, 18(3), 420-439.",
    "Iansiti, M., & Lakhani, K. R. (2020). Competing in the age of AI: Strategy and leadership when algorithms and networks run the world. Harvard Business Review Press.",
    "McIntyre, D. P., & Chintakananda, A. (2014). Competing in network markets: Can the winner take all? Business Horizons, 57(1), 117-125.",
    "Stelmaszak, M., Joshi, M., & Constantiou, I. (2025). Artificial intelligence as an organizing capability arising from human-algorithm relations. Journal of Management Studies.",
    "Teece, D. J. (2007). Explicating dynamic capabilities: The nature and microfoundations of (sustainable) enterprise performance. Strategic Management Journal, 28(13), 1319-1350.",
    "Varian, H. R. (2010). Computer mediated transactions. American Economic Review, 100(2), 1-10.",
    "Williamson, O. E. (1985). The economic institutions of capitalism: Firms, markets, relational contracting. Free Press.",
]


# ============================== helpers ==============================

EM_DASH_RE = re.compile(r"\s*[—–]\s*")  # em-dash, en-dash


def clean_text(text: str) -> str:
    """Strip em-dashes per course rule. Keep regular hyphens."""
    return EM_DASH_RE.sub(", ", text)


def _new(tag):
    return OxmlElement(tag)


def _set(el, attr, val):
    el.set(qn(attr), val)


def _add_size_rpr(rPr, *, with_cs: bool = False):
    """Append <w:sz w:val=22/> (and optionally <w:szCs/>) to an rPr element.
    22 half-points = 11 pt, the body-text size per Design.md §1.1."""
    sz = _new("w:sz")
    _set(sz, "w:val", "22")
    rPr.append(sz)
    if with_cs:
        szCs = _new("w:szCs")
        _set(szCs, "w:val", "22")
        rPr.append(szCs)


def _make_normal_pPr():
    """pPr that matches the existing pre-patch body-Normal format:
    <w:pPr><w:spacing w:line="240" w:lineRule="auto"/>
    <w:rPr><w:sz w:val="22"/></w:rPr></w:pPr>

    No explicit pStyle: Word treats unstyled paragraphs as Normal. No explicit
    before/after: the Normal style's defaults are correct for body spacing in
    this template, and forcing zero collapsed inter-paragraph gaps in the
    round-4-first-attempt patch."""
    pPr = _new("w:pPr")
    spacing = _new("w:spacing")
    _set(spacing, "w:line", "240")
    _set(spacing, "w:lineRule", "auto")
    pPr.append(spacing)
    rPr = _new("w:rPr")
    _add_size_rpr(rPr, with_cs=False)  # paragraph mark size, controls blank-line height
    pPr.append(rPr)
    return pPr


def _make_run(text: str, *, bold: bool = False, italic: bool = False):
    """Body-text run with explicit 11pt size. Bold/italic toggles supported."""
    r = _new("w:r")
    rPr = _new("w:rPr")
    if bold:
        rPr.append(_new("w:b"))
        rPr.append(_new("w:bCs"))
    if italic:
        rPr.append(_new("w:i"))
        rPr.append(_new("w:iCs"))
    _add_size_rpr(rPr, with_cs=True)
    r.append(rPr)
    t = _new("w:t")
    _set(t, "xml:space", "preserve")
    t.text = text
    r.append(t)
    return r


def _tokenize_inline(text: str):
    """Split on **bold** and *italic* markers; return list of (kind, text)."""
    tokens = []
    for chunk in re.split(r"(\*\*[^*]+\*\*)", text):
        if chunk.startswith("**") and chunk.endswith("**") and len(chunk) > 4:
            tokens.append(("bold", chunk[2:-2]))
        else:
            for sub in re.split(r"(\*[^*]+\*)", chunk):
                if sub.startswith("*") and sub.endswith("*") and len(sub) > 2:
                    tokens.append(("italic", sub[1:-1]))
                elif sub:
                    tokens.append(("plain", sub))
    return tokens


def make_paragraph(style_name: str, text: str, *, parse_inline: bool = True):
    """Build a <w:p>. For style_name='Normal', emits the body paragraph format
    that matches the existing pre-patch document exactly: no pStyle, line=240,
    explicit 11pt size on the paragraph mark and on every run.

    For style_name='ReferenceText', emits the existing reference-list format:
    pStyle=ReferenceText, w:spacing w:after=0, w:ind w:right=0, plain run
    (size inherited from style = 9.5 pt per Design.md §1.1)."""
    text = clean_text(text)
    p = _new("w:p")

    if style_name == "Normal":
        p.append(_make_normal_pPr())
        if parse_inline:
            tokens = _tokenize_inline(text)
        else:
            tokens = [("plain", text)] if text else []
        for kind, run_text in tokens:
            p.append(_make_run(run_text, bold=(kind == "bold"), italic=(kind == "italic")))
    elif style_name == "ReferenceText":
        pPr = _new("w:pPr")
        pStyle = _new("w:pStyle")
        _set(pStyle, "w:val", "ReferenceText")
        pPr.append(pStyle)
        spacing = _new("w:spacing")
        _set(spacing, "w:after", "0")
        pPr.append(spacing)
        ind = _new("w:ind")
        _set(ind, "w:right", "0")
        pPr.append(ind)
        p.append(pPr)
        # Plain run, no rPr — ReferenceText style provides size and font.
        r = _new("w:r")
        t = _new("w:t")
        _set(t, "xml:space", "preserve")
        t.text = text
        r.append(t)
        p.append(r)
    else:
        raise ValueError(f"unsupported style_name {style_name!r}")

    return p


def make_blank_normal():
    """Single blank Normal paragraph used as inter-content separator. Matches
    the existing pre-patch blank-Normal format exactly: no pStyle, line=240,
    explicit 11pt on paragraph mark, no run. The 11pt mark gives the blank
    paragraph the same height as one body line."""
    p = _new("w:p")
    p.append(_make_normal_pPr())
    return p


def find_paragraph(doc, *, style_name: str, text_match: str):
    """Return the <w:p> element of the first paragraph matching style+text."""
    for p in doc.paragraphs:
        if p.style is None or p.style.name != style_name:
            continue
        if p.text.strip() == text_match:
            return p._element
    return None


def find_paragraph_by_text(doc, text_match: str):
    """Return the <w:p> element of the first paragraph whose stripped text
    starts with text_match (any style)."""
    for p in doc.paragraphs:
        if p.text.strip().startswith(text_match):
            return p._element
    return None


def replace_section(doc, *, anchor_el, stop_predicate, new_paras, style_name: str = "Normal"):
    """Delete every <w:p> sibling after anchor_el up to (exclusive) the first
    sibling matching stop_predicate. Then insert the new paragraphs (with
    blank-Normal separators between them) immediately after anchor_el.

    stop_predicate(el) returns True for the next-section anchor element."""
    parent = anchor_el.getparent()
    # 1. Collect victims.
    victims = []
    cur = anchor_el.getnext()
    while cur is not None and not stop_predicate(cur):
        if cur.tag == qn("w:p"):
            victims.append(cur)
        cur = cur.getnext()
    for v in victims:
        parent.remove(v)
    # 2. Build new paragraphs with blank-Normal separators between content paras.
    insert_after = anchor_el
    for idx, text in enumerate(new_paras):
        p_el = make_paragraph(style_name, text)
        insert_after.addnext(p_el)
        insert_after = p_el
        if idx < len(new_paras) - 1:
            blank = make_blank_normal()
            insert_after.addnext(blank)
            insert_after = blank
    print(f"[replace] {len(victims)} paragraphs removed, {len(new_paras)} content + separators inserted")


def _style_id(el):
    pPr = el.find(qn("w:pPr"))
    if pPr is None:
        return None
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return None
    return pStyle.get(qn("w:val"))


def make_h1_predicate(*texts):
    """Predicate matching a <w:p> that is a Heading1 (XML styleId) with
    stripped text in `texts`."""
    target_texts = set(texts)

    def pred(el):
        if el.tag != qn("w:p"):
            return False
        if _style_id(el) != "Heading1":
            return False
        text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        return text in target_texts

    return pred


def make_h2_predicate(*texts):
    target_texts = set(texts)

    def pred(el):
        if el.tag != qn("w:p"):
            return False
        if _style_id(el) != "Heading2":
            return False
        text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        return text in target_texts

    return pred


def make_h1_or_refheading_predicate(*texts):
    """Stop when we hit either a Heading1 with one of the texts or any
    ReferenceHeading (used for §6 -> References)."""
    target_texts = set(texts)

    def pred(el):
        if el.tag != qn("w:p"):
            return False
        sid = _style_id(el)
        if sid == "ReferenceHeading":
            return True
        if sid != "Heading1":
            return False
        text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        return text in target_texts

    return pred


# ============================== main patch ==============================


def patch():
    if not TARGET.exists():
        raise SystemExit(f"missing target: {TARGET}")

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    backup = BACKUP_DIR / f"{TARGET.stem}_{stamp}_pre-patch-r4.docx"
    shutil.copy2(TARGET, backup)
    print(f"[backup] {backup.name}")

    doc = Document(str(TARGET))

    # --- §1 Introduction ---
    intro_el = find_paragraph(doc, style_name="Heading 1", text_match="Introduction")
    if intro_el is None:
        raise SystemExit("could not locate Introduction H1")
    print("[§1 Introduction]")
    replace_section(
        doc,
        anchor_el=intro_el,
        stop_predicate=make_h1_predicate("Background"),
        new_paras=INTRO_PARAS,
    )

    # --- §2 Background ---
    bg_el = find_paragraph(doc, style_name="Heading 1", text_match="Background")
    if bg_el is None:
        raise SystemExit("could not locate Background H1")
    print("[§2 Background]")
    replace_section(
        doc,
        anchor_el=bg_el,
        stop_predicate=make_h1_predicate("Method"),
        new_paras=BACKGROUND_PARAS,
    )

    # --- §3 Method ---
    method_el = find_paragraph(doc, style_name="Heading 1", text_match="Method")
    if method_el is None:
        raise SystemExit("could not locate Method H1")
    print("[§3 Method]")
    replace_section(
        doc,
        anchor_el=method_el,
        stop_predicate=make_h1_predicate("Results"),
        new_paras=METHOD_PARAS,
    )

    # --- §4 Results: SKIPPED per user instruction ---
    print("[§4 Results] skipped (interview tables retained)")

    # --- §5.x Discussion subsections. Each H2 anchors its own subsection. ---
    h2_5_1 = find_paragraph(
        doc, style_name="Heading 2", text_match="Transaction costs and the move to the middle"
    )
    if h2_5_1 is None:
        raise SystemExit("could not locate §5.1 H2")
    print("[§5.1]")
    replace_section(
        doc,
        anchor_el=h2_5_1,
        stop_predicate=make_h2_predicate(
            "Platform economics and concentration",
            "Dynamic capabilities and the new basis of differentiation",
            "Implications for banks and sector governance",
        ),
        new_paras=DISC_5_1_PARAS,
    )

    h2_5_2 = find_paragraph(
        doc, style_name="Heading 2", text_match="Platform economics and concentration"
    )
    if h2_5_2 is None:
        raise SystemExit("could not locate §5.2 H2")
    print("[§5.2]")
    replace_section(
        doc,
        anchor_el=h2_5_2,
        stop_predicate=make_h2_predicate(
            "Dynamic capabilities and the new basis of differentiation",
            "Implications for banks and sector governance",
        ),
        new_paras=DISC_5_2_PARAS,
    )

    h2_5_3 = find_paragraph(
        doc,
        style_name="Heading 2",
        text_match="Dynamic capabilities and the new basis of differentiation",
    )
    if h2_5_3 is None:
        raise SystemExit("could not locate §5.3 H2")
    print("[§5.3]")
    replace_section(
        doc,
        anchor_el=h2_5_3,
        stop_predicate=make_h2_predicate("Implications for banks and sector governance"),
        new_paras=DISC_5_3_PARAS,
    )

    h2_5_4 = find_paragraph(
        doc, style_name="Heading 2", text_match="Implications for banks and sector governance"
    )
    if h2_5_4 is None:
        raise SystemExit("could not locate §5.4 H2")
    print("[§5.4]")
    replace_section(
        doc,
        anchor_el=h2_5_4,
        stop_predicate=make_h1_predicate("Conclusion and limitations"),
        new_paras=DISC_5_4_PARAS,
    )

    # --- §6 Conclusion and limitations ---
    concl_el = find_paragraph(doc, style_name="Heading 1", text_match="Conclusion and limitations")
    if concl_el is None:
        raise SystemExit("could not locate Conclusion H1")
    print("[§6 Conclusion and limitations]")
    # Stop predicate: any Reference Heading or H1 named Appendix.
    replace_section(
        doc,
        anchor_el=concl_el,
        stop_predicate=make_h1_or_refheading_predicate("Appendix"),
        new_paras=CONCLUSION_PARAS,
    )

    # --- References list ---
    refs_el = find_paragraph(doc, style_name="Reference Heading", text_match="References")
    if refs_el is None:
        raise SystemExit("could not locate Reference Heading")
    print("[References]")
    replace_section(
        doc,
        anchor_el=refs_el,
        stop_predicate=make_h1_predicate("Appendix"),
        new_paras=REFERENCES_ENTRIES,
        style_name="ReferenceText",
    )

    # --- save ---
    doc.save(str(TARGET))
    print(f"[saved] {TARGET.name}")


if __name__ == "__main__":
    patch()
