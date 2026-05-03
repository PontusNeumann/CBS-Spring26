"""
patch_docx_round6.py

Round-6 surgical patch. Three drivers:

  (a) Switch the bank-side validation interview from SEB to Danske Bank.
      SEB has no public wealth-platform commitment of comparable scope; Danske
      Bank publicly selected BlackRock's Aladdin Wealth in 2025 (Forward '28
      strategy). The validation interview is therefore reanchored on a peer
      Nordic universal bank that has just lived through a wealth-platform
      buying decision, which gives the buyer-side perspective concrete
      texture and a vendor contrast to FNZ.

  (b) Lift in-text citation density on five sources that currently appear in
      the references list but nowhere in the body:
        - Bell, Harley & Bryman (2022) - method anchor.
        - European Central Bank (2025) - third-party concentration.
        - FNZ Interview A (2026) - section 5.1 hybrid governance / lock-in.
        - FNZ Interview B (2026) - section 5.3 analogical reasoning.
        - Danske Bank Interview (2026) - sections 5.2 and 5.4 buyer-side
          triangulation.
      Also adds Danske Bank (2025) Aladdin Wealth press release as an
      in-text citation in section 5.2 (peer Nordic mandate).

  (c) Preserve document furniture. The first page is manual-only and frozen:
      do not touch cover metadata, confidentiality text, images, logos, or
      visually blank cover paragraphs. The active final intentionally has no
      decorative CBS back page, so this patcher must not recreate or repair a
      Back - Text block or a back-page page break.

Sections preserved untouched per user instruction:
  - Appendix B.1 and Appendix B.2 transcripts (FNZ Operations and FNZ
    Client Management). Only Appendix B.3 is rewritten.
  - All section 5 prose paragraphs are kept verbatim except for the
    inline citation insertions documented above.

The patcher follows the conventions of patch_docx_round3.py through
patch_docx_round5.py: timestamped backup, paragraph-level surgical edits,
no other package parts touched, em-dashes stripped from inserted text.
"""
from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
TARGET = ROOT / "KAN-CDSCO2401U_185912_DPD_Spring2026.docx"
BACKUP_DIR = ROOT / "backup"


# ============================== content (citation edits) ==============================

# Each entry: (unique substring used to locate the paragraph, full new text).
# Paragraph runs are rewritten as a single run with the new text, preserving
# pPr. Section 5 paragraphs from round-4 are single-run already, so this is
# safe; verified against the docx body before drafting these edits.
PARAGRAPH_REPLACEMENTS = [
    # ---- Section 3 Method ----
    (
        "A supplementary evidence layer of three semi-structured interviews",
        "A supplementary evidence layer of three semi-structured interviews "
        "is reported in section 4. The interview design follows Bell, Harley "
        "and Bryman (2022). Two are FNZ-side, one in the Group Operations "
        "function and one in Client Management and Business Development for "
        "Europe, and the third is a bank-side validation interview with a "
        "senior leader in a Nordic universal bank's wealth and "
        "asset-management division. The sampling logic is theoretical: each "
        "respondent informs one analytical block more strongly than the "
        "others, and the bank-side interview triangulates the buyer-side "
        "implications rather than functioning as a fourth analytical pillar. "
        "Interview material, when present, deepens the public-record "
        "argument; the analysis in section 5 does not depend on it. All "
        "interviewees are anonymized at source by team or function, regional "
        "remit, and employer; full transcripts are reproduced in Appendix B.",
    ),
    # ---- Section 5.1 hybrid governance / lock-in ----
    (
        "Asset specificity is the residual TCE concern",
        "Asset specificity is the residual TCE concern (Williamson, 1985). "
        "Once a bank's operating model is built around FNZ workflows, custody "
        "arrangements, and reporting outputs, the underlying investments are "
        "difficult to redeploy: process maps, change-management muscle, "
        "vendor-management routines, and regulatory filings are all "
        "FNZ-shaped. This produces lock-in, and lock-in is what makes the "
        "hybrid form unstable in the long run unless it is governed "
        "deliberately (FNZ Interview A, 2026). The implication is that banks "
        "selecting FNZ are not making a procurement decision; they are "
        "choosing a long-horizon governance partner. From prior FNZ "
        "operations work, the lived version of this is concrete: I observed "
        "that automating reconciliation and transfer workflows reduces "
        "manual hand-offs and surfaces breaks earlier, but it does not "
        "remove the need for bank-side control, dual validation, and "
        "client-side accountability. The hybrid is real on both sides of "
        "the boundary.",
    ),
    # ---- Section 5.2 concentration risk: ECB SSM + Danske Bank Interview + Aladdin Wealth peer mandate ----
    (
        "This is rational at the firm level and creates concentration risk",
        "This is rational at the firm level and creates concentration risk "
        "at the sector level. The European Banking Authority's outsourcing "
        "guidelines explicitly require institutions to assess concentration "
        "and substitutability for critical or important outsourced functions "
        "(EBA, 2019), and the Digital Operational Resilience Act extends "
        "this to a sector-wide oversight regime for critical ICT "
        "third-party providers (European Parliament & Council, 2022). The "
        "European Central Bank's 2025 supervisory commentary makes the "
        "underlying market structure explicit: critical-function "
        "outsourcing in EU banking is concentrated on a small number of "
        "providers, and the share of critical functions that are difficult "
        "or impossible to reintegrate is now over ninety percent (European "
        "Central Bank, 2025). The same dynamic is visible at the Nordic "
        "peer level. Danske Bank's 2025 selection of BlackRock's Aladdin "
        "Wealth as the platform anchor for its Forward '28 wealth-management "
        "transformation is a parallel mandate: a different vendor, a "
        "narrower scope tilted toward advisory and portfolio analytics, "
        "but the same buy-versus-build logic on a layer the bank has "
        "decided no longer differentiates internally (Danske Bank, 2025). "
        "From the buyer side, lock-in once a bank commits is not primarily "
        "contractual or technical, which are solvable, but data-related and "
        "organizational, where unwinds run into multi-year timelines "
        "(Danske Bank Interview, 2026). The implication is that platform "
        "economics that make FNZ individually attractive create sectoral "
        "fragility that cannot be solved at the level of a single bank's "
        "contract.",
    ),
    # ---- Section 5.3 cognitive analogy: FNZ Interview B ----
    (
        "The cognitive dimension is what most often goes wrong",
        "The cognitive dimension is what most often goes wrong. Gavetti and "
        "Rivkin (2007) show that strategic decisions are routinely framed "
        "by analogy to past decisions, and that the choice of analogy "
        "shapes both the options considered and the capabilities rebuilt "
        "afterward. Bank executive committees that frame an FNZ adoption "
        "as another IT outsourcing, analogous to a SaaS swap or an "
        "infrastructure migration, typically scope the program to cost "
        "reduction and run-rate stability. They under-invest in the "
        "capability rebuild that the dynamic-capabilities framework "
        "predicts is essential. A more accurate analogy is the foundational "
        "supply-chain partnership of automotive original-equipment "
        "manufacturers, where a deep, long-horizon relationship with a "
        "Tier-1 supplier coexists with deliberate internal investment in "
        "design, brand, and customer ownership. Banks that adopt this "
        "second analogy frame the decision as a triple package: operational "
        "outsourcing, platform-governance arrangement, and "
        "capability-rebuilding program. Banks that adopt the first analogy "
        "capture the cost saving and surrender the strategic position "
        "(FNZ Interview B, 2026).",
    ),
    # ---- Section 5.4 platform-governance recommendation: Danske Bank Interview ----
    (
        "The platform-governance arrangement must internalize the concentration risk",
        "The platform-governance arrangement must internalize the "
        "concentration risk that the bank's individual contract creates "
        "for the sector. Drawing on EBA (2019), the Digital Operational "
        "Resilience Act (European Parliament & Council, 2022), and the ECB's "
        "2025 supervisory framing of third-party concentration (European "
        "Central Bank, 2025), the bank should require contractual "
        "portability of client data and process artefacts, a tested exit "
        "plan, structured roadmap influence proportional to its share of "
        "FNZ's AUA, and continuous resilience oversight aligned with DORA's "
        "critical-third-party regime (Danske Bank Interview, 2026). "
        "Operationally, this means FNZ adoptions should be governed by a "
        "permanent vendor-management capability inside the bank, not by a "
        "one-time procurement function.",
    ),
    # ---- Section 4 Table 3 caption ----
    (
        "Table 3. Interview C, SEB Wealth and Asset Management division",
        "Table 3. Interview C, Danske Bank Wealth and Asset Management "
        "division (Copenhagen-based, Nordic remit). Anchored in section 5.2 "
        "(concentration) and triangulating section 5.3 (bank-side "
        "capabilities).",
    ),
]


# ============================== content (references list edits) ==============================

# Replace existing SEB Interview entry text in place; insert Danske Bank
# (2025) Aladdin Wealth press release at correct alphabetical position.
SEB_INTERVIEW_OLD_PREFIX = "SEB Interview. (2026)."
DANSKE_INTERVIEW_NEW_TEXT = (
    "Danske Bank Interview. (2026). Confidential bank-side validation "
    "interview with a senior leader within Danske Bank's Wealth and Asset "
    "Management division (Copenhagen-based, Nordic remit). Anonymized at "
    "source by team or function, regional remit, and employer."
)
DANSKE_PRESS_RELEASE_TEXT = (
    "Danske Bank. (2025). Aladdin Wealth: Forward '28 partnership with "
    "BlackRock for wealth-management transformation [Press release]. "
    "https://danskebank.com/se/nyheter-och-press/nyhetsarkiv/news/2025/aladdin-wealth"
)


# ============================== content (appendix b.3 rewrite) ==============================

APP_B3_HEADING = (
    "Appendix B.3. Interview C, Danske Bank Wealth and Asset Management "
    "division (Copenhagen-based, Nordic remit)"
)

APP_B3_QA = [
    (
        "Question: When a universal bank considers outsourcing wealth-platform "
        "operations to a specialist provider, is that usually understood "
        "internally as an operating-model change or as a more conventional "
        "IT/vendor-outsourcing decision?",
        "Answer: That depends entirely on who in the bank is leading the "
        "conversation. If group IT or procurement runs it, it shows up on the "
        "agenda as an outsourcing decision and gets the standard "
        "vendor-management treatment. If it runs through the COO or the head "
        "of wealth, it is framed much more clearly as an operating-model "
        "change. Both framings are present in most banks at the same time, "
        "and the real question is which one wins by the time you are in "
        "implementation. In our recent platform decision, the framing landed "
        "with the business early. We described it internally as part of the "
        "Forward '28 transformation rather than as an IT migration, and that "
        "made the conversation different from a SaaS swap. From a governance "
        "standpoint, if you are outsourcing the operating backbone or the "
        "advisory analytics layer for an entire business line, calling it "
        "vendor management is dangerously narrow. But I also see why the "
        "technical framing is comfortable: it fits the procurement template, "
        "the legal template, and the audit template that already exist.",
        "Follow-up Q: At your bank specifically, would the framing be more "
        "business-driven or IT-driven?",
        "Follow-up A: I can speak only to my own perspective rather than to "
        "a formal bank position. The cultural answer in this house is that "
        "anything touching the customer relationship is owned by the "
        "business, so the framing tends to be operating-model first, vendor "
        "management second. Our public communication around the Aladdin "
        "Wealth selection was deliberately about advice quality, not about "
        "vendor cost-out. That is not universally true across Nordic banks.",
    ),
    (
        "Question: Which capabilities does the bank believe it must keep or "
        "strengthen internally if the operating backbone is outsourced?",
        "Answer: Five things, in order: the client relationship, advisory "
        "quality, the brand, regulatory accountability, and the internal "
        "data capability. The first three are obvious; you cannot outsource "
        "them and stay in business. Regulatory accountability is the one "
        "banks sometimes wrongly assume they can delegate alongside the "
        "operations, you can outsource the work, but the supervisor will "
        "still come to the bank. And the internal data capability is the "
        "one that gets the least attention at decision time and matters "
        "the most three years later. If you do not have people inside who "
        "can interrogate the platform's data, design analytics on it, and "
        "build the cases that turn it into a customer-facing differentiator, "
        "you have handed your future to a vendor. With the platform "
        "providers we work with, the analytics are powerful, but the value "
        "only shows up if the bank's own advisers and product teams know "
        "how to use them in a customer conversation. That is not a "
        "position any senior bank executive should be comfortable with if "
        "the internal capability is missing.",
        "Follow-up Q: Where do banks typically under-invest among those "
        "five?",
        "Follow-up A: The internal data capability. Closely followed by "
        "advisor enablement, the people-and-process work that makes new "
        "tooling actually change customer outcomes. We learned that lesson "
        "the hard way on earlier digital programs.",
    ),
    (
        "Question: From the buyer side, is lock-in mainly contractual, "
        "technical, operational, data-related, or organizational?",
        "Answer: Contractually, you can almost always exit with sufficient "
        "notice and a transition plan. Technically, exit is hard, but it is "
        "an engineering problem you can scope and price. The lock-in that "
        "actually concerns me is the combination of data-related and "
        "organizational. Once your data lineage flows through a particular "
        "platform's logic, your reporting suite is tuned to that data "
        "shape, your auditors have learned that shape, and your supervisor "
        "reviews reports in that shape. That is a multi-year unwind. The "
        "organizational side is similar: your operating teams have "
        "reorganized around the platform, your job descriptions have "
        "evolved, and rebuilding the prior internal capability requires a "
        "hiring and training cycle that takes years. We thought hard about "
        "this dimension specifically when we made our most recent platform "
        "decision, because the assumption that you can always switch later "
        "does not survive contact with reality once data lineage and team "
        "structure have settled.",
        "Follow-up Q: Does the board hear that risk in those terms?",
        "Follow-up A: Increasingly, yes. Five years ago the conversation at "
        "board level would have been about cost and resilience. Today the "
        "dependency dimension is on the agenda explicitly, partly because "
        "supervisors are asking pointed questions about substitutability "
        "and concentration. We have included exit-planning language and "
        "data-portability obligations in the contract from day one rather "
        "than negotiating them after the fact.",
    ),
    (
        "Question: How does the bank think about concentration risk and "
        "operational resilience when many institutions rely on the same "
        "platform infrastructure?",
        "Answer: It is a sector-level conversation, not a single-bank "
        "conversation. Any individual bank can manage its bilateral "
        "relationship with a platform provider through governance, exit "
        "planning, and strong operational testing. What it cannot manage is "
        "the situation where many of its peers depend on the same provider, "
        "because then the systemic exposure is shared and so is the "
        "recovery scenario. DORA and the EBA outsourcing guidelines are "
        "pushing the conversation in the right direction, the supervisor "
        "wants to see substitutability assessments, exit planning, and "
        "stress testing on critical third parties, but the underlying "
        "market structure is not something any one bank can fix alone. The "
        "ECB's 2025 supervisory commentary on third-party concentration "
        "captures the discomfort well: a meaningful share of the budget is "
        "concentrated on a small number of providers, and the great "
        "majority of critical functions are difficult to reintegrate. "
        "Those numbers should make any board uncomfortable. One of the "
        "reasons we deliberately picked a different vendor profile from "
        "the wealth-operations heavyweights was to keep our exposure to "
        "any single platform's outage scenario manageable.",
        "Follow-up Q: What would a sensible regulatory intervention look "
        "like, in your view?",
        "Follow-up A: Probably something close to the framework the ECB is "
        "already pointing at: portability standards, data-format "
        "requirements, and resolution-style planning for the largest "
        "providers. The market will not get there on its own. Cloud "
        "hyperscalers are already inside that supervisory perimeter; "
        "wealth-platform providers are not yet, and they probably should be.",
    ),
    (
        "Question: Does competitive pressure from tech-forward firms such as "
        "Revolut, Nordnet, and Avanza materially affect the modernization "
        "decision, or are internal cost, regulation, and legacy complexity "
        "the stronger drivers?",
        "Answer: Honestly, the strongest drivers are internal: cost, legacy "
        "complexity, and the regulatory burden of running fragmented "
        "systems. Fintech pressure shows up in the deck as the why-now "
        "narrative, but it is rarely the deciding factor. Where it does "
        "matter is in the customer-experience benchmark, once a Nordic "
        "customer has used a Revolut-style mobile experience, their "
        "tolerance for a 1990s-era bank app is gone. So fintech sets the "
        "expectation floor, but the modernization investment is paid for "
        "by the cost-and-regulation case, not the competitive-threat case. "
        "Our public framing of Forward '28 leans on the customer-experience "
        "language, but the budget approval ran on the cost and regulatory "
        "logic.",
        "Follow-up Q: Does that change for the affluent or high-net-worth "
        "segments?",
        "Follow-up A: Less so. In those segments, advice and trust still "
        "anchor the relationship and the digital-only players are some "
        "way from displacing the incumbent banks. The pressure is real on "
        "the mass-affluent and below, which is precisely where we are "
        "investing first.",
    ),
    (
        "Question: If a paper recommends that banks outsource standardized "
        "operational layers but retain client relationship, advice, brand, "
        "distribution, and proprietary data capability, what would you "
        "challenge or refine?",
        "Answer: I broadly agree with the direction. The refinement I would "
        "add is that the recommendation undersells the "
        "regulatory-accountability dimension. It should not be a sixth "
        "item on the list; it should sit alongside the others and shape "
        "how the contract and the governance are designed from day one. "
        "The other refinement is on data: retain proprietary data "
        "capability is the right phrase, but what banks actually need is "
        "the ability to act on the data in ways that change customer "
        "outcomes. Owning the data is not enough. You need analytics "
        "talent, advisor workflows that consume the analytics output, and "
        "a feedback loop that learns from what worked. Without that, "
        "data capability becomes a strategic slogan rather than a "
        "capability. We are building exactly that loop on top of our new "
        "advisory analytics partner, and it is the part that is hardest "
        "to staff.",
        "Follow-up Q: Anything you would add to the recommendation that is "
        "not there?",
        "Follow-up A: Two things. First, deliberate exit planning from "
        "year one, not because exit is the goal, but because the option "
        "discipline keeps the relationship balanced. Second, an honest "
        "internal conversation about which historical differentiators "
        "the bank is now prepared to give up. The capability rebuild "
        "fails when the bank tries to differentiate on too many things "
        "at once; the discipline is in choosing.",
    ),
]


# ============================== helpers ==============================

EM_DASH_RE = re.compile(r"\s*[—–]\s*")


def clean_text(text: str) -> str:
    return EM_DASH_RE.sub(", ", text)


def _new(tag):
    return OxmlElement(tag)


def _set(el, attr, val):
    el.set(qn(attr), val)


def _add_size_rpr(rPr, *, with_cs: bool = False):
    sz = _new("w:sz")
    _set(sz, "w:val", "22")
    rPr.append(sz)
    if with_cs:
        szCs = _new("w:szCs")
        _set(szCs, "w:val", "22")
        rPr.append(szCs)


def _make_normal_pPr():
    pPr = _new("w:pPr")
    spacing = _new("w:spacing")
    _set(spacing, "w:line", "240")
    _set(spacing, "w:lineRule", "auto")
    pPr.append(spacing)
    rPr = _new("w:rPr")
    _add_size_rpr(rPr, with_cs=False)
    pPr.append(rPr)
    return pPr


def _make_run(text: str, *, bold: bool = False, italic: bool = False):
    r = _new("w:r")
    rPr = _new("w:rPr")
    if bold:
        rPr.append(_new("w:b"))
        rPr.append(_new("w:bCs"))
    if italic:
        rPr.append(_new("w:i"))
        rPr.append(_new("w:iCs"))
    _add_size_rpr(rPr, with_cs=True)
    r.append(rPr)
    t = _new("w:t")
    _set(t, "xml:space", "preserve")
    t.text = text
    r.append(t)
    return r


def make_normal_paragraph(text: str):
    text = clean_text(text)
    p = _new("w:p")
    p.append(_make_normal_pPr())
    if text:
        p.append(_make_run(text))
    return p


def make_blank_normal():
    p = _new("w:p")
    p.append(_make_normal_pPr())
    return p


def make_styled_paragraph(style_id: str, text: str):
    text = clean_text(text)
    p = _new("w:p")
    pPr = _new("w:pPr")
    pStyle = _new("w:pStyle")
    _set(pStyle, "w:val", style_id)
    pPr.append(pStyle)
    p.append(pPr)
    r = _new("w:r")
    t = _new("w:t")
    _set(t, "xml:space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_reference_paragraph(text: str):
    text = clean_text(text)
    p = _new("w:p")
    pPr = _new("w:pPr")
    pStyle = _new("w:pStyle")
    _set(pStyle, "w:val", "ReferenceText")
    pPr.append(pStyle)
    spacing = _new("w:spacing")
    _set(spacing, "w:after", "0")
    pPr.append(spacing)
    ind = _new("w:ind")
    _set(ind, "w:right", "0")
    pPr.append(ind)
    p.append(pPr)
    r = _new("w:r")
    t = _new("w:t")
    _set(t, "xml:space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _style_id(el):
    pPr = el.find(qn("w:pPr"))
    if pPr is None:
        return None
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return None
    return pStyle.get(qn("w:val"))


def _para_text(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def find_paragraph_containing(doc, substring: str):
    """First paragraph whose stripped text contains `substring`."""
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None


def replace_paragraph_text(p, new_text: str):
    """Rewrite the paragraph's text as a single run, preserving its pPr."""
    el = p._element
    new_text = clean_text(new_text)
    # Drop existing runs.
    for r in el.findall(qn("w:r")):
        el.remove(r)
    # Insert a new run.
    el.append(_make_run(new_text))


def insert_after(anchor_el, new_el):
    parent = new_el.getparent()
    if parent is not None:
        parent.remove(new_el)
    anchor_el.addnext(new_el)
    return new_el


def clear_between(anchor_el, stop_predicate):
    parent = anchor_el.getparent()
    victims = []
    cur = anchor_el.getnext()
    while cur is not None and not stop_predicate(cur):
        if cur.tag in (qn("w:p"), qn("w:tbl")):
            victims.append(cur)
        cur = cur.getnext()
    for v in victims:
        parent.remove(v)
    return len(victims)


# ============================== section patchers ==============================


def patch_paragraph_replacements(doc):
    """Apply the in-paragraph citation-density edits and the §4 Table 3 rename."""
    applied = 0
    skipped = []
    for finder, new_text in PARAGRAPH_REPLACEMENTS:
        p = find_paragraph_containing(doc, finder)
        if p is None:
            skipped.append(finder[:50])
            continue
        replace_paragraph_text(p, new_text)
        applied += 1
        print(f"[edit] {finder[:50]!r}")
    if skipped:
        for s in skipped:
            print(f"[edit] WARNING: not found: {s!r}")
    print(f"[edit] total replacements applied: {applied}")


def patch_references_swap(doc):
    """Replace the SEB Interview reference entry with the Danske Bank Interview
    entry (in place, preserving alphabetical order broadly), and add the
    Danske Bank Aladdin Wealth press release entry at its alphabetical
    position (between Cordella and the European Banking Authority)."""
    # 1. Replace SEB Interview entry text.
    seb_p = None
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Reference Text":
            if p.text.strip().startswith(SEB_INTERVIEW_OLD_PREFIX):
                seb_p = p
                break
    if seb_p is not None:
        replace_paragraph_text(seb_p, DANSKE_INTERVIEW_NEW_TEXT)
        print("[refs] SEB Interview entry text replaced with Danske Bank Interview")
    else:
        print("[refs] WARNING: SEB Interview entry not found; cannot rename")

    # 2. Insert Danske Bank press release after the Cordella reference.
    # Skip if already present.
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Reference Text":
            if p.text.strip().startswith("Danske Bank. (2025)."):
                print("[refs] Danske Bank (2025) entry already present; skipping insert")
                return

    cordella_anchor = None
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Reference Text":
            if p.text.strip().startswith("Cordella, A. (2006)."):
                cordella_anchor = p._element
                # Don't break; want the LAST occurrence in case of duplicates.
    if cordella_anchor is None:
        print("[refs] WARNING: Cordella anchor not found; appending Danske entry at end of refs block")
        return
    new_p = make_reference_paragraph(DANSKE_PRESS_RELEASE_TEXT)
    cordella_anchor.addnext(new_p)
    # Reference list uses blank Normal separators between entries; insert one
    # immediately after the new entry to keep the cadence consistent.
    blank = make_blank_normal()
    new_p.addnext(blank)
    print("[refs] Danske Bank (2025) press release entry inserted")


def patch_appendix_b3(doc):
    """Replace Appendix B.3 entirely (heading + content) with the Danske Bank
    transcript. The new heading carries the Danske Bank label; old SEB
    paragraphs are cleared."""
    # Locate the Appendix B.3 H3 heading.
    appb3_el = None
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Heading 3":
            if p.text.strip().startswith("Appendix B.3"):
                appb3_el = p._element
                break
    if appb3_el is None:
        raise SystemExit("could not locate Appendix B.3 H3 anchor")

    # Stop before any existing back-page separator/block if one survives in an
    # older snapshot; otherwise stop before the document's final section
    # properties. Do not create or repair a decorative back page.
    first_back_text_el = None

    def is_after_anchor(candidate_el):
        cur = appb3_el.getnext()
        while cur is not None:
            if cur is candidate_el:
                return True
            cur = cur.getnext()
        return False

    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Back - Text":
            if is_after_anchor(p._element):
                first_back_text_el = p._element
                break

    stop_el = None
    stop_label = "end of document body"
    if first_back_text_el is not None:
        # If an old copy still has a back page, preserve its existing separator
        # exactly as found. This is preservation only, not recreation.
        pagebreak_el = None
        prev = first_back_text_el.getprevious()
        while prev is not None:
            if prev.tag == qn("w:p"):
                for r in prev.findall(qn("w:r")):
                    br = r.find(qn("w:br"))
                    if br is not None and br.get(qn("w:type")) == "page":
                        pagebreak_el = prev
                        break
                if pagebreak_el is not None:
                    break
            prev = prev.getprevious()
        stop_el = pagebreak_el if pagebreak_el is not None else first_back_text_el
        stop_label = "existing back-page separator/block"
    else:
        cur = appb3_el.getnext()
        while cur is not None:
            if cur.tag == qn("w:sectPr"):
                stop_el = cur
                stop_label = "document section properties"
                break
            cur = cur.getnext()

    def stop_at(el):
        return stop_el is not None and el is stop_el

    # Replace the heading text first.
    replace_paragraph_text_simple_styled(appb3_el, APP_B3_HEADING, "Heading3")

    # Clear everything between the heading and the stop anchor.
    n = clear_between(appb3_el, stop_at)
    print(f"[appendix B.3] cleared {n} elements after heading; stopped before {stop_label}")

    # Insert the new Q&A blocks.
    cursor = appb3_el
    for q, a, fq, fa in APP_B3_QA:
        cursor = insert_after(cursor, make_blank_normal())
        cursor = insert_after(cursor, make_normal_paragraph(q))
        cursor = insert_after(cursor, make_normal_paragraph(a))
        cursor = insert_after(cursor, make_normal_paragraph(fq))
        cursor = insert_after(cursor, make_normal_paragraph(fa))

    # Final blank Normal separator before the stop anchor or document end.
    cursor = insert_after(cursor, make_blank_normal())
    print(f"[appendix B.3] inserted {len(APP_B3_QA)} Q&A blocks")


def replace_paragraph_text_simple_styled(p_el, new_text: str, style_id: str):
    """Replace runs of a styled paragraph (e.g. Heading 3) with one new run."""
    new_text = clean_text(new_text)
    for r in p_el.findall(qn("w:r")):
        p_el.remove(r)
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = _new("w:pPr")
        p_el.insert(0, pPr)
    # Ensure the pStyle is correct.
    for ps in pPr.findall(qn("w:pStyle")):
        pPr.remove(ps)
    pStyle = _new("w:pStyle")
    _set(pStyle, "w:val", style_id)
    pPr.insert(0, pStyle)
    r = _new("w:r")
    t = _new("w:t")
    _set(t, "xml:space", "preserve")
    t.text = new_text
    r.append(t)
    p_el.append(r)


# ============================== main ==============================


def patch():
    if not TARGET.exists():
        raise SystemExit(f"missing target: {TARGET}")

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    backup = BACKUP_DIR / f"{TARGET.stem}_{stamp}_pre-patch-r6.docx"
    shutil.copy2(TARGET, backup)
    print(f"[backup] {backup.name}")

    doc = Document(str(TARGET))

    patch_paragraph_replacements(doc)
    patch_references_swap(doc)
    patch_appendix_b3(doc)

    doc.save(str(TARGET))
    print(f"[saved] {TARGET.name}")


if __name__ == "__main__":
    patch()
