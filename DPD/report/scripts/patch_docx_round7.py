"""
patch_docx_round7.py

Round-7 surgical patch. Five drivers, all upside given the headroom
under the 10-page body cap:

  (a) Section 4 Results: replace the roadmap-only intro with a synthesis
      paragraph that signposts the predict-observe-explain analytical
      method named in project_plan.md and summarizes the headline
      observations from the three interviews before the tables.

  (b) Section 5.2 Platform economics: add (i) one sentence positioning
      FNZ against the consumer sharing-economy typology of Constantiou,
      Marton, and Tuunainen (2017), and (ii) one sentence on data
      network effects from Gregory and colleagues (2021) to sharpen the
      network-effects argument beyond McIntyre and Chintakananda. Also
      specify Iansiti and Lakhani (2020, Ch. 2) where it is currently
      cited generically.

  (c) Section 2 Background: add one sentence positioning FNZ's
      end-to-end model as a systems-of-systems platform per Porter and
      Heppelmann (2015). Also specify Iansiti and Lakhani (2020, Ch. 1)
      for the AI-era framing.

  (d) Section 5.4 Implications: add a societal-dimension paragraph
      covering end-investor welfare under platform concentration, GDPR
      and data-controller responsibilities when client-data lineage
      runs through a third-party platform, and the labour-side
      reorganization of advisor work that Rosenblat and Stark (2016)
      document in algorithmic-management settings. This closes the LO3
      societal-impact gap.

  (e) Iansiti and Lakhani chapter specificity wherever the book is
      currently cited generically (Ch. 1 in section 2, Ch. 2 in section
      5.2, Ch. 4 in sections 5.3 and 5.4).

Reference list additions (placed at correct alphabetical positions):
  - Constantiou, Marton & Tuunainen (2017) - between Constantiou-Joshi-
    Stelmaszak (2023) and Cordella (2006).
  - Gregory, Henfridsson, Kaganer & Kyriakou (2021) - between Gavetti &
    Rivkin (2007) and Iansiti & Lakhani (2020).
  - Porter & Heppelmann (2015) - between McIntyre & Chintakananda
    (2014) and Rosenblat & Stark (2016).
  - Rosenblat & Stark (2016) - between Porter & Heppelmann (2015) and
    Stelmaszak, Joshi & Constantiou (2026).

The patcher follows round-6 conventions: timestamped backup, paragraph-
text rewrites as single runs preserving pPr, em-dashes stripped from
inserted text, no other package parts touched, no back-page handling
(the active final has no decorative back page per CLAUDE.md).
"""
from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
TARGET = ROOT / "KAN-CDSCO2401U_185912_DPD_Spring2026.docx"
BACKUP_DIR = ROOT / "backup"


# ============================== content (paragraph rewrites) ==============================

# Each entry: (unique substring used to locate the paragraph, full new text).
# Paragraphs in target areas are single-run already, so this is safe.
PARAGRAPH_REPLACEMENTS = [
    # ---- Section 2 Background paragraph 1: add Porter & Heppelmann (2015) ----
    (
        "FNZ's position in the industry helps explain why it is analytically useful.",
        "FNZ's position in the industry helps explain why it is analytically "
        "useful. The firm presents itself as an end-to-end wealth-management "
        "platform rather than a point-solution provider. Its public material "
        "emphasizes the consolidation of fragmented systems, the integration "
        "of technology with business and investment operations, and the "
        "ability to support banks, wealth managers, insurers, and asset "
        "managers on a common infrastructure (FNZ, n.d.). In practical "
        "terms, this means that the value proposition extends beyond "
        "software implementation. The platform becomes the operating "
        "backbone through which institutions process advice workflows, "
        "onboarding, portfolio administration, custody arrangements, "
        "reporting, and other core activities. This positioning fits Porter "
        "and Heppelmann's (2015) systems-of-systems logic: a wealth "
        "platform that integrates technology, operations, and regulated "
        "market connectivity becomes a substrate on which institutional "
        "clients construct their own wealth proposition, deepening both "
        "the value created and the dependency that the analysis in section "
        "5.1 traces.",
    ),
    # ---- Section 2 Background paragraph 3: specify Iansiti & Lakhani (2020, Ch. 1) ----
    (
        "The competitive pressure is not only about service speed.",
        "The competitive pressure is not only about service speed. Iansiti "
        "and Lakhani (2020, Ch. 1) characterize the current era as one in "
        "which competitive advantage in financial services accrues to "
        "firms that can convert routine operations into data-driven "
        "learning loops at scale. Tech-led entrants of the Revolut, "
        "Nordnet, and Avanza type are built around such loops from "
        "inception; legacy universal banks are not. Modernizing the "
        "operating backbone is therefore a precondition for competing on "
        "the layers (advisory quality, personalization, data-driven "
        "cross-sell) where AI-era competition is decided. A platform such "
        "as FNZ becomes attractive in this framing because it standardizes "
        "the layers that no longer differentiate, freeing the bank to "
        "invest in the layers that do. The strategic risk is that the bank "
        "standardizes the backbone without making the corresponding "
        "investment, ending up cheaper to run but no more competitive "
        "against tech-led entrants than before. Swedbank's public "
        "rationale illustrates this logic clearly. The bank described its "
        "investment in a new savings platform as part of a strategic "
        "effort to improve efficiency, simplify processes, and create a "
        "stable, standardized, and scalable foundation, while preserving "
        "the bank's direct customer relationship (FNZ, 2021). The same "
        "case also highlighted increased regulatory pressure and the "
        "benefit of sharing platform costs and technological development "
        "across institutions (FNZ, 2021).",
    ),
    # ---- Section 4 Results: rebuild as predict-observe-explain synthesis ----
    (
        "The results section is organized by interview rather than by chronology.",
        "The analytical method follows a predict-observe-explain pattern: "
        "each theoretical lens generates a prediction about the FNZ case, "
        "the public-record evidence and supplementary interview material "
        "show what the case reveals, and any divergence is explained "
        "either by a boundary condition of the theory or by a second "
        "theory that closes the gap. The results section is organized by "
        "interview rather than by chronology, matching the theoretical "
        "sampling logic. Each table summarizes the headline observation "
        "against the analytical purpose; the full transcripts are "
        "reproduced in Appendix B. Read together, the three interviews "
        "converge on three points that section 5 then tests against "
        "theory. First, the bank-FNZ relationship is a hybrid governance "
        "form that the parties run as a partnership rather than as a "
        "pure vendor contract, with day-to-day decisions taken in joint "
        "operating committees rather than through service-level disputes. "
        "Second, the layers banks once differentiated on, custody, "
        "reconciliation, and regulatory reporting, are now table stakes, "
        "with differentiation moving up the stack to advice, brand, "
        "distribution, and the proprietary use of bank-owned data. Third, "
        "lock-in once a bank commits is primarily data-related and "
        "organizational rather than contractual, with multi-year unwinds "
        "on data lineage and team reorganization. Where the FNZ-side and "
        "bank-side perspectives diverge, the bank-side respondent stresses "
        "the regulatory-accountability dimension and the lived "
        "underinvestment in internal data capability as the items boards "
        "consistently underestimate. These observations enter section 5 "
        "where each is tested against the prediction generated by the "
        "corresponding theoretical lens.",
    ),
    # ---- Section 5.2 paragraph 1: Constantiou-Marton-Tuunainen (2017) + Gregory et al. (2021) + Iansiti Ch. 2 ----
    (
        "Multi-sided platform economics explains why the wealth-platform layer concentrates",
        "Multi-sided platform economics explains why the wealth-platform "
        "layer concentrates on a small number of providers rather than "
        "supporting many parallel infrastructures. FNZ is not a classic "
        "two-sided consumer marketplace; it is a regulated B2B platform "
        "whose sides are institutional clients (banks, insurers, wealth "
        "managers) and the regulated market participants on which their "
        "workflows depend (custodians, exchanges, transfer agents, "
        "regulators). It also sits outside the consumer sharing-economy "
        "types that Constantiou, Marton, and Tuunainen (2017) classify by "
        "control and rivalry; FNZ's tight contractual control over "
        "institutional clients combined with low rivalry among them "
        "places it closer to the Principal model in spirit, though the "
        "typology was built for consumer settings and does not transfer "
        "cleanly to a regulated B2B operating backbone. The indirect "
        "network effects are weaker than in consumer platforms but still "
        "material. Each new institutional client raises the marginal "
        "value of FNZ's regulatory connectivity, jurisdictional coverage, "
        "and standardized integrations, because the cost of building "
        "those integrations is incurred once and amortized across the "
        "installed base (Iansiti & Lakhani, 2020, Ch. 2). A second "
        "mechanism is data network effects: as FNZ accumulates AUA and "
        "operating data across institutions, its ability to use that "
        "data to improve regulatory connectivity, fraud detection, and "
        "operational reliability compounds, and Gregory, Henfridsson, "
        "Kaganer, and Kyriakou (2021) argue that this kind of effect can "
        "produce winner-take-most outcomes even where classic two-sided "
        "network effects are weak. Recent client wins illustrate the "
        "cumulative effect: by November 2025, FNZ reported US$2.1 "
        "trillion of assets on platform and a US$650 million capital "
        "raise from existing institutional shareholders (FNZ, 2025a), "
        "and as of April 2026 its public site lists over US$2.4 trillion "
        "in assets on platform and nearly 30 million end investors "
        "(FNZ, n.d.).",
    ),
    # ---- Section 5.3 paragraph 2: specify Iansiti & Lakhani (2020, Ch. 4) ----
    (
        "What remains differentiating shifts toward the layers FNZ cannot replicate",
        "What remains differentiating shifts toward the layers FNZ cannot "
        "replicate: client relationship and trust, advisory quality, "
        "distribution reach, brand, and the proprietary use of bank-owned "
        "client data. Iansiti and Lakhani (2020, Ch. 4) describe the "
        "modern digital firm as built around an AI factory in which data "
        "pipelines and experimentation platforms convert routine "
        "operations into learning loops. For a bank that has outsourced "
        "the operating backbone, the data pipeline is partially co-owned "
        "with FNZ, and the experimentation platform must therefore be "
        "rebuilt at the layers the bank still controls: advisory tooling, "
        "personalization, cross-sell, and client lifecycle management. "
        "Constantiou, Joshi, and Stelmaszak (2023) extend this point by "
        "showing that data assets in platform ecosystems generate value "
        "only when the data-using firm builds the organizational "
        "capability to route, recombine, and act on them; without that "
        "capability, outsourcing the backbone leaves the bank with shared "
        "infrastructure and no proprietary data leverage. Stelmaszak, "
        "Joshi, and Constantiou (2026) develop the same logic for AI "
        "specifically, framing AI capability as an organizing capability "
        "that arises only from the deliberate cultivation of "
        "human-algorithm relations inside the firm, exactly the "
        "cultivation the bank cannot delegate to a platform partner.",
    ),
    # ---- Section 5.4 capability-rebuild paragraph: specify Iansiti & Lakhani (2020, Ch. 4) ----
    (
        "The capability-rebuilding program is the leg that the cognitive analogy most often hides.",
        "The capability-rebuilding program is the leg that the cognitive "
        "analogy most often hides. Banks should fund, from day one, a "
        "deliberate buildout of advisory tooling, data and AI capability "
        "on the layers they still control, and an experimentation "
        "platform on the client-facing edge (Iansiti & Lakhani, 2020, "
        "Ch. 4; Constantiou, Joshi, & Stelmaszak, 2023; Stelmaszak, "
        "Joshi, & Constantiou, 2026). Without this leg, the bank "
        "captures the run-rate cost saving and arrives in five years "
        "with reduced operational headcount, an FNZ-shaped operating "
        "model, and no new differentiating capability, the worst "
        "possible outcome under the dynamic-capabilities lens.",
    ),
]


# ---- Section 5.4 societal paragraph: NEW paragraph inserted before the closing sector-level paragraph ----

SOCIETAL_PARA_NEW = (
    "The recommendation also has societal dimensions that the bank-level "
    "and sector-level lenses partially obscure. Wealth-platform "
    "concentration affects end-investor welfare directly: when a single "
    "platform administers trillions in retail savings across regulated "
    "institutions, an outage or governance failure does not stay inside "
    "the bank-vendor relationship; it reaches the saver. Client-data "
    "lineage that flows through a third-party platform also creates "
    "GDPR and data-protection exposures for which the bank remains the "
    "legal data controller, regardless of where the data physically "
    "resides. Inside the bank, advisor labour reorganizes around "
    "platform-shaped workflows: suitability checks, product-fit logic, "
    "and onboarding journeys move from human discretion to platform "
    "routines, and the information asymmetries that Rosenblat and Stark "
    "(2016) document in algorithmic-management settings reappear in a "
    "regulated wealth context, where the bank both manages its advisors "
    "through platform-shaped tooling and is itself managed by FNZ's "
    "platform constraints. The implication is that the recommendation "
    "cannot be evaluated only on firm-level efficiency. Banks, "
    "regulators, and the European supervisor share an interest in "
    "protecting end-investor outcomes, data-protection rights, and "
    "advisor agency as the wealth-platform layer consolidates."
)

# Anchor used to insert the societal paragraph: it goes BEFORE the
# closing sector-level paragraph "At the sector level, the concentration
# externality identified in §5.2..." in section 5.4.
SOCIETAL_INSERT_ANCHOR_PREFIX = "At the sector level, the concentration externality"


# ============================== content (reference list additions) ==============================

# (anchor_after_prefix, new_text)
NEW_REFERENCE_ENTRIES = [
    (
        "Constantiou, I., Joshi, M., & Stelmaszak, M. (2023).",
        "Constantiou, I., Marton, A., & Tuunainen, V. K. (2017). Four "
        "models of sharing economy platforms. MIS Quarterly Executive, "
        "16(4), 231-251.",
    ),
    (
        "Gavetti, G., & Rivkin, J. W. (2007).",
        "Gregory, R. W., Henfridsson, O., Kaganer, E., & Kyriakou, H. "
        "(2021). The role of artificial intelligence and data network "
        "effects for creating user value. Academy of Management Review, "
        "46(3), 534-551.",
    ),
    (
        "McIntyre, D. P., & Chintakananda, A. (2014).",
        "Porter, M. E., & Heppelmann, J. E. (2015). How smart, connected "
        "products are transforming companies. Harvard Business Review, "
        "October 2015, 96-114.",
    ),
    (
        "Porter, M. E., & Heppelmann, J. E. (2015).",
        "Rosenblat, A., & Stark, L. (2016). Algorithmic labor and "
        "information asymmetries: A case study of Uber's drivers. "
        "International Journal of Communication, 10, 3758-3784.",
    ),
]


# ============================== helpers ==============================

EM_DASH_RE = re.compile(r"\s*[—–]\s*")


def clean_text(text: str) -> str:
    return EM_DASH_RE.sub(", ", text)


def _new(tag):
    return OxmlElement(tag)


def _set(el, attr, val):
    el.set(qn(attr), val)


def _add_size_rpr(rPr, *, with_cs: bool = False):
    sz = _new("w:sz")
    _set(sz, "w:val", "22")
    rPr.append(sz)
    if with_cs:
        szCs = _new("w:szCs")
        _set(szCs, "w:val", "22")
        rPr.append(szCs)


def _make_normal_pPr():
    pPr = _new("w:pPr")
    spacing = _new("w:spacing")
    _set(spacing, "w:line", "240")
    _set(spacing, "w:lineRule", "auto")
    pPr.append(spacing)
    rPr = _new("w:rPr")
    _add_size_rpr(rPr, with_cs=False)
    pPr.append(rPr)
    return pPr


def _make_run(text: str):
    r = _new("w:r")
    rPr = _new("w:rPr")
    _add_size_rpr(rPr, with_cs=True)
    r.append(rPr)
    t = _new("w:t")
    _set(t, "xml:space", "preserve")
    t.text = text
    r.append(t)
    return r


def make_normal_paragraph(text: str):
    text = clean_text(text)
    p = _new("w:p")
    p.append(_make_normal_pPr())
    if text:
        p.append(_make_run(text))
    return p


def make_blank_normal():
    p = _new("w:p")
    p.append(_make_normal_pPr())
    return p


def make_reference_paragraph(text: str):
    text = clean_text(text)
    p = _new("w:p")
    pPr = _new("w:pPr")
    pStyle = _new("w:pStyle")
    _set(pStyle, "w:val", "ReferenceText")
    pPr.append(pStyle)
    spacing = _new("w:spacing")
    _set(spacing, "w:after", "0")
    pPr.append(spacing)
    ind = _new("w:ind")
    _set(ind, "w:right", "0")
    pPr.append(ind)
    p.append(pPr)
    r = _new("w:r")
    t = _new("w:t")
    _set(t, "xml:space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def replace_paragraph_text(p, new_text: str):
    el = p._element
    new_text = clean_text(new_text)
    for r in el.findall(qn("w:r")):
        el.remove(r)
    el.append(_make_run(new_text))


def find_paragraph_containing(doc, substring: str):
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None


def _style_id(el):
    pPr = el.find(qn("w:pPr"))
    if pPr is None:
        return None
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return None
    return pStyle.get(qn("w:val"))


def _para_text(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


# ============================== section patchers ==============================


def patch_paragraph_replacements(doc):
    applied = 0
    skipped = []
    for finder, new_text in PARAGRAPH_REPLACEMENTS:
        p = find_paragraph_containing(doc, finder)
        if p is None:
            skipped.append(finder[:60])
            continue
        replace_paragraph_text(p, new_text)
        applied += 1
        print(f"[edit] {finder[:60]!r}")
    for s in skipped:
        print(f"[edit] WARNING: not found: {s!r}")
    print(f"[edit] total replacements applied: {applied}")


def patch_societal_insert(doc):
    """Insert the LO3 societal paragraph as a new Normal paragraph in section
    5.4, immediately before the closing sector-level paragraph. Pair it with
    a blank Normal separator on each side per Design.md spacing."""
    # Idempotency: skip if a paragraph already starts with the societal opener.
    opener = "The recommendation also has societal dimensions"
    for p in doc.paragraphs:
        if p.text.strip().startswith(opener):
            print("[societal] paragraph already present; skipping")
            return

    target = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(SOCIETAL_INSERT_ANCHOR_PREFIX):
            target = p._element
            break
    if target is None:
        print(f"[societal] WARNING: insert anchor not found: {SOCIETAL_INSERT_ANCHOR_PREFIX!r}")
        return

    # Insert: blank Normal -> societal paragraph -> blank Normal, all BEFORE target.
    new_para = make_normal_paragraph(SOCIETAL_PARA_NEW)
    blank_before = make_blank_normal()
    blank_after = make_blank_normal()

    # addprevious inserts before the target. Insert in order so the final
    # arrangement is: [..., blank_before, societal_paragraph, blank_after, target].
    target.addprevious(blank_before)
    target.addprevious(new_para)
    target.addprevious(blank_after)
    print("[societal] inserted LO3 societal paragraph before sector-level closing paragraph")


def patch_references(doc):
    """Insert new reference entries at correct alphabetical positions."""
    refs_heading_el = None
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Reference Heading":
            if p.text.strip() == "References":
                refs_heading_el = p._element
                break
    if refs_heading_el is None:
        raise SystemExit("could not locate Reference Heading")

    def iter_refs_block():
        cur = refs_heading_el.getnext()
        while cur is not None:
            if cur.tag == qn("w:p") and _style_id(cur) == "Heading1":
                return
            yield cur
            cur = cur.getnext()

    inserted = 0
    for after_prefix, new_text in NEW_REFERENCE_ENTRIES:
        # Idempotency.
        already = False
        for el in iter_refs_block():
            if _style_id(el) == "Reference Text" and _para_text(el).startswith(new_text[:40]):
                already = True
                break
        # ReferenceText style is sometimes stored as styleId "ReferenceText"
        # without space; check both.
        if not already:
            for el in iter_refs_block():
                if _style_id(el) == "ReferenceText" and _para_text(el).startswith(new_text[:40]):
                    already = True
                    break
        if already:
            print(f"[refs] skip (already present): {new_text[:40]!r}")
            continue

        anchor = None
        for el in iter_refs_block():
            sid = _style_id(el)
            if sid in ("ReferenceText", "Reference Text") and _para_text(el).startswith(after_prefix):
                anchor = el
        if anchor is None:
            print(f"[refs] WARNING: anchor not found for {after_prefix!r}; skipping")
            continue

        new_p = make_reference_paragraph(new_text)
        # Insert immediately after anchor's trailing blank Normal so the
        # cadence stays consistent. Walk from anchor to find the next
        # element; if it is a blank Normal, insert after it.
        next_el = anchor.getnext()
        if next_el is not None and next_el.tag == qn("w:p") and not next_el.findall(qn("w:r")):
            insert_anchor = next_el
        else:
            insert_anchor = anchor
        insert_anchor.addnext(new_p)
        # Add trailing blank Normal separator to keep the cadence.
        new_p.addnext(make_blank_normal())
        inserted += 1
        print(f"[refs] inserted: {new_text[:60]!r}")
    print(f"[refs] total inserted: {inserted}")


# ============================== main ==============================


def patch():
    if not TARGET.exists():
        raise SystemExit(f"missing target: {TARGET}")

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    backup = BACKUP_DIR / f"{TARGET.stem}_{stamp}_pre-patch-r7.docx"
    shutil.copy2(TARGET, backup)
    print(f"[backup] {backup.name}")

    doc = Document(str(TARGET))

    patch_paragraph_replacements(doc)
    patch_societal_insert(doc)
    patch_references(doc)

    doc.save(str(TARGET))
    print(f"[saved] {TARGET.name}")


if __name__ == "__main__":
    patch()
