"""
patch_docx_round2.py

Second surgical pass on the root docx. Confirmed scope (1)–(6) from the
prompt:

  1. Strip manual N.N. prefix from Heading 2 paragraphs in section 5
     (5.1, 5.2, 5.3, 5.4). Keep "Appendix A." prefix on the Appendix H2.
  2. Insert Interview D table after Table 3 in section 4 Results, header
     row filled with SEB green (007E40). 2-column layout mirroring
     Tables 1-3, with placeholder Theme rows from project_plan.md section 6.
     Add Appendix B "Interview transcripts" section with B.1, B.2, B.3
     (FNZ interviews) and B.4 (SEB Head of Wealth and Asset Management
     Division), all questions populated, all Answer fields empty.
  3. Force 11 pt on every run inside every paragraph in style "Normal".
     Reference Text paragraphs stay at the template-defined 9.5 pt.
  4. Override w:ind w:right="0" on every Reference Text paragraph so the
     reference list uses the full text-area width (style still carries the
     2835 right indent; we override it on the paragraph).
  5. Zero w:spacing w:before / w:after on every paragraph in styles
     Normal, Heading 1, Heading 2, Heading 3. Reference Text paragraphs
     keep their existing spacing per the user's instruction.
  6. Untouched: cover SDT block, back-page paragraphs (Back - Text),
     Reference Text font and spacing, Heading 1 numbering, captions.

Backs the source file up before saving.
"""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
TARGET = ROOT / "KAN-CDSCO2401U_185912_DPD_Spring2026.docx"
BACKUP_DIR = ROOT / "backup"

SEB_GREEN = "007E40"

# Heading 2 prefix strip targets (exact prefix to remove from the start of text).
H2_PREFIXES_TO_STRIP = ("5.1.", "5.2.", "5.3.", "5.4.")

# ---------- Interview D ----------

INTERVIEW_D_CAPTION = "Table 4. Interview D, bank-side validation interview."

INTERVIEW_D_TABLE = [
    ["Theme", "Observation"],
    ["Framing of the decision", ""],
    ["Retained capabilities", ""],
    ["Lock-in location", ""],
    ["Concentration risk and supervision", ""],
    ["Fintech-pressure salience", ""],
]

# Appendix B transcripts. Mirrors paper_body_draft.md Appendix B layout.
APPENDIX_B_INTRO = (
    "The full transcripts are reproduced below in Question / Answer "
    "format. Questions are presented exactly as they were asked. Answers "
    "are left empty in this draft and will be filled in once the "
    "interview material is locked."
)

INTERVIEW_A_BLOCKS = [
    ("Opening — rapport and context", [
        "How would you describe your seat at FNZ today and what falls inside your remit?",
        "How would you characterize FNZ's positioning in European wealth management today, in your own words rather than the corporate one?",
    ]),
    ("Block 1 — Network effects, multi-homing, and switching", [
        "When a new institutional client joins the platform, what changes, if anything, for the institutions already on it?",
        "In wealth management, banks often use multiple vendors around the edges. Why is the core operating backbone different?",
        "Thinking only about public or anonymizable examples, what would a realistic dual-platform or exit path involve for a large bank, and where do the main obstacles usually appear?",
    ]),
    ("Block 2 — Monetization and incentive alignment", [
        "FNZ often presents the AUA-linked model as part of a partnership rather than a standard vendor relationship. What does that pricing structure enable that a fixed-fee model would not?",
        "How, if at all, does that alignment shape day-to-day prioritization decisions?",
        "Where does that model create tension or misunderstanding for clients?",
    ]),
    ("Block 3 — Winner-take-most conditions in regulated B2B", [
        "What combination of factors makes European wealth-platform infrastructure consolidate onto a small number of providers rather than support many parallel platforms?",
        "Which matters more in practice: switching difficulty, the benefits of scale, or clients' limited appetite for operational variety?",
        "Looking only at public examples such as Swedbank, UniCredit Germany, or Raymond James Canada, what do those mandates reveal about what clients are buying from FNZ today?",
    ]),
    ("Closing — wind-down", [
        "If the European wealth-management sector continues on its current trajectory, what does the platform landscape look like in five years?",
        "Is there anything I have not asked about that you think is important for understanding FNZ's position in European wealth?",
    ]),
]

INTERVIEW_B_BLOCKS = [
    ("Opening — rapport and context", [
        "How would you describe your seat at FNZ today, and how has your time at the firm shaped how you read the bank-FNZ relationship?",
        "From your seat, where is the boundary between what the bank does and what FNZ does on a typical engagement?",
    ]),
    ("Block 1 — Coordination cost versus production cost", [
        "Which costs has technology and platform standardization actually reduced for the banks you work with, and which costs have stayed the same or increased?",
        "Public client announcements often emphasize standardization, governance, and scalability. In operational terms, what has to be true in practice for a bank to experience those benefits?",
        "Over your eleven years, how has that cost picture evolved? Where did the shift surprise you?",
    ]),
    ("Block 2 — Hybrid governance form", [
        "The bank-FNZ relationship is more involved than a simple vendor contract, but it is not internalized either. How would you describe that governance form in practice?",
        "Public cases such as Swedbank suggest banks still evaluate several providers before choosing one primary backbone partner. Why does the operating model so often settle on one core partner rather than several?",
        "Where do you see the line of authority on day-to-day operational decisions, and where does it shift over time?",
    ]),
    ("Block 3 — Asset specificity and lock-in", [
        "What investments does a bank make when it commits to the FNZ platform that become difficult to redeploy elsewhere?",
        "Is the lock-in primarily contractual, technical, operational, or organizational?",
        "In your experience, what do banks consistently underestimate when they enter a wealth-platform outsourcing arrangement?",
    ]),
    ("Closing — wind-down", [
        "If you were advising a universal bank today on which functions to keep and which to outsource to a platform like FNZ, where would you draw the line?",
        "Is there anything I have not asked about that you think is important for understanding the operational reality of this relationship?",
    ]),
]

INTERVIEW_C_BLOCKS = [
    ("Opening — rapport and context", [
        "How would you describe your seat at FNZ today, and what does this role see that other seats do not?",
        "When a major European bank signs with FNZ, what typically changes inside that bank over the next one to two years?",
    ]),
    ("Block 1 — Capability redefinition", [
        "Which capabilities in wealth management are now table stakes, and which still create meaningful differentiation?",
        "What signals do the better-performing banks pick up that the others miss when they think about what to do internally post-implementation?",
        "If operational excellence is increasingly standardized through a partner platform, where is the new basis of competition?",
    ]),
    ("Block 2 — Capability rebuild", [
        "Once the operational backbone has moved to FNZ, what do banks typically try to rebuild or strengthen on their own side first?",
        "What are the most common reasons they fail?",
        "The banks that have made the most of the relationship: what did they do differently from the others?",
    ]),
    ("Block 3 — Cognitive framing of the outsourcing decision", [
        "When a bank's executive committee discusses an FNZ outsourcing decision, what past transformations or outsourcing decisions do they usually compare it to?",
        "Is this framed as another IT outsourcing, or as something different?",
        "Where does that analogy serve them well, and where does it lead them astray?",
    ]),
    ("Closing — wind-down", [
        "If you were advising the CEO of a mid-sized European universal bank that has just signed with FNZ, what are the three things they need to fix internally in the first eighteen months?",
        "Is there anything I have not asked about that you think is important for understanding the bank-side reality of the FNZ relationship?",
    ]),
]

INTERVIEW_D_BLOCKS = [
    ("Opening — rapport and context", [
        "How would you describe your seat at SEB and what falls inside your remit on the wealth and asset-management side?",
        "When SEB looks at wealth-platform partnerships, where does the conversation typically start in the executive committee?",
    ]),
    ("Block 1 — Framing of the outsourcing decision", [
        "How is the decision to commit to a wealth-platform partner like FNZ framed inside the bank: as IT outsourcing, vendor management, or operating-model change?",
        "When the decision is presented at the executive level, which past transformations or outsourcing waves is it compared to?",
        "Where do those comparisons help, and where do they mislead?",
    ]),
    ("Block 2 — Capabilities to keep, and where lock-in sits", [
        "Once the operational backbone moves to a platform partner, which capabilities does the bank deliberately keep internal, and which does it accept becoming standardized through the partner?",
        "In your experience, where does the real lock-in sit: contractual, technical, operational, or organizational?",
        "Which capability rebuilds typically take longer than the bank expects?",
    ]),
    ("Block 3 — Sector view: concentration risk and competitive pressure", [
        "How is concentration risk in wealth-platform infrastructure understood inside the bank and inside Swedish supervision today?",
        "How much does competitive pressure from tech-led retail-finance players such as Revolut, Nordnet, and Avanza factor into the modernization argument?",
        "What governance protections do you think banks should require from a wealth-platform partner that they currently do not?",
    ]),
    ("Closing — wind-down", [
        "If you were advising a Nordic universal bank that has just signed with a wealth-platform partner, what are the three things they need to fix internally in the first eighteen months?",
        "Is there anything I have not asked about that you think is important for understanding the bank-side reality of these partnerships?",
    ]),
]


# ---------- helpers ----------

def get_or_add_pPr(p_el):
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    return pPr


def make_paragraph_el(style_id: str, text: str = "", italic=False, bold=False,
                      sz_half_pt: int | None = None):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)
    p.append(pPr)
    if text:
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if italic:
            rPr.append(OxmlElement("w:i"))
            rPr.append(OxmlElement("w:iCs"))
        if bold:
            rPr.append(OxmlElement("w:b"))
            rPr.append(OxmlElement("w:bCs"))
        if sz_half_pt is not None:
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(sz_half_pt))
            rPr.append(sz)
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), str(sz_half_pt))
            rPr.append(szCs)
        if len(rPr) > 0:
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
    return p


def set_zero_spacing(p_el):
    pPr = get_or_add_pPr(p_el)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:beforeAutospacing"), "0")
    spacing.set(qn("w:afterAutospacing"), "0")


def force_run_size_22(p_el):
    """Ensure every w:r in this paragraph carries w:sz=22 and w:szCs=22."""
    for r in p_el.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r.insert(0, rPr)
        for tag in ("w:sz", "w:szCs"):
            for el in rPr.findall(qn(tag)):
                rPr.remove(el)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "22")
        rPr.append(szCs)


def set_right_indent_zero(p_el):
    pPr = get_or_add_pPr(p_el)
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:right"), "0")


def set_keep_next(p_el):
    pPr = get_or_add_pPr(p_el)
    for el in pPr.findall(qn("w:keepNext")):
        pPr.remove(el)
    pPr.append(OxmlElement("w:keepNext"))


def set_cant_split(row):
    trPr = row._tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        row._tr.insert(0, trPr)
    for el in trPr.findall(qn("w:cantSplit")):
        trPr.remove(el)
    trPr.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:shd")):
        tcPr.remove(el)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def insert_table(doc, body_el, before_el, rows, header_fill):
    """Create a 2-col table with header_fill on row 0, move it before before_el."""
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if i == 0:
                run.bold = True
                if header_fill:
                    set_cell_shading(cell, header_fill)
            set_zero_spacing(p._element)
            force_run_size_22(p._element)
        set_cant_split(table.rows[i])
        if i < len(rows) - 1:
            for cell in table.rows[i].cells:
                for p in cell.paragraphs:
                    set_keep_next(p._element)
    tbl_el = table._element
    body_el.remove(tbl_el)
    before_el.addprevious(tbl_el)
    return table


def find_caption_paragraph(doc, prefix: str):
    for p in doc.paragraphs:
        if p.text.startswith(prefix):
            return p
    return None


# ---------- main ----------

def patch():
    if not TARGET.exists():
        raise SystemExit(f"missing target: {TARGET}")

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    backup = BACKUP_DIR / f"{TARGET.stem}_{stamp}_pre-patch-r2.docx"
    shutil.copy2(TARGET, backup)
    print(f"[backup] {backup.name}")

    doc = Document(str(TARGET))
    body_el = doc.element.body

    # ---- 1. Strip manual N.N. prefix from H2 in section 5 ----
    stripped = []
    for p in doc.paragraphs:
        if p.style.name != "Heading 2":
            continue
        text = p.text
        for prefix in H2_PREFIXES_TO_STRIP:
            if text.startswith(prefix):
                new_text = text[len(prefix):].lstrip()
                # rewrite single run so the title is clean
                for r in p._element.findall(qn("w:r")):
                    p._element.remove(r)
                r = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = new_text
                t.set(qn("xml:space"), "preserve")
                r.append(t)
                p._element.append(r)
                stripped.append(prefix)
                break
    print(f"[H2] stripped manual prefixes: {stripped}")

    # ---- 2a. Insert Table 4 caption + table after Table 3 ----
    table3_caption = find_caption_paragraph(doc, "Table 3. Interview C")
    if table3_caption is None:
        raise SystemExit("could not find Table 3 caption")
    # find the H1 immediately following Table 3's table block as the anchor
    anchor = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() == "Discussion":
            anchor = p._element
            break
    if anchor is None:
        raise SystemExit("could not find Discussion H1 anchor for Table 4")

    if find_caption_paragraph(doc, "Table 4. Interview D") is None:
        cap = make_paragraph_el("Normal", INTERVIEW_D_CAPTION, italic=True,
                                sz_half_pt=22)
        set_keep_next(cap)
        anchor.addprevious(cap)
        insert_table(doc, body_el, anchor, INTERVIEW_D_TABLE, SEB_GREEN)
        anchor.addprevious(make_paragraph_el("Normal", "", sz_half_pt=22))
        print("[§4] inserted Table 4 (Interview D, SEB green header)")
    else:
        print("[§4] Table 4 already present, skipping")

    # ---- 2b. Add Appendix B with B.1, B.2, B.3, B.4 transcripts ----
    # anchor = first back-page paragraph (carries pageBreakBefore in template).
    back_anchor = None
    for p in doc.paragraphs:
        if p.style.name == "Back - Text":
            back_anchor = p._element
            break
    if back_anchor is None:
        raise SystemExit("could not find back-page anchor")

    # Skip if Appendix B already present.
    appendix_b_exists = any(
        p.style.name == "Heading 2" and p.text.startswith("Appendix B.")
        for p in doc.paragraphs
    )
    if not appendix_b_exists:
        # Heading 2 "Appendix B. Interview transcripts"
        h2 = make_paragraph_el("Heading2", "Appendix B. Interview transcripts")
        back_anchor.addprevious(h2)
        # intro paragraph (Normal, 11 pt)
        intro = make_paragraph_el("Normal", APPENDIX_B_INTRO, sz_half_pt=22)
        back_anchor.addprevious(intro)
        # blank
        back_anchor.addprevious(make_paragraph_el("Normal", "", sz_half_pt=22))

        def write_block(persona_h3: str, blocks: list[tuple[str, list[str]]]):
            # Heading 3 for the persona
            back_anchor.addprevious(make_paragraph_el("Heading3", persona_h3))
            for block_label, questions in blocks:
                # bold Normal block label
                back_anchor.addprevious(
                    make_paragraph_el("Normal", block_label, bold=True,
                                      sz_half_pt=22)
                )
                # blank between block label and first Q
                back_anchor.addprevious(
                    make_paragraph_el("Normal", "", sz_half_pt=22)
                )
                for i, q in enumerate(questions):
                    back_anchor.addprevious(
                        make_paragraph_el("Normal", "Question: " + q,
                                          sz_half_pt=22)
                    )
                    back_anchor.addprevious(
                        make_paragraph_el("Normal", "Answer:", sz_half_pt=22)
                    )
                    if i < len(questions) - 1:
                        back_anchor.addprevious(
                            make_paragraph_el("Normal", "", sz_half_pt=22)
                        )
                # blank after the last Q&A of the block
                back_anchor.addprevious(
                    make_paragraph_el("Normal", "", sz_half_pt=22)
                )

        write_block("Appendix B.1. Interview A — Group Head of Europe",
                    INTERVIEW_A_BLOCKS)
        write_block("Appendix B.2. Interview B — Group Head of Operations",
                    INTERVIEW_B_BLOCKS)
        write_block(
            "Appendix B.3. Interview C — Managing Director, Client "
            "Management and Business Development",
            INTERVIEW_C_BLOCKS,
        )
        write_block(
            "Appendix B.4. Interview D — SEB Head of Wealth and Asset "
            "Management Division (bank-side validation)",
            INTERVIEW_D_BLOCKS,
        )
        print("[Appendix B] inserted B.1, B.2, B.3, B.4 transcripts")
    else:
        print("[Appendix B] already present, skipping")

    # ---- 3. Force 11 pt on every Normal paragraph run (excl. Reference Text) ----
    n_normal = 0
    for p in doc.paragraphs:
        if p.style.name == "Normal":
            force_run_size_22(p._element)
            n_normal += 1
    print(f"[font] forced 11 pt on {n_normal} Normal paragraphs")

    # ---- 4. Reference Text: full-width (right indent 0); leave font + spacing ----
    n_ref = 0
    for p in doc.paragraphs:
        if p.style.name == "Reference Text":
            set_right_indent_zero(p._element)
            n_ref += 1
    print(f"[refs] zeroed right indent on {n_ref} Reference Text paragraphs")

    # ---- 5. Zero spacing on Normal + Heading 1/2/3 (NOT RefText, NOT Cover/Back) ----
    spacing_styles = {"Normal", "Heading 1", "Heading 2", "Heading 3"}
    n_spacing = 0
    for p in doc.paragraphs:
        if p.style.name in spacing_styles:
            set_zero_spacing(p._element)
            n_spacing += 1
    # also zero spacing inside table cells
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_zero_spacing(p._element)
                    force_run_size_22(p._element)
    print(f"[spacing] zeroed before/after on {n_spacing} body+heading paragraphs")

    # ---- save ----
    doc.save(str(TARGET))
    print(f"[saved] {TARGET.name}")


if __name__ == "__main__":
    patch()
