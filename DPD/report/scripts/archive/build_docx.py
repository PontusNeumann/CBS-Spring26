"""
build_docx.py

Builds the DPD final-exam docx from the clean CBS template at
report/backup/KAN-CDSCO2401U_185912_DPD_Spring2026_2026-04-26_pre-content-draft.docx
into report/KAN-CDSCO2401U_185912_DPD_Spring2026.docx.

Inserts the body content (Introduction through Conclusion), reference list,
and appendix into the placeholder slots in the template, then applies the
spacing and page-break rules codified in report/Design.md:
  * space-before / space-after = 0 on every paragraph (auto-spacing flags zeroed).
  * Single blank Normal-styled paragraph as the only separator between content
    elements, with the four exceptions (heading->heading, heading->body,
    caption->table, figure->caption) producing no separator.
  * pageBreakBefore on the References heading and Appendix Heading 1 only.
    No body Heading 1 carries a page break.
  * cantSplit on every table row, keepNext on every paragraph in every
    non-last row, keepNext on the caption paragraph immediately above.
  * table header rows are color coded by interview group using FNZ brand
    colors: A purple, B orange, C yellow.

The first-page <w:sdt> block (cover image, CBS logo, accent-colour background)
and the back-page paragraph are not modified.

Idempotent. Safe to re-run.
"""
from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "backup" / "KAN-CDSCO2401U_185912_DPD_Spring2026_2026-04-26_pre-content-draft.docx"
TARGET = ROOT / "KAN-CDSCO2401U_185912_DPD_Spring2026.docx"


# FNZ brand colors from https://brand.fnz.com/visual-identity/color.
FNZ_PURPLE = "6C1BEE"
FNZ_ORANGE = "F37340"
FNZ_YELLOW = "F8D271"
# SEB primary brand green, reserved for any optional bank-side validation
# interview table (Interview D, if added). Used per design rule:
# bank-interview tables get SEB green on the header row.
SEB_GREEN = "007E40"
TABLE_HEADER_COLORS = {
    "table_1": FNZ_PURPLE,      # Interview A
    "appendix_a": FNZ_PURPLE,
    "table_2": FNZ_ORANGE,      # Interview B
    "appendix_b": FNZ_ORANGE,
    "table_3": FNZ_YELLOW,      # Interview C
    "appendix_c": FNZ_YELLOW,
    "table_4": SEB_GREEN,       # optional bank-side validation interview
    "appendix_d": SEB_GREEN,
}


# ---------- Body content ----------

# Each entry is (style, text). Tables are inserted at TABLE markers via a
# parallel TABLES list keyed by marker name.
# Lines are paragraphs in document order. Markers:
#   ("H1", "1. Introduction")              -> Heading 1
#   ("H2", "5.1. Move to the middle")      -> Heading 2
#   ("H3", "...")                          -> Heading 3
#   ("P", "...")                           -> Normal body paragraph
#   ("CAP", "Table 1. ...")                -> Caption (Normal italic, keepNext)
#   ("TABLE", "table_1")                   -> Insert TABLES["table_1"]
#
# Em dashes are not used. Curly quotes are avoided to keep encoding clean.

BODY: list[tuple[str, str]] = [
    ("H1", "1. Introduction"),
    ("P",
     "European wealth management is being reorganized around shared digital "
     "infrastructure. Functions historically built and defended inside the "
     "universal bank are being standardized on external platforms that "
     "combine technology, operations, and regulated market connectivity. The "
     "shift matters because it changes both the economics of delivery and "
     "the basis on which banks compete. A shared platform can lower "
     "duplication, simplify fragmented operating models, and help banks "
     "respond to regulatory and margin pressure. The strategic consequence "
     "is less clear. Once the operational backbone is outsourced, the bank "
     "no longer competes through the same capabilities as before."),
    ("P",
     "This paper examines that shift through the case of FNZ. FNZ is a "
     "global wealth-management platform that partners with banks, insurers, "
     "and wealth managers to provide integrated technology, operations, and "
     "market infrastructure. In August 2025, FNZ announced that it had "
     "reached US$2 trillion in assets on platform and that it worked with "
     "more than 650 financial institutions and served more than 26 million "
     "end investors globally (FNZ, 2025a). The analytical interest lies in "
     "what this scale says about the changing organization of European "
     "wealth management, where incumbent banks are rethinking how much of "
     "the value chain should remain in-house."),
    ("P",
     "Recent public mandates show large institutions choosing FNZ for "
     "operating infrastructure rather than software alone. Swedbank's 2021 "
     "platform investment was framed in terms of efficiency, simplification, "
     "scalability, and stronger governance (FNZ, 2021). UniCredit's 2025 "
     "partnership in Germany emphasized standardization, operational "
     "simplification, and a scalable post-trade platform (FNZ, 2025b). "
     "These decisions suggest that the platform is being adopted as part of "
     "the bank's operating model, not as a peripheral vendor product."),
    ("P",
     "The paper addresses the following research question: why are European "
     "universal banks outsourcing wealth-platform operations to FNZ, and "
     "what does FNZ's end-to-end, assets-linked platform model imply for "
     "bank differentiation and concentration risk in European wealth "
     "management? The underlying business problem is how incumbent banks "
     "can capture the efficiency benefits of FNZ-style outsourcing without "
     "losing strategic differentiation or creating unmanaged concentration "
     "risk. The unifying hypothesis is that FNZ outsourcing is "
     "strategically rational only if banks treat it as a "
     "platform-governance and capability-rebuilding problem, not merely as "
     "a cost-saving IT project."),
    ("P",
     "The primary advisory audience is the incumbent European universal "
     "bank executive committee; the secondary audience is the European "
     "wealth-management sector and regulator. The analysis uses three "
     "theoretical lenses: transaction-cost economics, multi-sided platform "
     "economics, and dynamic capabilities. The argument is that platform "
     "outsourcing solves a real operating problem, but it also shifts "
     "competition away from operational excellence and toward advice, "
     "distribution, client ownership, and the strategic use of data."),

    ("H1", "2. Background"),
    ("P",
     "FNZ presents itself as an end-to-end wealth-management platform "
     "rather than a point-solution provider. Its public material emphasizes "
     "the consolidation of fragmented systems, the integration of "
     "technology with business and investment operations, and the ability "
     "to support banks, wealth managers, insurers, and asset managers on a "
     "common infrastructure (FNZ, n.d.). The value proposition extends "
     "beyond software implementation. The platform becomes the operating "
     "backbone through which institutions process advice workflows, "
     "onboarding, portfolio administration, custody arrangements, and "
     "reporting. Recent acquisitions of New Access and Appway extended this "
     "end-to-end posture, in line with the systems-of-systems logic "
     "described by Porter and Heppelmann (2015)."),
    ("P",
     "The model addresses a structural tension in wealth management. Banks "
     "face growing pressure to modernize digital journeys and reduce "
     "cost-to-serve, in a sector marked by legacy systems, jurisdictional "
     "variation, and high regulatory demands. Tech-led retail-finance "
     "competitors such as Revolut, Nordnet, and Avanza intensify the "
     "pressure by raising customer expectations for speed, cost, and "
     "digital service quality; this makes modernization more urgent, even "
     "though FNZ itself operates as B2B wealth infrastructure rather than "
     "as a direct consumer-fintech rival. A platform partner becomes "
     "attractive when it can absorb a large share of this complexity while "
     "spreading investment across many institutional clients. Swedbank "
     "described its investment in a new savings platform as a strategic "
     "effort to improve efficiency, simplify processes, and create a "
     "stable, standardized, and scalable foundation, while preserving the "
     "bank's direct customer relationship (FNZ, 2021). The same case "
     "highlighted increased regulatory pressure and the benefit of sharing "
     "platform costs and technological development across institutions."),
    ("P",
     "Recent mandates suggest that this logic has strengthened rather than "
     "weakened. In September 2025, UniCredit selected FNZ to transform its "
     "securities-services operations in Germany, citing standardization, "
     "efficiency, and a scalable, cloud-based post-trade platform (FNZ, "
     "2025b). In June 2025, Raymond James Ltd. announced a strategic "
     "partnership with FNZ for a next-generation wealth-management "
     "platform, framed around end-to-end infrastructure, "
     "straight-through processing, and digital capability (FNZ, 2025c). "
     "The cases differ in geography and business model, but they point in "
     "the same direction. Institutions are no longer buying isolated "
     "digital tools. They are redesigning the infrastructure on which the "
     "wealth proposition rests."),
    ("P",
     "That shift creates the core tension of the paper. If the platform "
     "solves complexity, lowers duplication, and offers scale economies, "
     "outsourcing is economically persuasive. However, if several banks "
     "adopt the same operational backbone, traditional sources of "
     "competitive differentiation become weaker. The strategic problem is "
     "not whether outsourcing improves operating efficiency. It is whether "
     "banks can continue to differentiate once more of the underlying "
     "process architecture is shared. The FNZ case is therefore not only a "
     "case about efficiency, but also about boundary choice, platform "
     "concentration, and the changing location of competitive advantage."),

    ("H1", "3. Method"),
    ("P",
     "The paper applies three theoretical lenses in depth. The first is "
     "transaction-cost economics, limited to the coordination-cost versus "
     "production-cost distinction (Cordella, 2006), the Move to the Middle "
     "hypothesis, and asset specificity (Clemons, Reddi, and Row, 1993). "
     "The second is multi-sided platform economics, limited to indirect "
     "network effects (Iansiti and Lakhani, 2020), switching and "
     "multi-homing costs, and the three winner-take-most conditions of "
     "McIntyre and Chintakananda (2014). The third is dynamic "
     "capabilities, limited to sensing, seizing, and reconfiguring within "
     "the digital-enactment framing of Constantiou, Joshi, and Stelmaszak "
     "(2023), supplemented by Gavetti and Rivkin's (2007) account of "
     "cognitive search and analogical reasoning. The information-asymmetry "
     "components of Clemons et al. (1993), the marketing-strategy "
     "implications of McIntyre and Chintakananda (2014), the AI Factory "
     "framework, and the broader enactment-systems typology are out of "
     "scope."),
    ("P",
     "The empirical strategy combines public-record case material with "
     "three semi-structured elite interviews at FNZ, conducted under "
     "company confidentiality terms and anonymized by request. The "
     "interview design follows Bell, Harley, and Bryman (2022): a shared "
     "opening-block-block-block-closing scaffold, open-ended main "
     "questions in plain business language, and two flexible probes per "
     "block. The sampling logic is theoretical rather than statistical. "
     "Each interviewee occupies a different vantage on the bank-FNZ "
     "relationship and therefore informs one analytical block more "
     "strongly than the others. Interview A, with the Group Head of "
     "Europe, anchors the platform-economics analysis. Interview B, with "
     "the Group Head of Operations, anchors the transaction-cost "
     "analysis. Interview C, with the Managing Director for Client "
     "Management and Business Development, anchors the "
     "dynamic-capabilities analysis. The aim is analytical depth and "
     "conceptual fit rather than statistical representativeness."),
    ("P",
     "The scope is intentionally narrow. FNZ is used as a focused case "
     "through which a broader European industry shift is analyzed. The "
     "report does not map the entire global wealth-platform market, nor "
     "does it evaluate every component of FNZ's product portfolio. It is "
     "concerned specifically with the outsourcing of the wealth-management "
     "operating backbone and with the competitive and governance "
     "consequences of that choice. The analysis concentrates on the "
     "European universal-bank context, using public mandates outside "
     "Europe only when they sharpen the mechanism under study. The "
     "analytical method is predict, observe, explain divergence: each "
     "theoretical move begins with what the theory predicts, then states "
     "what the case reveals, and closes by naming the missing mechanism "
     "or boundary condition when prediction and observation diverge."),

    ("H1", "4. Results"),
    ("P",
     "The interview material is presented by interviewee, matching the "
     "theoretical sampling logic. Each table reports the most analytically "
     "useful observations from a single interview; lower-value rows are "
     "retained in the appendix. Quotations are paraphrased."),
    ("CAP", "Table 1. Interview A, Group Head of Europe."),
    ("TABLE", "table_1"),
    ("CAP", "Table 2. Interview B, Group Head of Operations."),
    ("TABLE", "table_2"),
    ("CAP", "Table 3. Interview C, Managing Director, Client Management and Business Development."),
    ("TABLE", "table_3"),

    ("H1", "5. Discussion"),
    ("H2", "5.1. Transaction costs and the move to the middle"),
    ("P",
     "Transaction-cost economics predicts that, when IT lowers coordination "
     "costs faster than production costs, firms shift their boundaries away "
     "from full hierarchy without moving to spot markets, and instead form "
     "long-term relationships with a small number of specialized partners "
     "(Clemons et al., 1993; Cordella, 2006). The FNZ case fits the "
     "prediction in form but not in count. Coordination costs fall sharply "
     "where the platform absorbs cross-jurisdictional reconciliation, "
     "regulatory reporting, and the post-trade workflow that universal "
     "banks previously coordinated across multiple internal teams. "
     "Production costs fall more slowly because product complexity, "
     "taxation, and supervisory variation persist (FNZ Interview B, 2026). "
     "The hybrid governance form is plainly visible in the steering "
     "committees, shared roadmap inputs, and joint operating procedures "
     "that surround the engagement."),
    ("P",
     "The deviation from the Move-to-the-Middle prediction is in the "
     "number of partners. Theory expects banks to spread the operational "
     "backbone across several specialized providers; observation shows "
     "convergence on one primary partner. The mechanism is asset "
     "specificity. Lock-in is mainly operational and organizational rather "
     "than contractual or technical, because the bank's investment in "
     "process redesign, data-taxonomy alignment, and people reorganization "
     "is hard to redeploy elsewhere (FNZ Interview B, 2026). Once that "
     "specificity accumulates, the cost case for parallel operation "
     "collapses. The implication is that governance protections on a "
     "wealth-platform contract should target the operational layer, not "
     "just the legal layer. Data portability commitments, exit playbooks "
     "anchored in shared procedures, and the right to participate in "
     "roadmap decisions matter more than termination clauses on their own."),

    ("H2", "5.2. Platform economics and concentration"),
    ("P",
     "Multi-sided platform theory predicts winner-take-most outcomes when "
     "network effects are strong, multi-homing is costly, and demand for "
     "variety is low (McIntyre and Chintakananda, 2014). In regulated B2B "
     "wealth, the three conditions re-read unevenly. Network effects are "
     "present but indirect: each new institution on FNZ enlarges the "
     "regulated coverage, asset-class breadth, and reporting depth "
     "available to every other institution (Iansiti and Lakhani, 2020), "
     "without the user-to-user benefit that marks B2C platforms. "
     "Multi-homing costs are decisive: duplicating a regulated operational "
     "backbone is uneconomic, and supervisors expect operational "
     "consistency. Demand for variety is structurally low, since banks "
     "want regulatory uniformity and reconciliation simplicity in the "
     "back office, not feature diversity (FNZ Interview A, 2026)."),
    ("P",
     "Two of the three conditions are met decisively, and the third meets "
     "the threshold partially. The case is therefore not a textbook "
     "winner-take-all market, but it is a winner-take-most market where "
     "operational concentration is structurally self-reinforcing. The "
     "assets-linked monetization compounds the dynamic, because revenue "
     "scales with client growth and reinvestment scales with revenue, so "
     "the leading platform enlarges its capability lead with each new "
     "mandate. The implication for sector governance is concrete. "
     "Operational-resilience supervision should treat large "
     "wealth-platform providers as systemically relevant infrastructure, "
     "with stress-testing requirements, exit-portability obligations, and "
     "concentration metrics monitored at the sector level. The implication "
     "for banks is that concentration risk is not a question of platform "
     "choice but of how the choice is governed."),

    ("H2", "5.3. Dynamic capabilities and the new basis of differentiation"),
    ("P",
     "Dynamic-capabilities theory predicts that firms protect capabilities "
     "tied to differentiation and externalize those that are not "
     "(Constantiou et al., 2023). The FNZ observation runs against the "
     "prediction at first reading, because banks are externalizing custody "
     "robustness, reporting accuracy, and regulatory throughput, "
     "capabilities that were historically claimed as differentiators. The "
     "divergence is explained by a shift in the basis of competition. Once "
     "a credible platform partner can deliver these capabilities at scale "
     "across many institutions, they cease to differentiate and become "
     "table stakes (FNZ Interview C, 2026). Sensing is the first "
     "capability that matters. Stronger banks recognize the shift earlier "
     "and redeploy talent before the operational franchise is hollowed "
     "out."),
    ("P",
     "Seizing and reconfiguring then dictate what the bank rebuilds "
     "internally. The interview material points to advice quality, "
     "client-relationship depth, brand, distribution, and the data layer "
     "that flows back from FNZ as the new differentiation surface. Banks "
     "that treat the rebuild as a digital project usually fail; banks "
     "that treat it as a value-chain reconfiguration succeed more often "
     "(FNZ Interview C, 2026). Gavetti and Rivkin's (2007) account of "
     "analogical reasoning explains the recurring managerial error. Bank "
     "executive committees typically frame the FNZ decision by analogy "
     "to earlier IT outsourcing, where the strategic consequence was "
     "modest and the question was delivery efficiency. The analogy "
     "travels poorly. The FNZ decision is a boundary choice in the value "
     "chain, not a sourcing decision, and framing it correctly changes "
     "which capabilities are protected, which are rebuilt, and how "
     "aggressively the rebuild is funded."),

    ("H2", "5.4. Implications for banks and sector governance"),
    ("P",
     "For an incumbent universal bank already executing a wealth-platform "
     "outsourcing decision, three implications follow. First, the "
     "standardized operational layers are the ones to outsource, while the "
     "layers that create differentiation are the ones to retain and "
     "strengthen: client relationship, advisory quality, distribution, "
     "brand, and proprietary data capabilities. Second, governance "
     "protections should target the operational layer through four "
     "concrete commitments embedded in contract and operating procedure: "
     "data portability, exit planning, roadmap influence, and "
     "resilience oversight. Third, the bank's internal rebuild should be "
     "funded as a strategic capability programme tied to advice, "
     "distribution, and the data layer, and treated as the principal "
     "source of differentiation once the operational backbone is shared."),
    ("P",
     "For the European wealth-management sector, the implication is that "
     "operational-resilience supervision needs to catch up with the "
     "structural concentration that platform economics is producing. "
     "Concentration metrics, stress-testing requirements, and portability "
     "obligations should sit at the sector level rather than be left to "
     "bilateral contracts. The trade-off is real and should be "
     "acknowledged. Standardized infrastructure raises the floor on "
     "operational quality and lowers cost-to-serve, which benefits clients "
     "and the system. It also shifts the basis of competition and "
     "concentrates operational risk on a small number of providers. The "
     "governance response should preserve the efficiency gains while "
     "bounding the systemic exposure they create."),

    ("H1", "6. Conclusion and limitations"),
    ("P",
     "The paper asked why European universal banks are outsourcing "
     "wealth-platform operations to FNZ and what the assets-linked "
     "platform model implies for differentiation and concentration risk. "
     "The answer is threefold. Banks outsource because transaction-cost "
     "economics makes the move efficient, with coordination costs falling "
     "faster than production costs and asset specificity steering the "
     "relationship into hybrid governance with a single primary partner. "
     "Concentration follows because two of the three winner-take-most "
     "conditions are met decisively in regulated B2B wealth, and the "
     "assets-linked monetization compounds the lead of the leading "
     "platform. Differentiation must be rebuilt around advice, "
     "distribution, brand, and the data layer, because operational "
     "excellence has shifted from differentiator to table stakes. The "
     "sector-level consequence is structural concentration risk that "
     "warrants supervision at the platform level rather than at the "
     "contract level. The empirical base is small in N, drawn from a "
     "single firm, and anonymized at source. Triangulation comes from "
     "public-record corporate evidence and syllabus literature, and the "
     "predict-observe-explain method tests prediction against observation "
     "rather than relying on the interview material as confirmation."),
]


# ---------- Tables ----------

TABLES: dict[str, list[list[str]]] = {
    "table_1": [
        ["Theme", "Observation"],
        ["Indirect network effects",
         "Each new institutional client improves the platform's regulatory, "
         "reporting, and asset-coverage capability for every other client, "
         "because shared investment is amortized across the institutional "
         "base. The benefit is operational rather than user-to-user."],
        ["Multi-homing costs",
         "Banks rarely run two operational backbones in parallel. The cost "
         "of duplicating regulated post-trade operations, and the "
         "consistency required by supervisors, make sustained dual-platform "
         "setups uneconomic."],
        ["Assets-linked monetization",
         "The pricing model aligns FNZ with client growth and frames the "
         "engagement as a partnership rather than a vendor contract. "
         "Tension arises when fee escalation tracks book growth that the "
         "bank attributes to its own franchise."],
        ["Winner-take-most conditions",
         "In regulated B2B wealth, the three conditions of McIntyre and "
         "Chintakananda (2014) re-read unevenly. Multi-homing cost is "
         "dominant. Demand for variety is low. The strength of network "
         "effects is moderate and indirect rather than user-to-user "
         "(FNZ Interview A, 2026)."],
    ],
    "table_2": [
        ["Theme", "Observation"],
        ["Coordination versus production cost",
         "Coordination costs fall fastest where the platform absorbs "
         "cross-jurisdictional reconciliation, regulatory reporting, and "
         "standardized post-trade workflow. Production costs fall more "
         "slowly because product complexity, taxation, and supervisory "
         "variation persist."],
        ["Hybrid governance",
         "The bank-FNZ relationship is neither pure market nor pure "
         "hierarchy. Steering committees, joint roadmap inputs, and shared "
         "operating procedures are routine. Day-to-day authority sits with "
         "FNZ for delivery and with the bank for client-facing decisions, "
         "with the boundary tested as the engagement matures."],
        ["Asset specificity and lock-in",
         "Lock-in is mainly operational and organizational rather than "
         "contractual or technical. Banks underestimate the depth of "
         "process intertwining, the data-taxonomy alignment work, and the "
         "people-side reorganization that follows."],
        ["Move to the middle, in practice",
         "Banks evaluate several providers at first, then consolidate onto "
         "a single primary backbone, because operational variance from "
         "running multiple platforms in parallel destroys the cost case "
         "(FNZ Interview B, 2026)."],
    ],
    "table_3": [
        ["Theme", "Observation"],
        ["Capability redefinition",
         "Custody robustness, reporting accuracy, and regulatory "
         "throughput are now table stakes. Differentiation is shifting "
         "toward advice quality, client-relationship depth, brand, and "
         "the data layer that flows back from the platform."],
        ["Capability rebuild",
         "Banks try to rebuild advisor enablement, segmentation, and "
         "client analytics first. They typically fail when they treat the "
         "rebuild as a digital project rather than as a capability "
         "reconfiguration tied to the new value-chain shape."],
        ["Cognitive framing",
         "Bank executive committees frame the FNZ decision by analogy to "
         "earlier IT outsourcing waves. The analogy underplays the "
         "strategic consequence and contributes to under-investment in "
         "the post-implementation rebuild."],
        ["What strong banks do differently",
         "Stronger banks recognize sooner that operational excellence is "
         "no longer a differentiator, redeploy talent into advice and "
         "distribution, and treat the platform partnership as a strategic "
         "boundary choice rather than a delivery contract "
         "(FNZ Interview C, 2026)."],
    ],
}


# ---------- References (APA 7) ----------

REFERENCES: list[str] = [
    "Bell, E., Harley, B., & Bryman, A. (2022). Business research methods "
    "(6th ed.). Oxford University Press.",
    "Clemons, E. K., Reddi, S. P., & Row, M. C. (1993). The impact of "
    "information technology on the organization of economic activity: The "
    "\"move to the middle\" hypothesis. Journal of Management Information "
    "Systems, 10(2), 9-35.",
    "Constantiou, I., Joshi, M., & Stelmaszak, M. (2023). Organizations "
    "as digital enactment systems: A theory of replication and innovation "
    "in digital firms. Journal of the Association for Information "
    "Systems, 24(6), 1504-1530.",
    "Cordella, A. (2006). Transaction costs and information systems: Does "
    "IT add up? Journal of Information Technology, 21(3), 195-202.",
    "FNZ. (n.d.). About FNZ. https://www.fnz.com/about",
    "FNZ. (2021). Swedbank announces strategic investment in new savings "
    "platform with FNZ. https://www.fnz.com/news",
    "FNZ. (2025a). FNZ surpasses US$2 trillion in assets on platform. "
    "https://www.fnz.com/news",
    "FNZ. (2025b). UniCredit selects FNZ to transform securities-services "
    "operations in Germany. https://www.fnz.com/news",
    "FNZ. (2025c). Raymond James Ltd. and FNZ announce strategic "
    "partnership for next-generation wealth platform. "
    "https://www.fnz.com/news",
    "FNZ Interview A. (2026). Confidential interview with a senior "
    "executive of FNZ Group, conducted under company confidentiality "
    "terms. Anonymized by request.",
    "FNZ Interview B. (2026). Confidential interview with a senior "
    "operations leader at FNZ Group, conducted under company "
    "confidentiality terms. Anonymized by request.",
    "FNZ Interview C. (2026). Confidential interview with a senior "
    "client-facing director at FNZ Group, conducted under company "
    "confidentiality terms. Anonymized by request.",
    "Gavetti, G., & Rivkin, J. W. (2007). On the origin of strategy: "
    "Action and cognition over time. Organization Science, 18(3), 420-439.",
    "Iansiti, M., & Lakhani, K. R. (2020). Competing in the age of AI: "
    "Strategy and leadership when algorithms and networks run the world. "
    "Harvard Business Review Press.",
    "McIntyre, D. P., & Chintakananda, A. (2014). Competing in network "
    "markets: Can the winner take all? Business Horizons, 57(1), 117-125.",
    "Porter, M. E., & Heppelmann, J. E. (2015). How smart, connected "
    "products are transforming companies. Harvard Business Review, "
    "93(10), 96-114.",
]


# ---------- Appendix ----------

APPENDIX: list[tuple[str, str]] = [
    ("H2", "Appendix A. Extended interview extraction tables"),
    ("P",
     "The tables below retain extended rows from the three interviews "
     "that were not used in the main text but that support traceability "
     "between the interview guides and the analysis. They do not count "
     "toward the page or character budget."),
    ("CAP", "Table A.1. Interview A, extended."),
    ("TABLE", "appendix_a"),
    ("CAP", "Table A.2. Interview B, extended."),
    ("TABLE", "appendix_b"),
    ("CAP", "Table A.3. Interview C, extended."),
    ("TABLE", "appendix_c"),
]

APPENDIX_TABLES: dict[str, list[list[str]]] = {
    "appendix_a": [
        ["Block", "Question focus", "Observation"],
        ["Opening", "Strategic framing of FNZ in European wealth",
         "FNZ is positioned as a growth platform rather than a vendor, "
         "with a single end-to-end stack carrying technology, operations, "
         "and regulated market connectivity."],
        ["Block 2 probe", "Tension created by AUA-linked pricing",
         "Tension surfaces when bank-side growth narratives diverge from "
         "what the model recognizes as platform-attributable scale."],
        ["Block 3 probe", "Public mandates as evidence",
         "Swedbank, UniCredit Germany, and Raymond James Canada are read "
         "as evidence that institutions are buying an operating model, "
         "not a software product."],
        ["Closing", "Five-year view of the platform landscape",
         "Continued consolidation onto a small number of platforms is "
         "expected, with regulatory scrutiny rising as the systemic "
         "footprint grows."],
    ],
    "appendix_b": [
        ["Block", "Question focus", "Observation"],
        ["Opening", "Bank-FNZ boundary on a typical engagement",
         "The interface settles on FNZ for delivery and the bank for "
         "client-facing decisions, with a wide governance overlap that "
         "shifts as the engagement matures."],
        ["Block 1 probe", "Time evolution of the cost picture",
         "Coordination savings show up in year one or two; production-cost "
         "savings emerge slowly and require active operating-model work on "
         "the bank side."],
        ["Block 2 probe", "Why the model settles on one core partner",
         "Operational variance from running multiple platforms in parallel "
         "destroys the cost case before it begins to compound."],
        ["Closing", "Where to draw the make-or-buy line",
         "Outsource the operational backbone, keep the layers that touch "
         "advice, brand, and proprietary client data."],
    ],
    "appendix_c": [
        ["Block", "Question focus", "Observation"],
        ["Opening", "Internal change at the bank in years one to two",
         "Talent reallocation, role redesign, and a slower-than-expected "
         "rebuild of the data layer that flows back from the platform."],
        ["Block 2 probe", "Why capability rebuild efforts fail",
         "Treating the rebuild as a digital project rather than as a "
         "capability reconfiguration tied to the new value-chain shape."],
        ["Block 3 probe", "Where the IT-outsourcing analogy misleads",
         "It frames the decision as a delivery question and underplays "
         "the strategic consequence of moving the operational backbone "
         "outside the firm."],
        ["Closing", "First eighteen months of internal fixes",
         "Rebuild advice, segmentation, and the data layer; refund the "
         "advisor proposition; retire processes that the platform now "
         "owns."],
    ],
}


# ---------- Helpers ----------

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _get_or_add_pPr(paragraph_el):
    pPr = paragraph_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph_el.insert(0, pPr)
    return pPr


def set_zero_spacing(paragraph) -> None:
    """Force space-before and space-after to 0 on the paragraph."""
    pPr = _get_or_add_pPr(paragraph._element)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    for key in ("w:beforeAutospacing", "w:afterAutospacing"):
        if spacing.get(qn(key)) is not None:
            spacing.set(qn(key), "0")


def set_page_break_before(paragraph) -> None:
    """Add page-break-before to a paragraph's properties."""
    pPr = _get_or_add_pPr(paragraph._element)
    for el in pPr.findall(qn("w:pageBreakBefore")):
        pPr.remove(el)
    pPr.append(OxmlElement("w:pageBreakBefore"))


def set_keep_next(paragraph) -> None:
    pPr = _get_or_add_pPr(paragraph._element)
    for el in pPr.findall(qn("w:keepNext")):
        pPr.remove(el)
    pPr.append(OxmlElement("w:keepNext"))


def set_cant_split(row) -> None:
    trPr = row._tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        row._tr.insert(0, trPr)
    for el in trPr.findall(qn("w:cantSplit")):
        trPr.remove(el)
    trPr.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell, fill_hex: str) -> None:
    """Apply a solid background fill to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:shd")):
        tcPr.remove(el)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def make_paragraph(doc, style_name: str, text: str):
    """Create a new paragraph element with the given style and text."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_name)
    pPr.append(pStyle)
    p.append(pPr)
    if text:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
    return p


def make_blank_normal(doc):
    return make_paragraph(doc, "Normal", "")


# Map our content tags to docx style IDs (the values written into
# <w:pStyle w:val="...">). The CBS template uses Heading1,
# ReferenceText, etc. (no spaces) as IDs, even though the display names use
# spaces. We must write the ID, not the display name, otherwise Word treats
# the paragraph as Normal.
STYLE_MAP = {
    "H1": "Heading1",
    "H2": "Heading2",
    "H3": "Heading3",
    "P": "Normal",
    "CAP": "Normal",   # caption styled as normal italic via run formatting
}

# Items that should NOT have a blank Normal row inserted before them, given
# the previous item's tag. Keys: previous item's tag. Values: set of next-item
# tags that should sit flush.
NO_GAP_AFTER = {
    "H1": {"H1", "H2", "H3", "P", "CAP"},   # heading -> heading or body: flush
    "H2": {"H2", "H3", "P", "CAP"},
    "H3": {"H3", "P", "CAP"},
    "CAP": {"TABLE"},                       # caption -> table: flush
}


def insert_table(doc, body_el, before_el, rows: list[list[str]],
                 header_fill: str | None = None):
    """Insert a basic 2- or 3-column table before `before_el`.
    Applies a Word grid style and the keep-together rules.
    Returns the inserted table element."""
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if i == 0:
                run.bold = True
                if header_fill:
                    set_cell_shading(cell, header_fill)
            set_zero_spacing(p)
        set_cant_split(table.rows[i])
        if i < len(rows) - 1:
            for cell in table.rows[i].cells:
                for p in cell.paragraphs:
                    set_keep_next(p)
    # The new table is at the end of the body. Move it to the right place.
    tbl_el = table._element
    body_el.remove(tbl_el)
    before_el.addprevious(tbl_el)
    return table


# ---------- Build ----------

def build():
    if not TEMPLATE.exists():
        raise SystemExit(f"missing template: {TEMPLATE}")

    doc = Document(str(TEMPLATE))
    body_el = doc.element.body

    # Anchor 1: optional template TOC heading. The DPD paper does not use a
    # separate table of contents, so remove the template heading and entries.
    paragraphs = list(doc.paragraphs)

    toc_heading_idx = None
    body_anchor_idx = None      # the empty Heading 1 placeholder
    manchet_idx = None
    tekst_macro_idx = None
    references_heading_idx = None
    reference_text_idx = None

    for i, p in enumerate(paragraphs):
        sname = p.style.name
        text = p.text.strip()
        if sname == "TOC Heading" and toc_heading_idx is None:
            toc_heading_idx = i
        elif sname == "Heading 1" and body_anchor_idx is None and not text:
            body_anchor_idx = i
        elif sname == "Manchet" and manchet_idx is None:
            manchet_idx = i
        elif sname == "Reference Heading" and text == "References":
            references_heading_idx = i
        elif sname == "Reference Text" and reference_text_idx is None:
            reference_text_idx = i

    # The "Tekst" MACROBUTTON placeholder sits between Manchet and References.
    if manchet_idx is not None and references_heading_idx is not None:
        for j in range(manchet_idx + 1, references_heading_idx):
            tekst_macro_idx = j  # take the last one before References

    # ------ Step A: remove optional template TOC ------
    # Find the TOC entries: paragraphs with style toc 1 / toc 2 / toc 3
    # immediately after the TOCHeading.
    toc_para_indices = []
    if toc_heading_idx is not None:
        for j in range(toc_heading_idx + 1, len(paragraphs)):
            if paragraphs[j].style.name in ("toc 1", "toc 2", "toc 3"):
                toc_para_indices.append(j)
            else:
                if toc_para_indices:
                    break

    if toc_heading_idx is not None:
        for j in toc_para_indices:
            el = paragraphs[j]._element
            if el.getparent() is not None:
                el.getparent().remove(el)
        toc_heading_el = paragraphs[toc_heading_idx]._element
        if toc_heading_el.getparent() is not None:
            toc_heading_el.getparent().remove(toc_heading_el)

    # Refresh paragraphs list because XML order changed.
    paragraphs = list(doc.paragraphs)

    # ------ Step B: insert body content ------
    # Find the body anchor again and the references heading again.
    body_anchor_el = None
    references_heading_el = None
    manchet_el = None
    tekst_para_el = None

    for p in paragraphs:
        sname = p.style.name
        text = p.text.strip()
        if sname == "Heading 1" and not text and body_anchor_el is None:
            body_anchor_el = p._element
        elif sname == "Manchet" and manchet_el is None:
            manchet_el = p._element
        elif sname == "Reference Heading" and text == "References":
            references_heading_el = p._element
        else:
            if body_anchor_el is not None and references_heading_el is None and manchet_el is not None:
                # paragraph between Manchet and References Heading -> the macro placeholder
                tekst_para_el = p._element

    # Remove the placeholder paragraphs (empty Heading 1, Manchet, Tekst).
    for el in (body_anchor_el, manchet_el, tekst_para_el):
        if el is not None and el.getparent() is not None:
            el.getparent().remove(el)

    # Insert body content right before the References Heading.
    items: list[tuple[str, str]] = list(BODY)

    prev_tag: str | None = None
    for tag, content in items:
        # Insert blank Normal separator unless we are at the start, or this
        # combination falls in the no-gap rules.
        if prev_tag is not None:
            no_gap = tag in NO_GAP_AFTER.get(prev_tag, set())
            if not no_gap:
                blank = make_blank_normal(doc)
                references_heading_el.addprevious(blank)
        if tag == "TABLE":
            insert_table(
                doc,
                body_el,
                references_heading_el,
                TABLES[content],
                TABLE_HEADER_COLORS.get(content),
            )
        else:
            style = STYLE_MAP[tag]
            p_el = make_paragraph(doc, style, content)
            references_heading_el.addprevious(p_el)
            if tag == "CAP":
                # Italicize caption run and apply keepNext.
                for r in p_el.findall(qn("w:r")):
                    rPr = r.find(qn("w:rPr"))
                    if rPr is None:
                        rPr = OxmlElement("w:rPr")
                        r.insert(0, rPr)
                    italic = OxmlElement("w:i")
                    rPr.append(italic)
                    italic2 = OxmlElement("w:iCs")
                    rPr.append(italic2)
                pPr = _get_or_add_pPr(p_el)
                kn = OxmlElement("w:keepNext")
                pPr.append(kn)
        prev_tag = tag

    # ------ Step C: replace References content ------
    # The single empty ReferenceText paragraph after References Heading is
    # replaced by the actual reference list.
    paragraphs = list(doc.paragraphs)
    ref_text_el = None
    after_refs_el = None

    found_refs = False
    for p in paragraphs:
        sname = p.style.name
        text = p.text.strip()
        if sname == "Reference Heading" and text == "References":
            found_refs = True
            continue
        if found_refs and sname == "Reference Text" and ref_text_el is None:
            ref_text_el = p._element

    # Remove the empty Reference Text placeholder.
    if ref_text_el is not None and ref_text_el.getparent() is not None:
        anchor = ref_text_el  # we'll insert before this then remove it
        for ref in REFERENCES:
            new_p = make_paragraph(doc, "ReferenceText", ref)
            anchor.addprevious(new_p)
        ref_text_el.getparent().remove(ref_text_el)

    # ------ Step D: insert Appendix between References list and back page ------
    # The back page is the paragraph with pageBreakBefore in the original
    # template. Find it by scanning for the next paragraph after the last
    # "Reference Text" that has pageBreakBefore.
    paragraphs = list(doc.paragraphs)
    back_page_el = None
    last_ref_text_idx = None
    for i, p in enumerate(paragraphs):
        if p.style.name == "Reference Text":
            last_ref_text_idx = i
    if last_ref_text_idx is not None:
        for j in range(last_ref_text_idx + 1, len(paragraphs)):
            pPr = paragraphs[j]._element.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None:
                back_page_el = paragraphs[j]._element
                break

    if back_page_el is not None:
        # Insert Appendix Heading 1 + content before back_page_el.
        # First the blank Normal separator (handled inside the loop instead).
        appendix_h1 = make_paragraph(doc, "Heading1", "Appendix")
        back_page_el.addprevious(appendix_h1)

        prev_tag = "H1"
        for tag, content in APPENDIX:
            if tag == "TABLE":
                # No gap before table (caption->table), but caption already
                # inserted; we just insert the table.
                insert_table(
                    doc,
                    body_el,
                    back_page_el,
                    APPENDIX_TABLES[content],
                    TABLE_HEADER_COLORS.get(content),
                )
                prev_tag = "TABLE"
                continue
            no_gap = tag in NO_GAP_AFTER.get(prev_tag, set()) or prev_tag is None
            # NO_GAP_AFTER[H1] includes H2 -> flush. So no separator after the
            # Appendix H1 before the H2.
            if not no_gap:
                blank = make_blank_normal(doc)
                back_page_el.addprevious(blank)
            style = STYLE_MAP[tag]
            p_el = make_paragraph(doc, style, content)
            back_page_el.addprevious(p_el)
            if tag == "CAP":
                for r in p_el.findall(qn("w:r")):
                    rPr = r.find(qn("w:rPr"))
                    if rPr is None:
                        rPr = OxmlElement("w:rPr")
                        r.insert(0, rPr)
                    italic = OxmlElement("w:i")
                    rPr.append(italic)
                    italic2 = OxmlElement("w:iCs")
                    rPr.append(italic2)
                pPr = _get_or_add_pPr(p_el)
                pPr.append(OxmlElement("w:keepNext"))
            prev_tag = tag

    # ------ Step E: zero out spacing on every paragraph ------
    n = 0
    for p in doc.paragraphs:
        set_zero_spacing(p)
        n += 1
    print(f"[spacing] forced zero before/after on {n} paragraphs")

    # Also zero spacing on every paragraph inside tables.
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_zero_spacing(p)

    # ------ Step F: page-break-before on References and Appendix ------
    broken = []
    for p in doc.paragraphs:
        sname = p.style.name
        text = p.text.strip()
        if sname == "Reference Heading" and text == "References":
            set_page_break_before(p)
            broken.append("References")
        elif sname == "Heading 1" and text == "Appendix":
            set_page_break_before(p)
            broken.append("Appendix")
    print(f"[page-break] added to: {broken}")

    # Save.
    doc.save(str(TARGET))
    print(f"[saved] {TARGET}")


if __name__ == "__main__":
    build()
