"""
patch_docx_round3.py

Third surgical pass. Fixes three issues caught in round-2 verification:

  1. Body Heading 1 paragraphs (Background, Method, Results, Discussion,
     Conclusion and limitations) currently break to a new page because the
     "Heading 1" style has pageBreakBefore baked in. Override with an
     explicit <w:pageBreakBefore w:val="false"/> on each body H1 so they
     flow inline. Introduction H1 is left alone because it carries the
     cover->body break naturally.

  2. References Heading paragraph has lost its explicit pageBreakBefore.
     Re-add it so References starts on a new page.

  3. Appendix B section is currently inserted between the back-page
     pageBreakBefore Normal paragraph and the first Back-Text paragraph,
     which puts Appendix B inside the back-page block. Move every
     Appendix B paragraph (Heading 2 "Appendix B. ...", intro, all Q&A
     and block-label paragraphs, all blanks) to sit immediately before
     the back-page break Normal so Appendix B is part of the body flow
     and the back page remains on its own page.

  Appendix H1 keeps its style-level pageBreakBefore (no change needed).
  Back-page Normal blank with pageBreakBefore is left untouched.
"""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
TARGET = ROOT / "KAN-CDSCO2401U_185912_DPD_Spring2026.docx"
BACKUP_DIR = ROOT / "backup"

BODY_H1_DISABLE_PBB = {
    "Background",
    "Method",
    "Results",
    "Discussion",
    "Conclusion and limitations",
}


def get_or_add_pPr(p_el):
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    return pPr


def set_page_break_before(p_el, on: bool):
    pPr = get_or_add_pPr(p_el)
    for el in pPr.findall(qn("w:pageBreakBefore")):
        pPr.remove(el)
    pbb = OxmlElement("w:pageBreakBefore")
    if not on:
        pbb.set(qn("w:val"), "false")
    pPr.append(pbb)


def patch():
    if not TARGET.exists():
        raise SystemExit(f"missing target: {TARGET}")

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    backup = BACKUP_DIR / f"{TARGET.stem}_{stamp}_pre-patch-r3.docx"
    shutil.copy2(TARGET, backup)
    print(f"[backup] {backup.name}")

    doc = Document(str(TARGET))
    body_el = doc.element.body

    # ---- 1. Override pbb=false on body H1s ----
    disabled = []
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() in BODY_H1_DISABLE_PBB:
            set_page_break_before(p._element, on=False)
            disabled.append(p.text.strip())
    print(f"[H1] pbb=false on body H1s: {disabled}")

    # ---- 2. Restore explicit pbb on References Heading ----
    for p in doc.paragraphs:
        if p.style.name == "Reference Heading" and p.text.strip() == "References":
            set_page_break_before(p._element, on=True)
            print("[refs] re-added explicit pageBreakBefore on References")
            break

    # ---- 3. Move Appendix B section before back-page-break Normal ----
    # Identify the back-page-break Normal: the Normal paragraph that carries
    # pageBreakBefore and sits immediately before the first Back-Text paragraph.
    backpage_break_el = None
    first_back_text_el = None
    for p in doc.paragraphs:
        if p.style.name == "Back - Text" and first_back_text_el is None:
            first_back_text_el = p._element
            break
    if first_back_text_el is None:
        raise SystemExit("could not find first Back-Text anchor")

    # Walk preceding siblings of first_back_text_el to find the Normal w/ pbb.
    prev = first_back_text_el.getprevious()
    while prev is not None:
        if prev.tag == qn("w:p"):
            pPr = prev.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None:
                backpage_break_el = prev
                break
        prev = prev.getprevious()
    if backpage_break_el is None:
        # Fall back: no separate break paragraph exists.
        # Use the first Back-Text paragraph as the anchor.
        backpage_break_el = first_back_text_el
        print("[appendix B] no separate back-page break; anchoring to first Back-Text")
    else:
        print("[appendix B] back-page break paragraph located")

    # Collect every paragraph between (exclusive) backpage_break_el and
    # first_back_text_el. These are the Appendix B paragraphs that were
    # inserted in the wrong position by round-2.
    misplaced = []
    cursor = backpage_break_el.getnext()
    while cursor is not None and cursor is not first_back_text_el:
        misplaced.append(cursor)
        cursor = cursor.getnext()

    if misplaced:
        print(f"[appendix B] moving {len(misplaced)} misplaced paragraphs before back-page break")
        # Detach in order, re-attach before backpage_break_el preserving order.
        for el in misplaced:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        for el in misplaced:
            backpage_break_el.addprevious(el)
    else:
        print("[appendix B] no paragraphs trapped after the back-page break")

    # ---- save ----
    doc.save(str(TARGET))
    print(f"[saved] {TARGET.name}")


if __name__ == "__main__":
    patch()
