"""
patch_docx_round5.py

FROZEN SNAPSHOT — DO NOT RUN.

This script reflects the 2026-05-03 mid-day state in which the bank-side
validation interview was at SEB Wealth and Asset Management. Later the
same day, after the SEB-vs-Danske discussion, the validation interview
moved to Danske Bank Wealth and Asset Management (Aladdin Wealth peer
mandate), with corresponding rewrites to Appendix B.3, the section 4
Table 3 caption, and the references list. Re-running this script would
revert: (a) Appendix B.3 to the SEB persona, (b) the references list
to "SEB Interview" instead of "Danske Bank Interview", and (c) the
section 4 Table 3 caption to SEB.

The current canonical patcher is scripts/patch_docx_round6.py.

To re-enable for forensic use, comment out the SystemExit guard below.
"""
import sys
sys.exit("patch_docx_round5.py is a frozen snapshot — do not run. Use patch_docx_round6.py.")
"""
Round-5 surgical patch. Aligns the active .docx with the 2026-05-03 decision
to (a) drop the FNZ Strategy interview, (b) keep two FNZ-side interviews
(Operations and Client Management), (c) add one bank-side validation
interview at SEB Wealth and Asset Management division, and (d) anonymize
all interviewees at the level of function/team plus regional remit plus
employer (no titles).

Sections patched:
  1. Section 3 Method - replace the interview paragraph (was: three FNZ;
     now: two FNZ plus one bank-side validation, anonymized).
  2. Section 4 Results - rebuild from the H1 anchor: new intro plus three
     compact summary tables (Question focus / Purpose / Headline observation),
     replacing the four Theme/Detail tables left over from earlier rounds.
  3. Section 6 Conclusion and limitations - update the limitations sentence
     to reflect the two-FNZ-plus-one-bank composition.
  4. References list - insert Bell Harley and Bryman (2022), European Central
     Bank (2025) supervisory newsletter, and three anonymized interview
     entries (FNZ Interview A, FNZ Interview B, SEB Interview), each placed
     at its alphabetical position.
  5. Appendix A captions - relabel the three extended-table captions to the
     new anonymized interview names; the 3x3 placeholder table contents
     remain as-is.
  6. Appendix B - rebuild from the H2 anchor: anonymized intro plus three
     H3-anchored sub-sections (FNZ Operations, FNZ Client Management, SEB
     Wealth and Asset Management) with six main Question and Answer pairs
     plus one Follow-up Q and A per question.

The patcher follows the conventions established by patch_docx_round3.py
and patch_docx_round4.py:
  - Timestamped backup in backup/ before any change.
  - Patches word/document.xml only via python-docx.
  - Inserted body paragraphs use the Normal-paragraph format that matches
    the existing pre-patch body (no pStyle, line=240 lineRule=auto, 11pt
    on the paragraph mark and on every run).
  - Inserted reference entries use the ReferenceText style (size and font
    inherit from style).
  - Inserted H3 sub-section headings use the existing Heading 3 style.
  - Em-dashes are stripped per the in-class instruction.
  - Inter-block spacing follows Design.md 1.2: one blank Normal paragraph
    between content paragraphs; no blank between caption and table; one
    blank between a table and the next caption.

Front page, back page, headers, footers, fonts, accent shapes, cover
image, and all other package parts are not touched.
"""
# `from __future__ import annotations` removed because this script is hard-
# disabled at the top of the file; it never executes past the sys.exit guard.

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
TARGET = ROOT / "KAN-CDSCO2401U_185912_DPD_Spring2026.docx"
BACKUP_DIR = ROOT / "backup"


# ============================== content (sections 3, 6 and references) ==============================

METHOD_INTERVIEW_PARA_NEW = (
    "A supplementary evidence layer of three semi-structured interviews is reported in section 4. "
    "Two are FNZ-side, one in the Group Operations function and one in Client Management and "
    "Business Development for Europe, and the third is a bank-side validation interview with a "
    "senior leader in a Nordic universal bank's wealth and asset-management division. The sampling "
    "logic is theoretical: each respondent informs one analytical block more strongly than the "
    "others, and the bank-side interview triangulates the buyer-side implications rather than "
    "functioning as a fourth analytical pillar. Interview material, when present, deepens the "
    "public-record argument; the analysis in section 5 does not depend on it. All interviewees "
    "are anonymized at source by team or function, regional remit, and employer; full transcripts "
    "are reproduced in Appendix B."
)

LIMITATIONS_PARA_NEW = (
    "The analysis is limited by its single-case design and by its reliance on public-record "
    "material complemented by reflexive case familiarity. The supplementary interviews reported "
    "in section 4 and Appendix B, two FNZ-side and one bank-side, add depth where the public "
    "record is thin, but no analytical claim, theoretical move, or recommendation in section 5 "
    "depends on them. Further research could test the capability-rebuilding argument across a "
    "larger sample of post-FNZ European banks and could quantify the concentration externality "
    "at the sector level using the European supervisor's third-party register."
)

# New reference entries to insert in alphabetical position. Each is a single
# Reference Text paragraph. The patcher places each entry immediately after
# the existing reference whose author/year sorts just before it.
NEW_REFERENCE_ENTRIES = [
    # (author_key_to_search_for, new_entry_text)
    # Bell goes before Clemons.
    (
        "<TOP>",
        "Bell, E., Harley, B., & Bryman, A. (2022). Business research methods (6th ed.). Oxford University Press.",
    ),
    # ECB goes after EBA, before European Parliament.
    (
        "European Banking Authority. (2019).",
        "European Central Bank. (2025, February 19). Outsourcing: Are banks managing third-party "
        "risks adequately? Supervision Newsletter. "
        "https://www.bankingsupervision.europa.eu/press/supervisory-newsletters/newsletter/2025/html/ssm.nl250219_2.en.html",
    ),
    # FNZ Interview A goes after the last FNZ corporate entry (FNZ 2025c).
    (
        "FNZ. (2025c, September).",
        "FNZ Interview A. (2026). Confidential interview with a senior leader within the FNZ "
        "Group Operations function (London-based, multi-client European remit). Anonymized at "
        "source by team or function, regional remit, and employer.",
    ),
    # FNZ Interview B goes immediately after Interview A.
    (
        "FNZ Interview A. (2026).",
        "FNZ Interview B. (2026). Confidential interview with a senior leader within FNZ Group "
        "Client Management and Business Development for Europe (London-based). Anonymized at "
        "source by team or function, regional remit, and employer.",
    ),
    # SEB Interview goes after McIntyre and before Stelmaszak.
    (
        "McIntyre, D. P., & Chintakananda, A. (2014).",
        "SEB Interview. (2026). Confidential bank-side validation interview with a senior leader "
        "within SEB's Wealth and Asset Management division (Stockholm-based). Anonymized at source "
        "by team or function, regional remit, and employer.",
    ),
]


# ============================== content (section 4 results tables) ==============================

RESULTS_INTRO_PARA = (
    "The results section is organized by interview rather than by chronology. This structure "
    "matches the theoretical sampling logic and makes it easier to move from empirical material "
    "to analysis. Each table below summarizes the headline observation against the analytical "
    "purpose; the full transcripts are reproduced in Appendix B."
)

RESULTS_TABLES = [
    {
        "caption": (
            "Table 1. Interview A, FNZ Group Operations function (London-based, multi-client "
            "European remit). Anchored in section 5.1."
        ),
        "headers": ["Question focus", "Purpose in the paper", "Headline observation"],
        "rows": [
            [
                "Bank/FNZ operating boundary",
                "Define the operating interface",
                "Clean perimeter on regulation, advice, and balance-sheet exposure. "
                "Reconciliation breaks emerge as the recurring boundary blur, surfacing in week "
                "three of every go-live.",
            ],
            [
                "Coordination cost vs production cost",
                "Apply Cordella's IT-cost decomposition",
                "Coordination cost falls hard; production cost falls only when the bank "
                "simplifies its own product variants. Swedbank's 2021 cost-sharing logic is "
                "invoked directly.",
            ],
            [
                "Hybrid governance",
                "Characterize the relationship as a hybrid rather than market or hierarchy",
                "Between long-term contracting and a bilateral platform consortium. Test for "
                "healthy partnership: who calls whom first when a regulator publishes a new "
                "technical standard.",
            ],
            [
                "Why one core partner",
                "Explain settlement on one backbone",
                "End-to-end regulatory reconciliation and single-point operational "
                "accountability. Two backbones add a third reconciliation surface that fails "
                "arithmetically.",
            ],
            [
                "Asset specificity and lock-in",
                "Identify lock-in mechanisms",
                "Reframed as dependency. Sharpest in data lineage and in the platform-shaped "
                "retained workforce. Contractual exit is solvable; data-lineage and "
                "organizational exit is not.",
            ],
            [
                "Keep-vs-outsource line (closing)",
                "Bridge into discussion implications",
                "Outsource regulated infrastructure layer; retain client interface, advice, "
                "brand, and proprietary data and analytics. Banks that keep too thin a retained "
                "team get a worse outcome than no outsourcing at all.",
            ],
        ],
    },
    {
        "caption": (
            "Table 2. Interview B, FNZ Group Client Management and Business Development, Europe "
            "(London-based). Anchored in section 5.3."
        ),
        "headers": ["Question focus", "Purpose in the paper", "Headline observation"],
        "rows": [
            [
                "Two-year framing post-signature",
                "Map early post-outsourcing organizational effects",
                "Visible change is the platform; less visible change is organizational "
                "redistribution of governance, prioritization, and compliance ownership. "
                "Plumbing modernizes in months; strategy in years.",
            ],
            [
                "Table stakes vs differentiation",
                "Identify the new basis of competition",
                "Operational excellence, digital onboarding, basic mobile, regulatory accuracy "
                "are now table stakes. Differentiation moved up the stack to advice quality, "
                "brand, distribution, and the data layer.",
            ],
            [
                "New basis of competition",
                "Link outsourcing to strategic repositioning",
                "Three places, in order: advisory quality at scale, distribution and brand, "
                "and the data and personalization layer. Few banks executing the third well "
                "today.",
            ],
            [
                "Capability rebuild and failure modes",
                "Operationalize seizing and reconfiguring",
                "Banks rebuild digital channels and analytics first. Failure: under-resourced "
                "product teams, IT-project framing, unchanged advisor incentives, and "
                "over-broad sequencing.",
            ],
            [
                "Cognitive framing / analogy",
                "Test the prevailing analogy",
                "Defaults to another IT outsourcing. Analogy is the trap because the decision "
                "is harder to reverse and more customer-visible. Better analogy: OEM-to-Tier-1 "
                "supplier relationship.",
            ],
            [
                "CEO advice in the first 18 months (closing)",
                "Direct route into discussion implications",
                "Senior business owner with seniority over procurement; retain more capability "
                "than the program plan suggests; explicitly redefine what the bank wins on.",
            ],
        ],
    },
    {
        "caption": (
            "Table 3. Interview C, SEB Wealth and Asset Management division (Stockholm-based). "
            "Anchored in section 5.2 (concentration) and triangulating section 5.3 (bank-side "
            "capabilities)."
        ),
        "headers": ["Question focus", "Purpose in the paper", "Headline observation"],
        "rows": [
            [
                "Operating-model vs vendor framing",
                "Test whether banks frame the decision as boundary change or as procurement",
                "Both framings co-exist; whichever wins by implementation determines the "
                "post-decision capability investment. Business-owned framing is healthier.",
            ],
            [
                "Retained capabilities",
                "Bound which capabilities the buyer believes must stay internal",
                "Five, ordered: client relationship, advisory quality, brand, regulatory "
                "accountability, internal data capability. Last two systematically "
                "under-invested.",
            ],
            [
                "Lock-in from the buyer side",
                "Refine the lock-in argument",
                "Contractual and technical exit is solvable. Data-related and organizational "
                "lock-in is the concerning combination. Multi-year unwind on data lineage and "
                "team reorganization.",
            ],
            [
                "Concentration and resilience",
                "Stress-test the platform-economics concentration claim",
                "Single-bank governance manages bilateral risk; sector-level concentration "
                "cannot be solved bilaterally. ECB and DORA framing supplies the right policy "
                "direction.",
            ],
            [
                "Why now: fintech vs internal drivers",
                "Test the why-now narrative",
                "Strongest drivers are internal: cost, legacy complexity, regulation. Fintechs "
                "reset the customer-experience floor but do not pay for the modernization "
                "investment.",
            ],
            [
                "Recommendation test (closing)",
                "Stress-test the paper's recommendation",
                "Broadly endorses the recommendation. Adds: regulatory accountability sits "
                "alongside the differentiation list, not after it; owning data is hollow "
                "without analytics talent and advisor workflows that consume it.",
            ],
        ],
    },
]


# ============================== content (appendix b transcripts) ==============================

APP_B_INTRO = (
    "The three interviews were conducted under company confidentiality and have been anonymized "
    "at source. Identifying details have been removed; interviewees are referenced by function, "
    "regional remit, and employer only. Each transcript follows a Question / Answer format with "
    "conditional follow-ups where the original line of questioning developed in real time. The "
    "semi-structured guide design follows Bell, Harley and Bryman (2022): six main questions "
    "per interview, anchored to the focal theory of the corresponding report section, with "
    "probes deployed flexibly."
)

# Each interview is (heading, [(q, a, fup_q, fup_a), ...])
APP_B_INTERVIEWS = [
    (
        "Appendix B.1. Interview A, FNZ Group Operations function (London-based, multi-client "
        "European remit)",
        [
            (
                "Question: From your seat, where is the boundary between what the bank does "
                "and what FNZ does on a typical engagement?",
                "Answer: The clean answer is the obvious one: the client owns the regulatory "
                "perimeter, the client relationship, the advice, and the balance-sheet "
                "exposure. We own the operating layer underneath that, order management, "
                "custody connectivity, corporate-actions processing, the regulatory-reporting "
                "plumbing, the data model. Where it gets less clean is everything that touches "
                "data ownership and customer-facing communication. Some clients want to draft "
                "every customer-facing document; others delegate the templates to us with bank "
                "approval. We adapt. There is no single operating model.",
                "Follow-up Q: Where does that boundary blur most often in practice?",
                "Follow-up A: Reconciliation breaks. The bank assumes we own them because we "
                "run the platform; we assume they own them because the position is on their "
                "book. Every implementation surfaces that conversation in week three of "
                "go-live, without exception.",
            ),
            (
                "Question: Which costs has technology and platform standardization actually "
                "reduced for the banks you work with, and which costs have stayed the same or "
                "increased?",
                "Answer: Coordination cost falls hard, particularly anything that used to "
                "require a person to chase a status across three systems. Manual "
                "reconciliations, regulatory submissions, intra-day position checks, that work "
                "moves into automated workflow. Production cost is more mixed. The unit cost of "
                "a transaction falls because the platform does the same work for many clients "
                "at once. But total cost depends on what the bank brings to the table. If they "
                "keep three legacy product variants because their advisors prefer them, the "
                "platform absorbs that complexity but charges for it. The Swedbank framing in "
                "2021 was honest on this, they spoke about sharing the cost of platforms and "
                "technological development across institutions because the regulatory burden "
                "was making the standalone model uneconomic. That logic has not weakened.",
                "Follow-up Q: Anything that surprised you over the years?",
                "Follow-up A: How long change-the-bank programs take. We can stand up the "
                "platform in twelve to eighteen months. The bank's own product simplification "
                "often takes another two to three years on top of that. The cost saving is "
                "real, but it shows up later than the slide pack promised.",
            ),
            (
                "Question: The bank-FNZ relationship is more involved than a simple vendor "
                "contract, but it is not internalized either. How would you describe that "
                "governance form in practice?",
                "Answer: It is a long-horizon partnership with a contractual frame, but the "
                "day-to-day is run through joint operating committees rather than through "
                "service-level disputes. We share the roadmap, we share incident management, "
                "we share regulatory horizon-scanning. Calling it outsourcing understates it. "
                "Calling it a joint venture overstates it. The honest description is something "
                "between long-term contracting and a bilateral platform consortium. The "
                "cleanest test for whether it is working as a partnership is who picks up the "
                "phone first when a regulator publishes a new technical standard. If it is us "
                "calling them, that is a healthy signal.",
                "Follow-up Q: When does it tilt back toward a vendor relationship?",
                "Follow-up A: When the client's procurement function takes over the "
                "relationship. Procurement frames everything as substitutable. That framing "
                "works for a stationery contract and breaks for an operating backbone. The "
                "relationships that work have a senior business sponsor on the bank side who "
                "keeps procurement in scope rather than in charge.",
            ),
            (
                "Question: Why does the operating model so often settle on one core backbone "
                "partner rather than several?",
                "Answer: Two reasons that people don't always say out loud. First, regulatory "
                "reporting has to reconcile end-to-end. If you split the backbone across two "
                "providers, you create a third reconciliation surface and that surface fails, "
                "almost arithmetically. Second, accountability. The bank's chief operating "
                "officer needs one phone to pick up at three in the morning when something is "
                "broken. Two backbones means two phones and a coordination problem inside the "
                "worst possible window. The peripheral choices around the edges, execution "
                "venues, market-data sources, KYC tools, those can absolutely be multi-vendor. "
                "The core platform almost never is.",
                "Follow-up Q: Have you ever seen a serious dual-backbone attempt?",
                "Follow-up A: Once or twice, usually after a merger. Both unwound within four "
                "years. The integration cost of running two platforms in parallel is genuinely "
                "brutal once you account for the data reconciliation overhead, and the bank's "
                "own people end up exhausted by it.",
            ),
            (
                "Question: What investments does a bank make when it commits to the FNZ "
                "platform that become difficult to redeploy elsewhere, and would you call that "
                "lock-in?",
                "Answer: I would call it dependency, which is more accurate than lock-in but "
                "uncomfortable for the same reasons. The boring answer is data lineage. Once "
                "your reporting flows through a particular platform's data model, recreating "
                "those flows somewhere else is a multi-year project, and the regulator does "
                "not pause while you do it. The less boring answer is the internal team. A "
                "bank that has spent five years building people who understand how to operate "
                "alongside us, change managers, vendor governance leads, product owners, has "
                "a workforce that is genuinely platform-shaped. They can move on, but the "
                "institution's muscle memory is built around our operating model. "
                "Contractually exit is solvable; data-wise and organizationally it bites the "
                "hardest.",
                "Follow-up Q: Of those, which dimension does management usually underestimate "
                "at signature?",
                "Follow-up A: The data lineage. They model the technical migration cost; they "
                "do not model the years of audit history that has to be reproduced on a new "
                "platform's data shape, and the regulator's familiarity with the old shape, "
                "before exit is genuinely complete.",
            ),
            (
                "Question: If you were advising a universal bank today on which functions to "
                "keep and which to outsource to a platform like FNZ, where would you draw the "
                "line?",
                "Answer: Outsource the regulated infrastructure layer, custody, settlement, "
                "post-trade, regulatory-reporting plumbing, the platform itself. Keep the "
                "client interface, the advice, the brand, and the proprietary data and "
                "analytics. The middle layer, onboarding, suitability, complaints, is "
                "genuinely contestable and should be designed with both sides at the table "
                "from the start, not handed over and then re-negotiated when something goes "
                "wrong.",
                "Follow-up Q: One thing banks consistently underestimate when they sign?",
                "Follow-up A: How much capability they need to retain in-house to be a good "
                "client. Outsourcing the work does not outsource the accountability, and a "
                "thin retained team produces a worse outcome than no outsourcing at all. The "
                "banks that get the most out of the relationship invest in their own "
                "platform-management muscle, not less.",
            ),
        ],
    ),
    (
        "Appendix B.2. Interview B, FNZ Group Client Management and Business Development, "
        "Europe (London-based)",
        [
            (
                "Question: When a major European bank signs with FNZ, what typically changes "
                "inside that bank over the next one to two years?",
                "Answer: A surprising amount changes, and not always what the executive "
                "committee predicted at signature. The technology piece is the visible part. "
                "The less visible part is organizational. A wealth COO function suddenly owns "
                "vendor governance at a scale they have not run before. Product teams that "
                "used to debate features now debate prioritization sequences on a shared "
                "roadmap. A compliance function that used to own the process end-to-end "
                "becomes one of three parties in the same conversation. The first twelve "
                "months are mostly about absorbing that. The second year is when the more "
                "interesting strategic conversations start: what do we now do with the "
                "freed-up capacity, what do we build on top of what we no longer have to "
                "operate ourselves.",
                "Follow-up Q: What changes faster than people expect, and what changes slower?",
                "Follow-up A: Faster: the platform itself, the underlying data quality, the "
                "speed of regulatory delivery. Slower: anything that touches advisor "
                "behavior, internal incentives, and the bank's own brand and customer "
                "proposition. The plumbing modernizes in months. The strategy modernizes in "
                "years.",
            ),
            (
                "Question: Which capabilities in wealth management are now table stakes, and "
                "which still create meaningful differentiation?",
                "Answer: Operational excellence in the back office is firmly table stakes. So "
                "is digital onboarding, basic mobile experience, and regulatory-reporting "
                "accuracy. Five years ago each of those could differentiate a wealth "
                "proposition; today they are price of entry. Differentiation has moved up the "
                "stack: the quality and breadth of advice, the relationship the advisor can "
                "sustain at scale, brand trust, distribution reach, and increasingly the "
                "bank's ability to use the data flowing back from the platform to do something "
                "the next bank cannot. That last point is underweighted in most strategy decks "
                "I see.",
                "Follow-up Q: So data is still a differentiator if a bank treats it right?",
                "Follow-up A: It can be. But the bank has to build the analytics capability "
                "and the use cases inside, otherwise the data sits in a warehouse unused and "
                "the differentiator never materializes. You cannot outsource your way to a "
                "data advantage.",
            ),
            (
                "Question: If operational excellence is increasingly standardized through a "
                "partner platform, where is the new basis of competition?",
                "Answer: Three places, in this order. Advisory quality at scale, the ability "
                "to deliver good, personalized advice without the cost structure of a private "
                "bank. Distribution and brand, trust and reach in the home market, which the "
                "regional incumbents still own and the digital challengers do not. And the "
                "data and personalization layer that sits on top of the platform, that is "
                "where the smarter players are starting to invest, although honestly only a "
                "handful are executing it well today. The fintechs reset price and "
                "digital-experience expectations, but they have not yet displaced the "
                "incumbent in advice or trust.",
                "Follow-up Q: Where do digital-only players such as Revolut, Nordnet, or "
                "Avanza enter that picture?",
                "Follow-up A: They reset customer expectations on price, speed, and digital "
                "experience, which forces incumbent banks to modernize the layers we operate. "
                "They are less of a direct competitor for affluent and high-net-worth "
                "segments where advice and trust still drive the relationship, at least for "
                "now. The pressure they create is real but indirect.",
            ),
            (
                "Question: Once the operational backbone has moved to FNZ, what do banks "
                "typically try to rebuild or strengthen on their own side first, and where do "
                "they fail?",
                "Answer: They try to rebuild the digital channel and the analytics layer "
                "first. The failure points are relatively predictable. Under-resourcing the "
                "internal product team. Treating the new capabilities as IT projects rather "
                "than as business capabilities. Not changing the incentive structure for "
                "advisors so that the new tooling actually gets used in customer "
                "conversations. The other failure mode is sequencing: banks that try to "
                "rebuild everything at once spread their best people too thin and end up with "
                "three half-built capabilities instead of one strong one.",
                "Follow-up Q: What do the banks that get the most out of the relationship do "
                "differently?",
                "Follow-up A: They keep a senior business owner on the bank side who treats "
                "us as a long-term partner rather than as a cost center. They invest in their "
                "own internal product muscle from day one. And they are disciplined about "
                "saying we will not try to differentiate on this any more so the focus goes "
                "to the layers where they actually can.",
            ),
            (
                "Question: When a bank's executive committee discusses an FNZ decision, what "
                "past transformations or outsourcing decisions do they compare it to, and is "
                "this framed as another IT outsourcing or as something different?",
                "Answer: Most often, they compare it to a previous IT outsourcing, core "
                "banking modernization, an infrastructure-as-a-service migration, sometimes a "
                "custody outsourcing. The comparison is partly useful: it brings procurement "
                "discipline and a sensible vendor-governance template. But it also encodes a "
                "mental model that under-scopes the problem. A core banking migration is "
                "reversible in principle and largely invisible to the customer. This decision "
                "is harder to reverse and changes what the customer experiences. Calling it "
                "another IT outsourcing lets the steering committee feel familiar, and that "
                "familiarity is the trap.",
                "Follow-up Q: Have you seen anyone use a better analogy?",
                "Follow-up A: A few. Some of the more thoughtful banks frame it as something "
                "close to an automotive OEM-to-Tier-1 supplier relationship, long-horizon, "
                "mutual investment, deliberate retention of design and brand internally. That "
                "is a more accurate frame. It also gets harder push-back from procurement "
                "because it does not fit the standard vendor-management playbook.",
            ),
            (
                "Question: If you were advising the CEO of a mid-sized European universal "
                "bank that has just signed with FNZ, what are the three things they need to "
                "fix internally in the first eighteen months?",
                "Answer: First, put a senior business owner on the relationship and give them "
                "the seniority to override procurement when it matters. Second, keep more "
                "capability in-house than the program plan suggests is necessary, "
                "particularly product ownership, data and analytics, and the advisor-facing "
                "change function. Third, redefine what the bank is actually trying to win on. "
                "If you cannot articulate the answer to why us and not the next bank with the "
                "same backbone inside eighteen months, the cost saving will arrive but the "
                "strategic position will not.",
                "Follow-up Q: Anything you would warn them away from?",
                "Follow-up A: Treating the cost saving as the headline KPI for the program. "
                "Cost-out is a useful starting metric and a terrible sustaining one. The "
                "programs that overcommit to the cost-out narrative end up cutting exactly "
                "the retained capability they need to make the next phase work.",
            ),
        ],
    ),
    (
        "Appendix B.3. Interview C, SEB Wealth and Asset Management division (Stockholm-based)",
        [
            (
                "Question: When a universal bank considers outsourcing wealth-platform "
                "operations to a specialist provider, is that usually understood internally "
                "as an operating-model change or as a more conventional IT/vendor-outsourcing "
                "decision?",
                "Answer: That depends entirely on who in the bank is leading the "
                "conversation. If group IT or procurement runs it, it shows up on the agenda "
                "as an outsourcing decision and gets the standard vendor-management "
                "treatment. If it runs through the COO or the head of wealth, it is framed "
                "much more clearly as an operating-model change. Both framings are present in "
                "most banks at the same time, and the real question is which one wins by the "
                "time you are in implementation. From a governance standpoint, if you are "
                "outsourcing the operating backbone for an entire business line, calling it "
                "vendor management is dangerously narrow. But I also see why the technical "
                "framing is comfortable: it fits the procurement template, the legal "
                "template, and the audit template that already exist.",
                "Follow-up Q: At your bank specifically, would the framing be more "
                "business-driven or IT-driven?",
                "Follow-up A: I can speak only to my own perspective rather than to a formal "
                "bank position. The cultural answer in this house is that anything touching "
                "the customer relationship is owned by the business, so the framing tends to "
                "be operating-model first, vendor management second. That is not universally "
                "true across Nordic banks.",
            ),
            (
                "Question: Which capabilities does the bank believe it must keep or "
                "strengthen internally if the operating backbone is outsourced?",
                "Answer: Five things, in order: the client relationship, advisory quality, "
                "the brand, regulatory accountability, and the internal data capability. The "
                "first three are obvious; you cannot outsource them and stay in business. "
                "Regulatory accountability is the one banks sometimes wrongly assume they can "
                "delegate alongside the operations, you can outsource the work, but the "
                "supervisor will still come to the bank. And the internal data capability is "
                "the one that gets the least attention at decision time and matters the most "
                "three years later. If you do not have people inside who can interrogate the "
                "platform's data, design analytics on it, and build the cases that turn it "
                "into a customer-facing differentiator, you have handed your future to a "
                "vendor. That is not a position any senior bank executive should be "
                "comfortable with.",
                "Follow-up Q: Where do banks typically under-invest among those five?",
                "Follow-up A: The internal data capability. Closely followed by advisor "
                "enablement, the people-and-process work that makes new tooling actually "
                "change customer outcomes.",
            ),
            (
                "Question: From the buyer side, is lock-in mainly contractual, technical, "
                "operational, data-related, or organizational?",
                "Answer: Contractually, you can almost always exit with sufficient notice and "
                "a transition plan. Technically, exit is hard, but it is an engineering "
                "problem you can scope and price. The lock-in that actually concerns me is "
                "the combination of data-related and organizational. Once your data lineage "
                "flows through a particular platform's logic, your reporting suite is tuned "
                "to that data shape, your auditors have learned that shape, and your "
                "supervisor reviews reports in that shape. That is a multi-year unwind. The "
                "organizational side is similar: your operating teams have reorganized "
                "around the platform, your job descriptions have evolved, and rebuilding the "
                "prior internal capability requires a hiring and training cycle that takes "
                "years.",
                "Follow-up Q: Does the board hear that risk in those terms?",
                "Follow-up A: Increasingly, yes. Five years ago the conversation at board "
                "level would have been about cost and resilience. Today the dependency "
                "dimension is on the agenda explicitly, partly because supervisors are "
                "asking pointed questions about substitutability and concentration.",
            ),
            (
                "Question: How does the bank think about concentration risk and operational "
                "resilience when many institutions rely on the same platform infrastructure?",
                "Answer: It is a sector-level conversation, not a single-bank conversation. "
                "Any individual bank can manage its bilateral relationship with a platform "
                "provider through governance, exit planning, and strong operational testing. "
                "What it cannot manage is the situation where many of its peers depend on "
                "the same provider, because then the systemic exposure is shared and so is "
                "the recovery scenario. DORA and the EBA outsourcing guidelines are pushing "
                "the conversation in the right direction, the supervisor wants to see "
                "substitutability assessments, exit planning, and stress testing on critical "
                "third parties, but the underlying market structure is not something any one "
                "bank can fix alone. The recent ECB supervisory commentary on third-party "
                "concentration captures the discomfort well: a meaningful share of the "
                "budget is concentrated on a small number of providers, and the great "
                "majority of critical functions are difficult to reintegrate. Those numbers "
                "should make any board uncomfortable.",
                "Follow-up Q: What would a sensible regulatory intervention look like, in "
                "your view?",
                "Follow-up A: Probably something close to the framework the ECB is already "
                "pointing at: portability standards, data-format requirements, and "
                "resolution-style planning for the largest providers. The market will not "
                "get there on its own.",
            ),
            (
                "Question: Does competitive pressure from tech-forward firms such as Revolut, "
                "Nordnet, and Avanza materially affect the modernization decision, or are "
                "internal cost, regulation, and legacy complexity the stronger drivers?",
                "Answer: Honestly, the strongest drivers are internal: cost, legacy "
                "complexity, and the regulatory burden of running fragmented systems. "
                "Fintech pressure shows up in the deck as the why-now narrative, but it is "
                "rarely the deciding factor. Where it does matter is in the "
                "customer-experience benchmark, once a Nordic customer has used a "
                "Revolut-style mobile experience, their tolerance for a 1990s-era bank app "
                "is gone. So fintech sets the expectation floor, but the modernization "
                "investment is paid for by the cost-and-regulation case, not the "
                "competitive-threat case.",
                "Follow-up Q: Does that change for the affluent or high-net-worth segments?",
                "Follow-up A: Less so. In those segments, advice and trust still anchor the "
                "relationship and the digital-only players are some way from displacing the "
                "incumbent banks. The pressure is real on the mass-affluent and below.",
            ),
            (
                "Question: If a paper recommends that banks outsource standardized "
                "operational layers but retain client relationship, advice, brand, "
                "distribution, and proprietary data capability, what would you challenge or "
                "refine?",
                "Answer: I broadly agree with the direction. The refinement I would add is "
                "that the recommendation undersells the regulatory-accountability "
                "dimension. It should not be a sixth item on the list; it should sit "
                "alongside the others and shape how the contract and the governance are "
                "designed from day one. The other refinement is on data: retain proprietary "
                "data capability is the right phrase, but what banks actually need is the "
                "ability to act on the data in ways that change customer outcomes. Owning "
                "the data is not enough. You need analytics talent, advisor workflows that "
                "consume the analytics output, and a feedback loop that learns from what "
                "worked. Without that, data capability becomes a strategic slogan rather "
                "than a capability.",
                "Follow-up Q: Anything you would add to the recommendation that is not there?",
                "Follow-up A: Two things. First, deliberate exit planning from year one, "
                "not because exit is the goal, but because the option discipline keeps the "
                "relationship balanced. Second, an honest internal conversation about which "
                "historical differentiators the bank is now prepared to give up. The "
                "capability rebuild fails when the bank tries to differentiate on too many "
                "things at once; the discipline is in choosing.",
            ),
        ],
    ),
]


# ============================== content (appendix a captions) ==============================

APP_A_CAPTION_RENAMES = [
    # (old_caption_prefix_to_match, new_caption_text)
    (
        "Table A.1. Interview A, extended.",
        "Table A.1. Interview A, FNZ Group Operations function, extended.",
    ),
    (
        "Table A.2. Interview B, extended.",
        "Table A.2. Interview B, FNZ Group Client Management and Business Development for "
        "Europe, extended.",
    ),
    (
        "Table A.3. Interview C, extended.",
        "Table A.3. Interview C, SEB Wealth and Asset Management division, extended.",
    ),
]


# ============================== helpers ==============================

EM_DASH_RE = re.compile(r"\s*[—–]\s*")  # em-dash, en-dash


def clean_text(text: str) -> str:
    """Strip em-dashes per course rule. Keep regular hyphens."""
    return EM_DASH_RE.sub(", ", text)


def _new(tag):
    return OxmlElement(tag)


def _set(el, attr, val):
    el.set(qn(attr), val)


def _add_size_rpr(rPr, *, with_cs: bool = False):
    sz = _new("w:sz")
    _set(sz, "w:val", "22")
    rPr.append(sz)
    if with_cs:
        szCs = _new("w:szCs")
        _set(szCs, "w:val", "22")
        rPr.append(szCs)


def _make_normal_pPr():
    pPr = _new("w:pPr")
    spacing = _new("w:spacing")
    _set(spacing, "w:line", "240")
    _set(spacing, "w:lineRule", "auto")
    pPr.append(spacing)
    rPr = _new("w:rPr")
    _add_size_rpr(rPr, with_cs=False)
    pPr.append(rPr)
    return pPr


def _make_run(text: str, *, bold: bool = False, italic: bool = False):
    r = _new("w:r")
    rPr = _new("w:rPr")
    if bold:
        rPr.append(_new("w:b"))
        rPr.append(_new("w:bCs"))
    if italic:
        rPr.append(_new("w:i"))
        rPr.append(_new("w:iCs"))
    _add_size_rpr(rPr, with_cs=True)
    r.append(rPr)
    t = _new("w:t")
    _set(t, "xml:space", "preserve")
    t.text = text
    r.append(t)
    return r


def make_normal_paragraph(text: str):
    """Body Normal paragraph with the existing pre-patch format."""
    text = clean_text(text)
    p = _new("w:p")
    p.append(_make_normal_pPr())
    if text:
        p.append(_make_run(text))
    return p


def make_blank_normal():
    """Blank Normal separator paragraph."""
    p = _new("w:p")
    p.append(_make_normal_pPr())
    return p


def make_page_break_paragraph():
    """A Normal paragraph carrying <w:br w:type="page"/> inside its run.
    Used at the end of Appendix B to push the back page onto its own page,
    matching the structure the user manually added after round-5's first run."""
    p = _new("w:p")
    pPr = _new("w:pPr")
    rPr_in_pPr = _new("w:rPr")
    _add_size_rpr(rPr_in_pPr, with_cs=False)
    pPr.append(rPr_in_pPr)
    p.append(pPr)
    r = _new("w:r")
    rPr = _new("w:rPr")
    _add_size_rpr(rPr, with_cs=False)
    r.append(rPr)
    br = _new("w:br")
    _set(br, "w:type", "page")
    r.append(br)
    p.append(r)
    return p


def make_styled_paragraph(style_id: str, text: str):
    """Build a paragraph with an explicit pStyle (e.g. Heading 3)."""
    text = clean_text(text)
    p = _new("w:p")
    pPr = _new("w:pPr")
    pStyle = _new("w:pStyle")
    _set(pStyle, "w:val", style_id)
    pPr.append(pStyle)
    p.append(pPr)
    r = _new("w:r")
    t = _new("w:t")
    _set(t, "xml:space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_reference_paragraph(text: str):
    """ReferenceText-styled paragraph."""
    text = clean_text(text)
    p = _new("w:p")
    pPr = _new("w:pPr")
    pStyle = _new("w:pStyle")
    _set(pStyle, "w:val", "ReferenceText")
    pPr.append(pStyle)
    spacing = _new("w:spacing")
    _set(spacing, "w:after", "0")
    pPr.append(spacing)
    ind = _new("w:ind")
    _set(ind, "w:right", "0")
    pPr.append(ind)
    p.append(pPr)
    r = _new("w:r")
    t = _new("w:t")
    _set(t, "xml:space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_table(doc, headers, rows):
    """Build a Table Grid table at the end of the document and return its
    underlying XML element. Caller is responsible for moving it into place."""
    n_rows = 1 + len(rows)
    n_cols = len(headers)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"
    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        para = hdr_cells[i].paragraphs[0]
        run = para.add_run(clean_text(h))
        run.bold = True
        run.font.size = None  # let cell style decide
    # Body rows
    for r_idx, row_data in enumerate(rows, start=1):
        cells = tbl.rows[r_idx].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ""
            para = cells[c_idx].paragraphs[0]
            para.add_run(clean_text(val))
    return tbl._element


def find_paragraph(doc, *, style_name: str, text_match: str):
    for p in doc.paragraphs:
        if p.style is None or p.style.name != style_name:
            continue
        if p.text.strip() == text_match:
            return p._element
    return None


def find_paragraph_starting_with(doc, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p._element
    return None


def _style_id(el):
    pPr = el.find(qn("w:pPr"))
    if pPr is None:
        return None
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return None
    return pStyle.get(qn("w:val"))


def _para_text(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def make_h1_predicate(*texts):
    target = set(texts)

    def pred(el):
        if el.tag != qn("w:p"):
            return False
        if _style_id(el) != "Heading1":
            return False
        return _para_text(el) in target

    return pred


def make_h2_predicate(*texts):
    target = set(texts)

    def pred(el):
        if el.tag != qn("w:p"):
            return False
        if _style_id(el) != "Heading2":
            return False
        return _para_text(el) in target

    return pred


def make_appendix_or_h1_predicate(*texts):
    """Stop on Heading1 with text in `texts` OR a ReferenceHeading."""
    target = set(texts)

    def pred(el):
        if el.tag != qn("w:p"):
            return False
        sid = _style_id(el)
        if sid == "ReferenceHeading":
            return True
        if sid == "Heading1" and _para_text(el) in target:
            return True
        return False

    return pred


def clear_between(anchor_el, stop_predicate):
    """Remove every <w:p> and <w:tbl> sibling between anchor_el (exclusive)
    and the first sibling matching stop_predicate (exclusive)."""
    parent = anchor_el.getparent()
    victims = []
    cur = anchor_el.getnext()
    while cur is not None and not stop_predicate(cur):
        if cur.tag in (qn("w:p"), qn("w:tbl")):
            victims.append(cur)
        cur = cur.getnext()
    for v in victims:
        parent.remove(v)
    return len(victims)


def insert_after(anchor_el, new_el):
    """Insert new_el immediately after anchor_el. Detach new_el first if it
    has a parent (so add_table elements can be moved)."""
    parent = new_el.getparent()
    if parent is not None:
        parent.remove(new_el)
    anchor_el.addnext(new_el)
    return new_el


# ============================== section patchers ==============================


def patch_method(doc):
    """Replace the interview paragraph in section 3 Method."""
    target_prefix = "A supplementary evidence layer of three semi-structured"
    found = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(target_prefix):
            found = p._element
            break
    if found is None:
        print("[section 3] WARNING: could not locate the interview paragraph in Method")
        return
    parent = found.getparent()
    new_p = make_normal_paragraph(METHOD_INTERVIEW_PARA_NEW)
    parent.replace(found, new_p)
    print("[section 3] interview paragraph replaced")


def patch_results(doc):
    """Rebuild section 4 Results: clear everything between Results H1 and
    Discussion H1, insert intro paragraph and three caption+table blocks."""
    results_h1 = find_paragraph(doc, style_name="Heading 1", text_match="Results")
    if results_h1 is None:
        raise SystemExit("could not locate Results H1")
    n = clear_between(results_h1, make_h1_predicate("Discussion"))
    print(f"[section 4] cleared {n} elements")

    cursor = results_h1
    # Intro paragraph
    intro_p = make_normal_paragraph(RESULTS_INTRO_PARA)
    cursor = insert_after(cursor, intro_p)

    for idx, spec in enumerate(RESULTS_TABLES):
        # blank Normal separator before the next caption (and before the
        # first caption, to match the existing pre-patch spacing).
        cursor = insert_after(cursor, make_blank_normal())
        # caption
        cap_p = make_normal_paragraph(spec["caption"])
        cursor = insert_after(cursor, cap_p)
        # table
        tbl_el = make_table(doc, spec["headers"], spec["rows"])
        cursor = insert_after(cursor, tbl_el)
    print(f"[section 4] inserted intro + {len(RESULTS_TABLES)} caption+table blocks")


def patch_limitations(doc):
    """Replace the limitations paragraph in section 6."""
    target_prefix = "The analysis is limited by its single-case design"
    for p in doc.paragraphs:
        if p.text.strip().startswith(target_prefix):
            new_p = make_normal_paragraph(LIMITATIONS_PARA_NEW)
            p._element.getparent().replace(p._element, new_p)
            print("[section 6] limitations paragraph replaced")
            return
    print("[section 6] WARNING: could not locate the limitations paragraph")


def patch_references(doc):
    """Insert new reference entries at their alphabetical positions."""
    # Find the References Heading.
    refs_heading_el = None
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Reference Heading":
            if p.text.strip() == "References":
                refs_heading_el = p._element
                break
    if refs_heading_el is None:
        raise SystemExit("could not locate Reference Heading")

    # Walk forward from the references heading until we hit a Heading1
    # (Appendix). Reference entries are interleaved with blank Normal
    # separators, so we cannot stop on first non-ReferenceText.
    def iter_refs_block():
        cur = refs_heading_el.getnext()
        while cur is not None:
            if cur.tag == qn("w:p") and _style_id(cur) == "Heading1":
                return
            yield cur
            cur = cur.getnext()

    inserted = 0
    for after_prefix, new_text in NEW_REFERENCE_ENTRIES:
        # Idempotency: skip if entry already present.
        already = False
        for el in iter_refs_block():
            if _style_id(el) == "ReferenceText" and _para_text(el).startswith(new_text[:40]):
                already = True
                break
        if already:
            print(f"[references] skip (already present): {new_text[:40]!r}")
            continue

        if after_prefix == "<TOP>":
            anchor = refs_heading_el
        else:
            anchor = None
            for el in iter_refs_block():
                if _style_id(el) == "ReferenceText" and _para_text(el).startswith(after_prefix):
                    anchor = el
                    # Don't break: keep last match in case of duplicates.
            if anchor is None:
                print(f"[references] WARNING: anchor not found for {after_prefix!r}; skipping")
                continue
        new_p = make_reference_paragraph(new_text)
        anchor.addnext(new_p)
        inserted += 1
        print(f"[references] inserted: {new_text[:60]!r}")
    print(f"[references] total inserted: {inserted}")


def patch_appendix_a_captions(doc):
    """Rename the three Appendix A extended-table captions."""
    renamed = 0
    for old_prefix, new_text in APP_A_CAPTION_RENAMES:
        for p in doc.paragraphs:
            if p.text.strip().startswith(old_prefix):
                # Replace the paragraph with a new one preserving the same
                # formatting (Normal, no italic). Easiest path: clear the
                # runs and insert fresh text.
                el = p._element
                # Remove all existing runs
                for r in el.findall(qn("w:r")):
                    el.remove(r)
                # Add a fresh run
                el.append(_make_run(clean_text(new_text)))
                renamed += 1
                break
    print(f"[appendix A] captions renamed: {renamed}")


def patch_appendix_b(doc):
    """Rebuild Appendix B: clear everything between the Appendix B H2 anchor
    and the back-page break (Normal paragraph with pageBreakBefore preceding
    the first Back - Text paragraph), then insert the new content."""
    # Locate Appendix B H2 anchor.
    app_b_h2 = find_paragraph(
        doc, style_name="Heading 2", text_match="Appendix B. Interview transcripts"
    )
    if app_b_h2 is None:
        raise SystemExit("could not locate Appendix B Heading 2 anchor")

    # The page break to the back page is style-level on Back - Text in this
    # template, not a separate <w:pageBreakBefore> paragraph. So the safe
    # stop anchor is the first Back - Text paragraph itself, and we clear
    # everything between Appendix B H2 (exclusive) and the first Back -
    # Text paragraph (exclusive). Any blank Normal padding paragraphs
    # immediately preceding Back - Text are cleared along with the old
    # Appendix B Q&A content; new spacing is inserted by the rebuild.
    first_back_text_el = None
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Back - Text":
            first_back_text_el = p._element
            break
    if first_back_text_el is None:
        raise SystemExit("could not find first Back-Text anchor")

    def stop_at_back_text(el):
        return el is first_back_text_el

    n = clear_between(app_b_h2, stop_at_back_text)
    print(f"[appendix B] cleared {n} elements")

    # Build the new content.
    cursor = app_b_h2
    # Blank separator after the H2.
    cursor = insert_after(cursor, make_blank_normal())
    # Intro paragraph.
    cursor = insert_after(cursor, make_normal_paragraph(APP_B_INTRO))

    for heading, qa_blocks in APP_B_INTERVIEWS:
        # Blank before each H3.
        cursor = insert_after(cursor, make_blank_normal())
        # H3 heading.
        cursor = insert_after(cursor, make_styled_paragraph("Heading3", heading))
        # Q&A blocks.
        for q, a, fq, fa in qa_blocks:
            cursor = insert_after(cursor, make_blank_normal())
            cursor = insert_after(cursor, make_normal_paragraph(q))
            cursor = insert_after(cursor, make_normal_paragraph(a))
            cursor = insert_after(cursor, make_normal_paragraph(fq))
            cursor = insert_after(cursor, make_normal_paragraph(fa))

    # Trailing structure before the first Back - Text paragraph: one blank
    # Normal, one Normal carrying <w:br w:type="page"/>, one blank Normal.
    # The page-break paragraph is required because the back-page page break
    # in this template is style-level on the first content run of the back
    # page rather than a paragraph-level pageBreakBefore, so without an
    # explicit page break here the back page would flow inline.
    cursor = insert_after(cursor, make_blank_normal())
    cursor = insert_after(cursor, make_page_break_paragraph())
    cursor = insert_after(cursor, make_blank_normal())
    print(f"[appendix B] inserted intro + {len(APP_B_INTERVIEWS)} interviews + back-page break")


# ============================== main ==============================


def patch():
    if not TARGET.exists():
        raise SystemExit(f"missing target: {TARGET}")

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    backup = BACKUP_DIR / f"{TARGET.stem}_{stamp}_pre-patch-r5.docx"
    shutil.copy2(TARGET, backup)
    print(f"[backup] {backup.name}")

    doc = Document(str(TARGET))

    patch_method(doc)
    patch_results(doc)
    patch_limitations(doc)
    patch_references(doc)
    patch_appendix_a_captions(doc)
    patch_appendix_b(doc)

    doc.save(str(TARGET))
    print(f"[saved] {TARGET.name}")


if __name__ == "__main__":
    patch()
